import io
import os
import re
import shutil
import time
import logging
from docx import Document
import json
import pandas as pd
from io import StringIO
import streamlit as st
from datetime import datetime
from docxcompose.composer import Composer
from dotenv import load_dotenv
from langchain_openai import AzureChatOpenAI
from openai import AzureOpenAI

from src.document_generate.dynamic_template_PMF import handle_user_message
from src.document_ingestion.data_collection import data_extraction,convert_all_doc_to_docx_in_folder,extract_text_from_xlsx
from src.document_analyzer.json_converter import extract_text_from_word
from src.document_generate.Assembling_appendix_PMF import assemble_appendices,convert_word_to_pdf
from src.document_analyzer.contents import refresh_toc_with_word,extract_headings_with_tables
from src.document_retriever.Vector_db import DocumentRetriever
from src.eval.eval_config import get_eval_rules
from src.eval.eval_utils import evaluate_run
from src.eval.eval_store import save_eval_run
from src.eval.eval_metrics import compute_all_metrics
from src.eval.eval_judge import PMFJudge
from src.eval.eval_rag import RAGEvaluator
from healthark_eval.suite import EvalSuite



def _derive_section_title(key: str, value: str) -> str:
    """Extract a human-readable section title from a template value.

    Template sections have a type prefix as the key (e.g. "Text: -",
    "Static_text:-") while the actual section heading is the first
    meaningful line of the value (e.g. "1.0 GENERAL INFORMATION:").
    This is used to produce meaningful section keys for the eval dashboard.
    """
    for line in (value or "").split("\n"):
        clean = line.strip().replace("**", "").replace("*", "").strip()
        if clean and len(clean) > 3:
            return clean.rstrip(":").strip()[:70]
    return key  # fallback to the type key if value is empty


# Function to split the template text into two parts based on the section header
def Template_to_list(text):
    if text is None:
        return [], []     # Handle the case where text is None by returning empty lists
    
    sections = [section.strip() for section in text.split('$') if section.strip()]
    split_index = None
    for i, item in enumerate(sections):
        if "DEVICE DESCRIPTION & PRODUCT SPECIFICATION".lower() in item.lower():
            split_index = i
            break

    if split_index is not None:
        part1 = sections[:split_index]
        part2 = sections[split_index:]
    else:
        part1 = sections
        part2 = []

    return part1, part2


# Convert the list of strings into a dictionary with unique keys
def convert_dict(list):

    result_dict = {}
    key_count = {}

    for item in list:
    # Split into key and value at first newline
        parts = item.split('\n', 1)
        if len(parts) == 2:
            key, value = parts[0].strip(), parts[1].strip()
        else:
            # If no newline found, use whole as key and empty value
            key, value = parts[0].strip(), ""

        # Handle duplicate keys by counting
        if key in key_count:
            key_count[key] += 1
            new_key = f"{key}_{key_count[key]}"
        else:
            key_count[key] = 1
            new_key = key

        result_dict[new_key] = value
    return result_dict



def _run_extended_evaluation(run_artifacts, eval_suite):
    """Run extended evaluation: DeepEval RAG Triad + Opik-style + MLflow logging.

    Per section: appends 'extended_eval' (DeepEval results) and
    'opik_eval' (hallucination / answer_relevance / regulatory_tone).
    Document level: saves 'extended_eval_summary' and logs to MLflow.
    """
    from src.eval.eval_opik_style import OpikStyleScorer
    from src.eval.eval_mlflow_tracker import MLflowTracker
    from openai import AzureOpenAI

    logging.info("Running extended evaluation (DeepEval + Opik-style + MLflow)...")

    # Build shared Azure client for Opik scorer
    _azure_key = run_artifacts.get("_azure_key", "")
    _azure_ep = run_artifacts.get("_azure_endpoint", "")
    _azure_ver = run_artifacts.get("_azure_version", "2024-06-01")
    _azure_name = run_artifacts.get("model_name", "gpt-4o")

    opik_scorer = None
    try:
        if _azure_key and _azure_ep:
            _client = AzureOpenAI(api_key=_azure_key, api_version=_azure_ver, azure_endpoint=_azure_ep)
            opik_scorer = OpikStyleScorer(llm_client=_client, model=_azure_name)
    except Exception as exc:
        logging.warning("Opik scorer init failed: %s", exc)

    section_results = []
    _eval_phase_start = time.perf_counter()

    for section in run_artifacts.get("sections", []):
        if section.get("is_static", False):
            continue
        generated = section.get("generated_text", "")
        if not generated.strip():
            continue

        section_key = section.get("section_key", "")
        instruction = section.get("prompt_text", "")
        retrieved_chunks = []
        for p in section.get("retrieved_paths", []):
            try:
                p_lower = str(p).lower()
                if p_lower.endswith(".pdf"):
                    import fitz
                    doc_fitz = fitz.open(p)
                    text = "".join(page.get_text() for page in doc_fitz)[:4000]
                elif p_lower.endswith(".docx"):
                    import docx as _docx_mod
                    d = _docx_mod.Document(p)
                    text = "\n".join(para.text for para in d.paragraphs)[:4000]
                else:
                    with open(p, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()[:4000]
                if text.strip():
                    retrieved_chunks.append(text)
            except Exception:
                pass

        merged_context = "\n\n".join(retrieved_chunks)
        _sec_eval_start = time.perf_counter()

        # ── DeepEval (Judge + RAG Triad) ─────────────────────────────
        try:
            result = eval_suite.run(
                generated=generated,
                retrieved=retrieved_chunks,
                reference="",
                section_key=section_key,
                section_instruction=instruction,
                site_name=run_artifacts.get("site_name", ""),
            )
            section["extended_eval"] = result.to_dict()
            section_results.append(result)
        except Exception as exc:
            logging.warning("DeepEval failed for '%s': %s", section_key, exc)

        # ── Opik-style (Hallucination + Answer Relevance + Regulatory Tone) ──
        if opik_scorer:
            try:
                opik_result = opik_scorer.evaluate_section(
                    section_key=section_key,
                    output=generated,
                    context=merged_context,
                    instruction=instruction,
                )
                section["opik_eval"] = opik_result
            except Exception as exc:
                logging.warning("Opik scoring failed for '%s': %s", section_key, exc)

        # Store eval timing back into the section's timing dict
        _sec_eval_ms = round((time.perf_counter() - _sec_eval_start) * 1000, 1)
        if "timing" not in section or section["timing"] is None:
            section["timing"] = {}
        section["timing"]["eval_ms"] = _sec_eval_ms
        section["timing"]["total_ms"] = round(
            (section["timing"].get("retrieval_ms") or 0)
            + (section["timing"].get("generation_ms") or 0)
            + _sec_eval_ms,
            1,
        )

    # ── Document-level aggregation ────────────────────────────────────
    from collections import Counter
    from healthark_eval.suite import _grade

    def _mean(vals):
        clean = [float(v) for v in vals if v is not None]
        return round(sum(clean) / len(clean), 4) if clean else None

    if section_results:
        composites = [r.composite_score for r in section_results]
        mean_composite = round(sum(composites) / len(composites), 2)
        grade_dist = dict(Counter(r.grade for r in section_results))

        mean_judge = _mean([
            r.judge_scores.get("normalized_score")
            for r in section_results
            if r.judge_scores and not r.judge_scores.get("judge_error")
        ])
        mean_rag = _mean([
            r.rag_scores.get("rag_triad_score") or r.rag_scores.get("ragas_score")
            for r in section_results if r.rag_scores
        ])
        mean_faith = _mean([
            r.rag_scores.get("faithfulness")
            for r in section_results if r.rag_scores
        ])
    else:
        mean_composite, grade_dist = 0.0, {}
        mean_judge = mean_rag = mean_faith = None

    # Opik document-level aggregation
    opik_sections = [s.get("opik_eval", {}) for s in run_artifacts.get("sections", []) if s.get("opik_eval")]
    mean_hallucination = _mean([s.get("hallucination_score") for s in opik_sections])
    mean_answer_rel = _mean([s.get("answer_relevance_score") for s in opik_sections])
    mean_reg_tone = _mean([s.get("regulatory_tone_score") for s in opik_sections])
    mean_opik_composite = _mean([s.get("opik_composite") for s in opik_sections])

    run_artifacts["extended_eval_summary"] = {
        "mean_composite": mean_composite,
        "overall_grade": _grade(mean_composite),
        "sections_evaluated": len(section_results),
        "grade_distribution": grade_dist,
        # DeepEval metrics
        "mean_judge_normalized": mean_judge,
        "mean_rag_triad_score": mean_rag,
        "mean_faithfulness": mean_faith,
        "mean_bertscore_f1": None,
        # Opik-style metrics
        "mean_hallucination_score": mean_hallucination,
        "mean_answer_relevance_score": mean_answer_rel,
        "mean_regulatory_tone_score": mean_reg_tone,
        "mean_opik_composite": mean_opik_composite,
        # Backwards-compat
        "mean_ragas": mean_rag,
        "framework": "deepeval_rag_triad + opik_style",
    }

    # ── Finalise overall pipeline timing ─────────────────────────────
    _total_eval_ms = round((time.perf_counter() - _eval_phase_start) * 1000, 1)
    if "timing" in run_artifacts and run_artifacts["timing"] is not None:
        run_artifacts["timing"]["total_eval_ms"] = _total_eval_ms
        gen_phase = run_artifacts["timing"].get("generation_phase_ms") or 0
        run_artifacts["timing"]["total_pipeline_ms"] = round(gen_phase + _total_eval_ms, 1)

    # Re-evaluate with all extended data (rule scores + extended metrics)
    rules = get_eval_rules()
    evaluation = evaluate_run(run_artifacts, rules)

    # ── Performance analysis (latency + failures + improvements) ─────
    try:
        from src.eval.eval_performance import PerformanceAnalyzer
        perf_report = PerformanceAnalyzer().analyze(run_artifacts, evaluation)
        run_artifacts["performance_report"] = perf_report.to_dict()
        logging.info(
            "Performance: %s | failures=%d | improvements=%d",
            perf_report.summary_technical[:120],
            len(perf_report.failures),
            len(perf_report.improvements),
        )
    except Exception as exc:
        logging.warning("Performance analysis failed: %s", exc)

    # ── MLflow tracking ───────────────────────────────────────────────
    try:
        tracker = MLflowTracker()
        mlflow_run_id = tracker.log_run(
            run_artifacts=run_artifacts,
            eval_summary=evaluation,
            extended_summary=run_artifacts["extended_eval_summary"],
        )
        run_artifacts["mlflow_run_id"] = mlflow_run_id
        run_artifacts["mlflow_ui_url"] = tracker.run_url(mlflow_run_id)
        logging.info("MLflow run logged: %s", mlflow_run_id)
    except Exception as exc:
        logging.warning("MLflow logging failed: %s", exc)

    # ── Save everything to disk (after performance + MLflow are set) ──
    save_eval_run(run_artifacts, evaluation)

    logging.info(
        "Evaluation complete — composite=%.1f grade=%s | judge=%.1f | "
        "rag_triad=%.3f | hallucination=%.3f | reg_tone=%.3f",
        mean_composite, _grade(mean_composite),
        mean_judge or 0.0, mean_rag or 0.0,
        mean_hallucination or 0.0, mean_reg_tone or 0.0,
    )


def extraction_pmf(template_file_path):

    load_dotenv()
    AZURE_KEY = os.getenv('AZURE_KEY', '')
    AZURE_ENDPOINT = os.getenv('AZURE_ENDPOINT', '')
    AZURE_NAME = os.getenv('AZURE_NAME', '')
    AZURE_VERSION = os.getenv('AZURE_VERSION', '')

    llm = None
    client = None
    try:
        llm = AzureChatOpenAI(
            azure_deployment=AZURE_NAME,
            api_key=AZURE_KEY,
            azure_endpoint=AZURE_ENDPOINT,
            api_version=AZURE_VERSION,
            temperature=0.1,
        )

        client = AzureOpenAI(
        api_key=AZURE_KEY,  
        api_version=AZURE_VERSION,
        azure_endpoint=AZURE_ENDPOINT,
)
    except Exception as e:
        print("Failed to initialize AzureChatOpenAI!")
        print("Error:", e)


                
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    run_artifacts = {
        "timestamp": timestamp,
        "template_file": template_file_path,
        "site_name": st.session_state.get("SiteName", ""),
        "model_name": AZURE_NAME,
        "sections": [],
    }
    

 
   
    # ------------------------------------- Devide json into 2 Part--------------------------------------------------------------------------------------------
    
    template_content = extract_text_from_word(template_file_path)
    part1,part2=Template_to_list(template_content)
    template_json=convert_dict(part1)
    
                    
#--------------------------------------------------------------------------------------------------------------  
    
    # Output doc with footer header
    template_path_1 = fr"templates\\output_template_1.docx"
    template_path = fr"templates\\output_template_PMF.docx"
    # Create a new file name
    new_file_name = fr"PMF_Output_{timestamp}.docx"
    new_file_name_pdf = fr"PMF_Output_{timestamp}.pdf"
    
    
    
    first_doc = Document(template_path)
    for paragraph in first_doc.paragraphs:
        for run in paragraph.runs:
            if "[Site Name]" in run.text:
                # Replace the text
                run.text = run.text.replace("[Site Name]", st.session_state.SiteName)
                # Make the replaced text bold
                run.bold = True
    
    
    doc = Document(template_path_1)

   


    
    

# --------------------------------Iterate over after executive Summary--------------------------------------------------------------------
    doc_1=doc
    
    base_folder = fr"data/artifacts/Extracted_folder"
    folders = [f for f in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, f))]
    st.write(folders)
    if len(folders)<1:
        base_folder = rf"data/artifacts/Extracted_folder"
    else:
        base_folder = rf"data/artifacts/Extracted_folder/{folders[0]}"


    # convert all doc files to docx in the folder
    convert_all_doc_to_docx_in_folder(base_folder) 

    # Initialize the DocumentRetriever with the base folder and store the vector database in the specified path
    retriever = DocumentRetriever(base_folder)
    retriever.process_documents()
    
    st.write("Extracting text from documents...")
    _pipeline_start = time.perf_counter()


    # Check if the uploaded file is a CSV or Excel file and read it into a DataFrame
    data_as_string = ""
    if st.session_state.uploaded_excel is not None:
   
        data_as_string = extract_text_from_xlsx(st.session_state.uploaded_excel)

        st.write(data_as_string)
    



    # TEST MODE: limit to first 5 sections for fast iteration — remove this line for full runs
    template_json = dict(list(template_json.items())[:5])

    for key,value in template_json.items():
       
        value = value.replace("[Site Name]", st.session_state.SiteName)
        value_ls= value.split('@!')
        
        flag=1
        index=1
        
        st.write(key)
        first_25 = value[:80]
        st.write(first_25)
        input_file_path=''
        st.write("")
        st.write("")

        section_json={key:value}
        _sec_start = time.perf_counter()
        _ret_ms = None
        _gen_ms = None
        try:

            input_file_text=""
            paths = []
            # Retrive the documents name from the vector database
            if len(value_ls)>1:
                st.write(value_ls[1])
                _ret_t0 = time.perf_counter()
                results = retriever.search(value_ls[1], top_k=5)
                _ret_ms = (time.perf_counter() - _ret_t0) * 1000

                for i, (path, filename, score) in enumerate(results):
                    paths.append(path)
                st.write("Retrieved documents:")
                st.write(paths)
                    
                input_file_text +="excel data:\n" + data_as_string + "\n\n"+ "########################################\n\n"
                input_file_text += data_extraction(paths)

                st.write("-----------------------------------------------------------------------------------------------------------------------------")
              
            elif "static" not in key.lower(): 
                base_folder_ls=[base_folder]

                input_file_text +="excel data:\n" + data_as_string + "\n\n"+ "########################################\n\n"
                input_file_text += data_extraction(base_folder_ls)
            
            

            # try:   
            if "static" not in key.lower():
                _gen_t0 = time.perf_counter()
                response_data = handle_user_message(llm,client,key,value_ls[0],doc_1,flag,index,input_file_text,input_file_path)
                _gen_ms = (time.perf_counter() - _gen_t0) * 1000
            else:
                _gen_t0 = time.perf_counter()
                response_data = handle_user_message(llm,client,key,value_ls[0],doc_1,flag,index)
                _gen_ms = (time.perf_counter() - _gen_t0) * 1000

            # Use a meaningful title derived from the section value instead of
            # the template type-prefix key ("Text: -", "Static_text:-", etc.)
            display_key = _derive_section_title(key, value_ls[0] if value_ls else value)

            run_artifacts["sections"].append(
                {
                    "section_key": display_key,
                    "template_key": key,          # original type prefix kept for debugging
                    "prompt_text": value_ls[0] if len(value_ls) > 0 else "",
                    "retrieval_query": value_ls[1] if len(value_ls) > 1 else "",
                    "retrieved_paths": paths,
                    "input_text_size": len(input_file_text or ""),
                    "is_static": "static" in key.lower(),
                    "agent_result": response_data,
                    "generated_text": (response_data or {}).get("result", {}).get("generated_text", ""),
                    "timing": {
                        "retrieval_ms": round(_ret_ms, 1) if _ret_ms is not None else None,
                        "generation_ms": round(_gen_ms, 1) if _gen_ms is not None else None,
                        "eval_ms": None,
                        "total_ms": round((time.perf_counter() - _sec_start) * 1000, 1),
                    },
                }
            )
        except Exception as e:
            st.write("Error in processing the section:")
            st.write(f"An error occurred: {e}")
            _sec_err_ms = round((time.perf_counter() - _sec_start) * 1000, 1)
            run_artifacts["sections"].append({
                "section_key": _derive_section_title(key, value_ls[0] if value_ls else value),
                "template_key": key,
                "prompt_text": value_ls[0] if len(value_ls) > 0 else "",
                "retrieval_query": value_ls[1] if len(value_ls) > 1 else "",
                "retrieved_paths": [],
                "input_text_size": 0,
                "is_static": "static" in key.lower(),
                "agent_result": None,
                "generated_text": "",
                "generation_error": str(e),
                "timing": {
                    "retrieval_ms": round(_ret_ms, 1) if _ret_ms is not None else None,
                    "generation_ms": round(_gen_ms, 1) if _gen_ms is not None else None,
                    "eval_ms": None,
                    "total_ms": _sec_err_ms,
                },
            })
            continue
                
           
            st.write("----------------------------------------------------------------------------------------------------------------------------")    

         

   

    # Store overall pipeline timing (generation phase only; eval_ms added later)
    _generation_phase_ms = round((time.perf_counter() - _pipeline_start) * 1000, 1)
    run_artifacts["timing"] = {
        "generation_phase_ms": _generation_phase_ms,
        "total_pipeline_ms": None,   # updated after eval completes
        "total_generation_ms": round(sum(
            (s.get("timing") or {}).get("generation_ms") or 0
            for s in run_artifacts["sections"]
        ), 1),
        "total_retrieval_ms": round(sum(
            (s.get("timing") or {}).get("retrieval_ms") or 0
            for s in run_artifacts["sections"]
        ), 1),
        "total_eval_ms": None,       # filled by _run_extended_evaluation
    }

    # Save the document to a new file----------------------------------------------------------------------------------------------------
    new_file_path = fr"data/artifacts/generated output file\\{new_file_name}" 
    new_file_path_temp =rf"data/artifacts/generated output file\Temp_{new_file_name}"
    final_document_doc = rf"data/artifacts/generated output file\\Final_{new_file_name}"
    new_file_path_pdf =rf"data/artifacts/generated output file\\{new_file_name_pdf}"
    final_document = fr"data/artifacts/generated output file\\Final_{new_file_name_pdf}" 

    

    # Generated file save in doc
    composer = Composer(doc_1)
    composer.save(new_file_path_temp)


    # add Content to the document
    doc_t=extract_headings_with_tables(new_file_path_temp,0, final_document_doc)
    file_path_abs = os.path.abspath(final_document_doc) 
    refresh_toc_with_word(file_path_abs)

    
    # Add cover page to the document
    doc1 = Document(final_document_doc)
    composer1 = Composer(first_doc)
    composer1.append(doc1)
    composer1.save(new_file_path_temp)
    file_path_abs = os.path.abspath(new_file_path_temp)

    # create a link to the document
    link_doc = Document(new_file_path_temp)
    output = io.BytesIO()
    link_doc.save(output)
    output.seek(0)  

    run_artifacts["final_doc_path"] = new_file_path_temp
    rules = get_eval_rules()
    evaluation = evaluate_run(run_artifacts, rules)
    eval_file = save_eval_run(run_artifacts, evaluation)
    st.session_state["last_eval_file"] = eval_file
    st.session_state["last_eval_score"] = evaluation.get("document_scores", {}).get("overall_score")

    # Always run extended evaluation using the Azure credentials already in .env.
    # No checkbox needed — Judge + DeepEval RAG Triad run automatically after generation.
    try:
        eval_suite = EvalSuite(
            task="pmf",
            llm_provider="azure_openai",
            llm_model=AZURE_NAME,
            api_key=AZURE_KEY,
            azure_endpoint=AZURE_ENDPOINT,
            azure_api_version=AZURE_VERSION,
            run_judge=True,
            run_rag=True,
            run_semantic=False,   # BERTScore skipped — torch DLL issue on Windows
            run_lexical=False,    # No reference text available during generation
        )
        run_artifacts["_azure_key"] = AZURE_KEY
        run_artifacts["_azure_endpoint"] = AZURE_ENDPOINT
        run_artifacts["_azure_version"] = AZURE_VERSION
        _run_extended_evaluation(run_artifacts, eval_suite)
        ext = run_artifacts.get("extended_eval_summary", {})
        st.session_state["last_extended_composite"] = ext.get("mean_composite", 0)
        st.session_state["last_extended_grade"] = ext.get("overall_grade", "?")
        st.session_state["last_extended_judge"] = ext.get("mean_judge_normalized")
        st.session_state["last_extended_rag"] = ext.get("mean_rag_triad_score")
        st.session_state["last_mlflow_run_id"] = run_artifacts.get("mlflow_run_id")
        st.session_state["last_mlflow_url"] = run_artifacts.get("mlflow_ui_url", "http://localhost:5000")
    except Exception as _ext_exc:
        logging.warning("Extended evaluation failed: %s", _ext_exc)
    

   



  


# #------------------------Appendices assembling---------------------------------------------------------
    
    # # Define the base folder path
    # base_folder = r"data/artifacts/Extracted_folder"
    # folders = [f for f in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, f))]
    # if len(folders)<1:
    #     base_folder = rf"data/artifacts/Extracted_folder"
    # else:
    #     base_folder = rf"data/artifacts/Extracted_folder/{folders[0]}"

    # # file_path_abs=r"C:\Users\mihir.vasoya\OneDrive - Thermo Fisher Scientific\Desktop\Git-Integration\RegulatoryDocGen\generated output file\PMF_German_9_7_25.docx"
    # convert_word_to_pdf(file_path_abs,new_file_path_pdf)
    
    

    # res=assemble_appendices(base_folder, new_file_path_pdf, final_document)
    # print(res)









    # final_document=os.path.abspath(final_document) 
   

    # output=None
    return output,final_document,new_file_name_pdf

    


