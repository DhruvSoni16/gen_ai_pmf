import io
from docx import Document
import base64
import docx
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import pandas as pd
from docx.shared import Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
import io
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_ALIGN_PARAGRAPH

import streamlit as st



def set_table_border(table):
    # Loop through rows and cells to set borders for each cell
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')  # Create table cell borders element
            
            # Define border styles (top, bottom, left, right)
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Border type: single line
                border.set(qn('w:sz'), '4')       # Border size: 4 (1/8 pt)
                border.set(qn('w:space'), '0')    # No space between border and content
                border.set(qn('w:color'), '000000')  # Border color: black
                borders.append(border)  # Append border element
            
            tcPr.append(borders)  # Add borders to cell properties

def set_cell_background(cell, color):
    """Set cell background shading color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Create a new shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # Set fill color
    tcPr.append(shd)

def process_text_to_docx(doc,text):
    

    # Split the input text by lines
    lines = text.splitlines()

    for line in lines:
        # Handle headings (## for H2)
        if line.startswith('#'):
            line = line.replace("#", " ")
            paragraph = doc.add_paragraph(line[3:])
            run1 = paragraph.runs[0]  # Access the first run in the paragraph
            run1.bold = True
            
        # Handle bold text (** for bold)
        elif '**' in line:
            p = doc.add_paragraph()
            bold_parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in bold_parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True  # Add bold text
                else:
                    p.add_run(part)  # Add regular text

        # Handle bullet points (* for bullet lists)
        elif line.startswith('*'):
            doc.add_paragraph(line[2:], style='ListBullet')

        # Handle other normal paragraphs
        else:
            if line.strip() != "":  # Avoid adding empty lines
                doc.add_paragraph(line)

# download doc file from link
def generate_word_download_link(doc_data, filename):
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">Click here to download the Word document</a>'
    return href

def generate_pdf_download_link(pdf_path, filename):
    with open(pdf_path, "rb") as f:
        pdf_data = f.read()
    b64 = base64.b64encode(pdf_data).decode()
    href = f'<a href="data:application/pdf;base64,{b64}" download="{filename}.pdf">Click here to download the PDF</a>'
    return href

# dynamic
def save_text_in_document_1(input,doc,flag,index=0,value=""):
    # print(text)
    
    style = doc.styles["Normal"]

    font = style.font
    font.name = "Times New Roman"
    
    if flag==0:
        process_text_to_docx(doc, input)

    elif flag==1:
        # num=value[0]
        image_extraction_paragraph = doc.add_paragraph(style='Normal')
        # run = image_extraction_paragraph.add_run(f"Image Extraction")
        # run.bold = True
        if input is not None:
        
            try:
                st.image(input, caption="")
                doc.add_picture(input, width=Inches(5.7))
                
                image_extraction_paragraph = doc.add_paragraph(style='Normal')
                image_extraction_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = image_extraction_paragraph.add_run(f"{value}")
                run.bold = True
            except:
                image_text = "**Image is not found**"
                process_text_to_docx(doc, image_text)

                image_extraction_paragraph = doc.add_paragraph(style='Normal')
                image_extraction_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = image_extraction_paragraph.add_run(f"{value}")
                run.bold = True
                

        else:
            image_text = "Image is not found"
            process_text_to_docx(doc, image_text)


# -------------------------------Table add into output doc-------------------------------------------------------------------
    elif flag==2:
        i = 0
        
        for key,df in input.items():
            table_name=None
            # extract heading and table name from the table json
            try:
                table_name=df[0]["table_name"]
            except:
                pass

            table=df[0]["columns"]



         
            # Json table convert to df
            if len(table)>1:
                df0=pd.DataFrame(table)
            else:
                df0=pd.DataFrame(list(table[0].items()), columns=["Attribute", "Value"])

            

            # Add table name and heading into output documents
            try:
                if table_name is not None:

                        paragraph = doc.add_paragraph(f"{table_name.strip()}")
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run1 = paragraph.runs[0]  # Access the first run in the paragraph
                        run1.bold = True
                else:
                    paragraph = doc.add_paragraph(f"{table_name.strip()}")
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run1 = paragraph.runs[0]  # Access the first run in the paragraph
                    run1.bold = True
            except:
                pass

            # Add table into output doc
            if df0.empty:
                print("The DataFrame is empty. No table will be created.")
            else:
                # Create table in document

                if "Value" not in df0.columns:

                    if "col_0" in df0.columns:
                        table = doc.add_table(rows=0, cols=len(df0.columns))
                
                    # Add rows from DataFrame
                        for index, row in df0.iterrows():
                            row_cells = table.add_row().cells
                            for j, cell in enumerate(row):
                                row_cells[j].text = str(cell)

                    else:
                        first_col_values = df0.iloc[:, 0].unique()
                        second_col_values = df0.iloc[:, 1].unique()
                        if len(first_col_values) == 1 and len(second_col_values) == 1 and len(df0.columns)>3:
                            first_2_column= df0.iloc[:, :2]
                            df0= df0.iloc[:, 2:]



                            table = doc.add_table(rows=0, cols=len(first_2_column.columns))

                            row_cells = table.add_row().cells
                            row_cells[0].text = str(first_2_column.columns[0])
                            row_cells[0].paragraphs[0].runs[0].font.bold = True
                            row_cells[1].text = str("\n".join(first_2_column.iloc[:, 0].unique()))

                            row_cells = table.add_row().cells
                            row_cells[0].text = str(first_2_column.columns[1])
                            row_cells[0].paragraphs[0].runs[0].font.bold = True
                            row_cells[1].text = str("\n".join(first_2_column.iloc[:, 1].unique()))

                            set_table_border(table)


                    
                        # Add rows from DataFrame
                        
                            

                        table = doc.add_table(rows=1, cols=len(df0.columns))

                        # Add column headers
                        hdr_cells = table.rows[0].cells
                        for j, col in enumerate(df0.columns):
                            if col is not None:  # Check for None values
                                hdr_cells[j].text = str(col)  # Ensure col is a string
                                hdr_cells[j].paragraphs[0].runs[0].font.bold = True
                                set_cell_background(hdr_cells[j], "F2F2F2")
                                

                        # Add rows from DataFrame
                        for index, row in df0.iterrows():
                            # Condition: More than 3 columns and more than 5 rows
                            if df0.shape[1] > 3 and df0.shape[0] > 5:
                                first_col_value = str(row[0]).strip().lower()
                                user_input_count = sum(1 for val in row if str(val).strip().lower() == "user input required")

                                # If first column is "user input required" OR count of "user input required" > 3 → skip
                                if first_col_value == "user input required" or user_input_count > 2:
                                    continue

                            # Add row to Word table
                            row_cells = table.add_row().cells
                            for j, cell in enumerate(row):
                                row_cells[j].text = str(cell)

         
                else:
                    table = doc.add_table(rows=0, cols=len(df0.columns))
                
                    # Add rows from DataFrame
                    for index, row in df0.iterrows():
                        row_cells = table.add_row().cells
                        for j, cell in enumerate(row):
                            row_cells[j].text = str(cell)

                set_table_border(table)
            doc.add_paragraph("")
            i += 1
    elif flag==3:
        # static
        pass
    else:
        # Web
        pass









