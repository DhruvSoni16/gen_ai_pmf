import streamlit as st
import os
import io
import logging
import time
import shutil

from src.document_analyzer.Extraction_module_PMF import extraction_pmf
from src.document_generate.doc_generate import generate_word_download_link
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import zipfile
import stat
import gc
import pandas as pd
from src.eval.eval_store import list_runs, load_run_by_file

# streamlit run app.py --server.maxUploadSize=1000

# Reset logging configuration
for handler in logging.root.handlers[:]:
    logging.root.removeHandler(handler)

logger = logging.getLogger()

# Check if handlers are already present
if not logger.hasHandlers():
    # Set logging level
    logger.setLevel(logging.INFO)

    # Create handlers
    file_handler = logging.FileHandler('logs/app.log')
    console_handler = logging.StreamHandler()

    # Set levels for handlers
    file_handler.setLevel(logging.INFO)
    console_handler.setLevel(logging.INFO)

    # Create formatter and add to handlers
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)

    # Add handlers to logger
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
# logging.basicConfig(filename='app.log', level=logging.INFO, filemode='a',
#                     format='%(asctime)s - %(levelname)s - %(message)s')
logging.info("App started")
def add_logo_to_docx(doc_data, logo_path="pdf_logo.png", text="The world leader in serving science"):
    """Adds a logo to the left side of the top header section and specified text to the right side.

    Args:
        doc_data (bytes): The content of the Word document as a byte stream.
        logo_path (str, optional): The path to the logo image file. Defaults to "pdf_logo.png".
        text (str, optional): The text to add to the header. Defaults to "world leader in serving science".

    Returns:
        bytes: The modified Word document content as a byte stream.
    """
    
    document = Document(doc_data)

    # Get the first section's header
    section = document.sections[0]
    header = section.header

    # Create a table with two cells: one for the logo, one for the text
    table = header.add_table(rows=1, cols=2, width=Inches(6))
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align table to left

    # Adjust column widths
    table.columns[0].width = Inches(1.5)  # Width for the logo
    table.columns[1].width = Inches(4.5)  # Width for the text

    # Add logo to the first cell
    logo_cell = table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.25))

    # Add text to the second cell, aligning it to the right
    text_cell = table.cell(0, 1)
    text_paragraph = text_cell.paragraphs[0]
    text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Create a run with a manual line break
    bold_run = text_paragraph.add_run("The world leader\nin serving science")
    bold_run.font.name = "Times New Roman"
    bold_run.bold = True  # Make the text bold

    # Optionally set margins for all sections
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Save the document and return as byte stream
    document.save("temp.docx")
    with open("temp.docx", "rb") as f:
        return f.read()


def cleanup_environment():
    """
    Clean up all temporary variables, cached data, and free up memory for the next run.
    """
    print("Cleaning up the environment...")

    # Clear global variables
    global all_tokens, competitor_analysis
    all_tokens = []
    competitor_analysis = []

    # Clear any other global variables or temporary data
    globals_to_clear = [var for var in globals() if not var.startswith("__") and not callable(globals()[var])]
    for var in globals_to_clear:
        if isinstance(globals()[var], (pd.DataFrame, list, dict, set)):
            globals()[var] = None  # Reset to None
            print(f"Cleared variable: {var}")

    # Force garbage collection
    gc.collect()
    print("Environment cleanup completed.")


# Title of the app
st.markdown(
    "<h1 style='text-align: center; color: red;'>Document Generator</h1>",
    unsafe_allow_html=True
)
st.markdown("""
    <style>
    /* Styling for the main submit button */
    div.stButton > button {
        background-color: red;  /* Red background */
        width: 200px;           /* Adjust width */
        height: 40px;           /* Adjust height */
        border: none;           /* Remove border */
        cursor: pointer;        /* Change cursor on hover */
        border-radius: 25px;    /* Rounded corners */
        color: Black;           /* White text color */
        border: 2px solid white; /* Border */
        margin-top: 5px;
    }

    /* Hover effect for the buttons */
    div.stButton > button:hover {
        background-color: white;
        color: red;
        border: 2px solid red;
    }
    </style>
""", unsafe_allow_html=True)

image_name=[]
# Example: Add your logo (local or online image path)
logo_url = r"static\logo1.png"  

# Add the logo to the sidebar
st.sidebar.image(logo_url, width=200,caption=None )

# Initialize session state variables if they don't exist
if 'selected_feature' not in st.session_state:
    st.session_state.selected_feature = None
    
col1, col2, col3 = st.columns(3)


#---------------------------------------Side Bar------------------------------------------------------------------------------

st.sidebar.header("Document Generator Menu")
with st.sidebar:
    
    if st.button("Technical File"):
        st.session_state.selected_feature = 'Technical'

    if st.button("CER File"):
        st.session_state.selected_feature = 'CER'

    if st.button("Plant Master File"):
        st.session_state.selected_feature = 'Plant'

    if st.button("Evaluation Dashboard"):
        st.session_state.selected_feature = 'Eval'

    

def handle_remove_readonly(func, path, excinfo):
    """
    Handle read-only file permission error by changing the file permissions
    and then retrying the operation.
    """
    import stat
    os.chmod(path, stat.S_IWRITE)  # Change file permissions to writeable
    func(path)  # Retry the operation

def clear_extracted_folder(folder_path):
    """
    Clears all files and subdirectories in the specified folder.
    """
    if os.path.exists(folder_path):
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # Remove file or symbolic link
                elif os.path.isdir(file_path):
                    # Use shutil.rmtree with error handling for permissions
                    shutil.rmtree(file_path, onerror=handle_remove_readonly)
            except Exception as e:
                print(f"Failed to delete {file_path}. Reason: {e}")
    else:
        os.makedirs(folder_path)




# Technical File
if st.session_state.selected_feature == 'Technical':
   pass

if st.session_state.selected_feature == 'Eval':
    try:
        from app_eval_dashboard import render_eval_dashboard
        render_eval_dashboard()
    except ImportError:
        # Fallback to basic dashboard if extended dashboard unavailable
        st.subheader("LLM Evaluation Dashboard")
        st.write("Review historical PMF generation runs and rule-based scores.")

        runs = list_runs()
        if not runs:
            st.info("No evaluation runs found yet. Generate a PMF document first.")
        else:
            run_df = pd.DataFrame(runs)
            st.dataframe(run_df[["timestamp", "site_name", "overall_score", "template_file"]], use_container_width=True)

            run_labels = [
                f"{row.get('timestamp', '')} | {row.get('site_name', '')} | score={row.get('overall_score', '')}"
                for row in runs
            ]
            selected_idx = st.selectbox("Select a run", options=list(range(len(runs))), format_func=lambda i: run_labels[i])
            selected_run = runs[selected_idx]

            run_payload = load_run_by_file(selected_run["run_file"])
            eval_data = run_payload.get("evaluation", {}).get("document_scores", {})

            col1, col2, col3 = st.columns(3)
            col1.metric("Overall Score", eval_data.get("overall_score", 0))
            col2.metric("Section Count", eval_data.get("section_count", 0))
            col3.metric("Retrieval Coverage (%)", eval_data.get("retrieval_coverage", 0))

            missing = eval_data.get("missing_required_sections", [])
            if missing:
                st.warning("Missing required sections: " + ", ".join(missing))
            else:
                st.success("All required sections present.")

            section_rows = eval_data.get("sections", [])
            if section_rows:
                section_df = pd.DataFrame(
                    [
                        {
                            "section_key": row.get("section_key"),
                            "score": row.get("score"),
                            "char_len": row.get("char_len"),
                            "min_chars": row.get("required_min_chars"),
                            "missing_keywords": ", ".join(row.get("missing_keywords", [])),
                            "checks": str(row.get("checks", {})),
                        }
                        for row in section_rows
                    ]
                )
                st.dataframe(section_df, use_container_width=True)

                st.write("Score Trend")
                trend_df = run_df[["timestamp", "overall_score"]].copy()
                trend_df["overall_score"] = pd.to_numeric(trend_df["overall_score"], errors="coerce")
                trend_df = trend_df.sort_values("timestamp")
                st.line_chart(trend_df.set_index("timestamp"))

if st.session_state.selected_feature =='Plant':
    st.subheader("📄 Plant Master File Generator")
    st.write("")

    uploaded_file = st.file_uploader("Upload a ZIP file containing your folder", type="zip")


    # File uploader widget
    st.session_state.uploaded_excel = st.file_uploader("Choose a Reference Excel File", type=["xlsx", "xls"])

    st.session_state.SiteName = st.text_input("Enter the site name", value="")

    template_doc = r"templates\PMF_Template_With_vector_DB - Copy.docx"
        
    
  
    st.write("")
    st.write("")

    if st.button("Submit"):
            clear_extracted_folder("data/artifacts/Extracted_folder")
            if uploaded_file is None :
                st.warning("🚨 Please fill in all required fields and upload files before submitting.")
            
            
            
            else:
                # Extract the uploaded Zip files in Extracted_folder
                if uploaded_file:
                    with open("temp.zip", "wb") as f:
                        f.write(uploaded_file.getbuffer())
                
                    # Extract the ZIP file
                    clear_extracted_folder("data/artifacts/Extracted_folder")
                    extract_dir = "data/artifacts/Extracted_folder"
                    # Clean up previous extractions
                    with zipfile.ZipFile("temp.zip", "r") as zip_ref:
                        zip_ref.extractall(extract_dir)

                    folder_structure = {}
                
                    # Recursively walk through the extracted folder
                    
                    for root, dirs, files in os.walk(extract_dir):
                        # Store the relative folder path
                        relative_root = os.path.relpath(root, extract_dir)
                        if relative_root == ".":
                            relative_root = "Root"  # Rename the top-level folder for clarity
                        
                        # Store the folder and its files in the dictionary
                        folder_structure[relative_root] = files
                        
                        # Display files in this folder, if any
                        if files:
                            for file in files:
                                relative_path = os.path.join(relative_root, file)
                        else:
                            pass
                    
                    os.remove("temp.zip")



#---------------------------------------------------------------------------------------------------------
                start_time = time.time()

                text,pdf_path,pdf_name = extraction_pmf(template_doc)
                
                response_time = time.time() - start_time
               
                logging.info(f"Response generation time: {response_time:.2f} seconds")

                # file_name = os.path.splitext(documnet_name)[0]
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            
                prefix = "PMF_Document" 
               
                if pdf_path: 
                    st.subheader("📥 The generated file has been saved at the following folder path:")

                    st.html(f"<p><strong>Response generation time:- </strong>{response_time:.2f} seconds</p>")
                    
                    st.html(f"<p><strong>PDF Path:- </strong>{pdf_path}</p>")
                    
                    # Generate the Word document download link
                    final_file_name = f"PMF_{timestamp}.docx"
                    if text: 
                        st.subheader("📥 You can download generated word file here: ")
                    else:  
                        st.write("No text available for download.")
                        st.button("Download", disabled=True)
                    st.markdown(generate_word_download_link(text.getvalue(),final_file_name), unsafe_allow_html=True)

                    # ── Evaluation summary ──────────────────────────────────────
                    _rule_score = st.session_state.get("last_eval_score")
                    _ext_composite = st.session_state.get("last_extended_composite")
                    _ext_judge = st.session_state.get("last_extended_judge")
                    _ext_rag = st.session_state.get("last_extended_rag")
                    _ext_grade = st.session_state.get("last_extended_grade")
                    _mlflow_run_id = st.session_state.get("last_mlflow_run_id")
                    _mlflow_url = st.session_state.get("last_mlflow_url", "http://localhost:5000")

                    # Use composite from extended eval if available, else rule score
                    _display_score = float(_ext_composite) if _ext_composite is not None else (float(_rule_score) if _rule_score is not None else 0.0)
                    _display_grade = _ext_grade if _ext_grade else ("A" if _display_score >= 90 else "B" if _display_score >= 75 else "C" if _display_score >= 60 else "D" if _display_score >= 45 else "F")
                    _grade_clr = "#1D9E75" if _display_score >= 75 else "#BA7517" if _display_score >= 60 else "#D85A30"
                    _health = "Good Quality" if _display_score >= 75 else "Acceptable" if _display_score >= 60 else "Needs Improvement"

                    _judge_str = f"{float(_ext_judge):.1f}/100" if _ext_judge is not None else "computing..."
                    _rag_str = f"{float(_ext_rag):.3f}" if _ext_rag is not None else "computing..."
                    _rule_str = f"{float(_rule_score):.1f}/100" if _rule_score is not None else "—"

                    _mlflow_line = ""
                    if _mlflow_run_id:
                        _mlflow_line = (
                            f'<br><span style="font-size:0.82rem;">'
                            f'📊 <a href="{_mlflow_url}" target="_blank" style="color:#5340C0;">'
                            f'View in MLflow UI</a> &nbsp;'
                            f'<span style="color:#9ca3af;">(run: {_mlflow_run_id[:8]}…)</span>'
                            f'</span>'
                        )

                    st.markdown(
                        f'<div style="background:{_grade_clr}22;border-left:4px solid {_grade_clr};'
                        f'padding:14px 18px;border-radius:8px;margin:16px 0;">'
                        f'<span style="font-size:1.1rem;font-weight:700;color:{_grade_clr};">'
                        f'Grade {_display_grade} — {_health} &nbsp;({_display_score:.1f}/100)</span><br><br>'
                        f'<span style="color:#374151;font-size:0.9rem;">'
                        f'Rule Score: <strong>{_rule_str}</strong> &nbsp;|&nbsp; '
                        f'Judge Score: <strong>{_judge_str}</strong> &nbsp;|&nbsp; '
                        f'RAG Triad: <strong>{_rag_str}</strong>'
                        f'</span><br>'
                        f'<span style="color:#6b7280;font-size:0.82rem;">'
                        f'Open the <strong>Evaluation Dashboard</strong> in the sidebar for full DeepEval + Opik metrics.'
                        f'</span>'
                        f'{_mlflow_line}'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

                    # shutil.rmtree("Extracted_folder")
                    if 'folder_structure' in st.session_state:
                        del st.session_state['folder_structure']
                    if 'extract_dir' in st.session_state:
                        del st.session_state['extract_dir']
                    clear_extracted_folder("data/artifacts/Extracted_folder")
                    cleanup_environment()