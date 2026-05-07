"""
Extended Evaluation Dashboard — Streamlit Application
======================================================

Production-grade 5-tab dashboard for the Healthark GenAI Evaluation Framework.

Tabs:
  1. Run Overview     — metric cards, grade badge, run selector
  2. Section Heatmap  — multi-metric heatmap, section detail drill-down
  3. Trend Analysis   — time-series charts, regression alerts
  4. Model Comparison — file upload, radar chart, cost analysis
  5. Benchmark Mgmt   — statistics, case table, add/export/run

Run standalone:
    streamlit run app_eval_dashboard.py

Or import into app.py:
    from app_eval_dashboard import render_eval_dashboard
    render_eval_dashboard()
"""

from __future__ import annotations

import json
import logging
import os
import sys
from datetime import datetime
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
import streamlit as st

PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from src.eval.eval_store import list_runs, load_run_by_file
from src.eval.eval_config import get_eval_rules
from src.eval.eval_utils import score_section

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Safe imports
# ---------------------------------------------------------------------------

try:
    import plotly.express as px
    import plotly.graph_objects as go
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False

# Lexical metrics (BLEU/ROUGE) — sacrebleu + rouge-score, no torch dependency
try:
    from src.eval.eval_metrics import LexicalMetrics
    HAS_LEXICAL = True
except (ImportError, OSError):
    HAS_LEXICAL = False
    logger.warning("LexicalMetrics unavailable — BLEU/ROUGE disabled")

# Semantic metrics (BERTScore) — bert-score + sentence-transformers, needs torch
try:
    from src.eval.eval_metrics import SemanticMetrics, compute_all_metrics
    HAS_SEMANTIC = True
except (ImportError, OSError):
    HAS_SEMANTIC = False
    logger.warning("SemanticMetrics unavailable — BERTScore disabled")

HAS_METRICS = HAS_LEXICAL or HAS_SEMANTIC

try:
    from src.eval.benchmark_loader import BenchmarkLoader
    HAS_BENCHMARK = True
except ImportError:
    HAS_BENCHMARK = False

BENCHMARK_DIR = os.path.join(PROJECT_ROOT, "data", "benchmark")
EVAL_RESULTS_DIR = os.path.join(PROJECT_ROOT, "eval_results")

# ═══════════════════════════════════════════════════════════════════════════
# DESIGN TOKENS
# ═══════════════════════════════════════════════════════════════════════════

CLR_PRIMARY = "#5340C0"   # purple
CLR_SUCCESS = "#1D9E75"   # teal
CLR_WARNING = "#BA7517"   # amber
CLR_DANGER  = "#D85A30"   # coral

GRADE_COLORS: Dict[str, str] = {
    "A": CLR_SUCCESS,        # teal
    "B": "#3B82F6",          # blue
    "C": CLR_WARNING,        # amber
    "D": CLR_DANGER,         # coral
    "F": "#DC2626",          # red
}

CUSTOM_CSS = f"""
<style>
  .dashboard-header {{
    color: {CLR_PRIMARY};
    font-size: 1.8rem;
    font-weight: 700;
    margin-bottom: 0.25rem;
  }}
  .grade-badge {{
    display: inline-block;
    padding: 4px 16px;
    border-radius: 6px;
    color: white;
    font-weight: 700;
    font-size: 1.4rem;
  }}
  .metric-label {{
    font-size: 0.75rem;
    color: #6b7280;
    text-transform: uppercase;
    letter-spacing: 0.05em;
  }}
  div[data-testid="stMetric"] label {{
    color: {CLR_PRIMARY} !important;
    font-weight: 600 !important;
  }}
  .stTabs [data-baseweb="tab-list"] button {{
    color: {CLR_PRIMARY};
  }}
  .stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {{
    border-bottom-color: {CLR_PRIMARY};
    color: {CLR_PRIMARY};
    font-weight: 700;
  }}
</style>
"""


def _grade_badge_html(grade: str) -> str:
    """Return an HTML span for a letter grade badge."""
    color = GRADE_COLORS.get(grade, "#6b7280")
    return (
        f'<span class="grade-badge" style="background:{color};">'
        f'{grade}</span>'
    )


def _letter_grade(score: float) -> str:
    if score >= 90: return "A"
    if score >= 75: return "B"
    if score >= 60: return "C"
    if score >= 45: return "D"
    return "F"


def _fmt(val: Any, decimals: int = 2) -> str:
    """Format a numeric value for display, or '—' if None."""
    if val is None:
        return "—"
    try:
        return f"{float(val):.{decimals}f}"
    except (ValueError, TypeError):
        return str(val)


# ═══════════════════════════════════════════════════════════════════════════
# DATA LOADING (cached)
# ═══════════════════════════════════════════════════════════════════════════

@st.cache_data(ttl=30)
def _load_runs() -> List[Dict[str, Any]]:
    return list_runs()

@st.cache_data(ttl=30)
def _load_run_payload(run_file: str) -> Dict[str, Any]:
    return load_run_by_file(run_file)

@st.cache_data(ttl=30)
def _load_evalsuite_results() -> List[Dict[str, Any]]:
    """Load all JSON result files saved by EvalSuite.save_results()."""
    results: List[Dict[str, Any]] = []
    if not os.path.isdir(EVAL_RESULTS_DIR):
        return results
    for fname in sorted(os.listdir(EVAL_RESULTS_DIR)):
        if fname.endswith(".json"):
            try:
                with open(os.path.join(EVAL_RESULTS_DIR, fname), "r", encoding="utf-8") as f:
                    results.append(json.load(f))
            except (json.JSONDecodeError, OSError):
                pass
    return results


def _extract_extended(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Pull extended_eval_summary from a run payload if present."""
    return payload.get("run_artifacts", payload).get("extended_eval_summary", {})


def _section_extended(section: Dict[str, Any]) -> Dict[str, Any]:
    """Pull extended_eval from a section dict."""
    return section.get("extended_eval", {})


# ═══════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════

def _render_sidebar() -> None:
    """Global sidebar: about, API keys, model selector, metric toggles, export."""
    with st.sidebar:
        st.markdown(
            f"<h3 style='color:{CLR_PRIMARY};'>About</h3>",
            unsafe_allow_html=True,
        )
        st.markdown(
            "- **Multi-layer evaluation** — Combines rule-based checks, "
            "lexical metrics (BLEU/ROUGE), semantic similarity (BERTScore), "
            "and LLM-as-judge scoring into one composite quality grade.\n"
            "- **RAG pipeline monitoring** — Measures faithfulness, context "
            "precision, and answer relevancy to detect hallucination and "
            "retrieval failures before they reach production.\n"
            "- **Regression tracking** — Automatically compares every run "
            "against a stored baseline and alerts when quality drops, "
            "enabling continuous improvement across model and template updates."
        )
        st.divider()

        st.markdown(f"<h3 style='color:{CLR_PRIMARY};'>Settings</h3>",
                    unsafe_allow_html=True)

        # Pre-fill keys from .env so the user doesn't have to retype them
        _default_anthropic = os.getenv("ANTHROPIC_API_KEY", "")
        _default_azure = os.getenv("AZURE_KEY", os.getenv("AZURE_OPENAI_API_KEY", ""))

        st.text_input(
            "Anthropic API Key",
            value=_default_anthropic,
            type="password",
            key="sidebar_anthropic_key",
            help="Used for LLM-as-Judge. Auto-loaded from .env if ANTHROPIC_API_KEY is set.",
        )
        st.text_input(
            "Azure OpenAI API Key",
            value=_default_azure,
            type="password",
            key="sidebar_azure_key",
            help="Auto-loaded from .env (AZURE_KEY).",
        )

        st.selectbox("Judge Model", [
            "claude-sonnet-4-6", "gpt-4o",
        ], key="sidebar_judge_model")

        st.markdown("**Quick Eval Toggles** *(for the text box below)*")
        st.checkbox(
            "Run LLM Judge",
            value=False,
            key="toggle_judge",
            help="Calls an LLM to score the pasted text on 5 PMF criteria (needs API key above).",
        )
        st.checkbox(
            "Run DeepEval RAG Triad",
            value=False,
            key="toggle_rag",
            help="Faithfulness + Contextual Precision + Answer Relevancy (DeepEval methodology, needs LLM API key).",
        )
        st.info("Judge + RAG Triad run **automatically** after every document generation — no checkbox needed there.")

        st.divider()

        # Export current run as JSON download
        if st.session_state.get("_current_payload"):
            data_str = json.dumps(st.session_state["_current_payload"], indent=2, default=str)
            st.download_button(
                "Export Current Run (JSON)",
                data_str,
                file_name="eval_run_export.json",
                mime="application/json",
            )


# ═══════════════════════════════════════════════════════════════════════════
# TAB 1 — RUN OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════

def _render_tab_overview(runs: List[Dict[str, Any]]) -> None:
    """Metric cards and grade for the current (latest) run."""

    if not runs:
        st.info(
            "No evaluation runs found yet. "
            "Generate a PMF document using the **Plant Master File** option in the sidebar."
        )
        return

    # Always use the latest run — no selector
    run_meta = runs[0]
    run_file = run_meta.get("run_file", "")
    if not run_file or not os.path.exists(run_file):
        st.warning("Run file not found.")
        return

    payload = _load_run_payload(run_file)
    st.session_state["_current_payload"] = payload

    eval_data = payload.get("evaluation", {}).get("document_scores", {})
    ext = _extract_extended(payload)

    rule_score = eval_data.get("overall_score", 0)
    judge_score = ext.get("mean_judge_normalized")
    rag_triad = ext.get("mean_rag_triad_score") or ext.get("mean_ragas")
    mean_faith = ext.get("mean_faithfulness")
    composite = ext.get("mean_composite", rule_score)

    # ── 1. What do these metrics mean? (at the very top) ─────────────────
    with st.expander("What do these metrics mean?"):
        st.markdown(
            "**Rule Score** — Structural quality checks: minimum character length, "
            "required sections present, site name referenced.\n\n"
            "**Judge Score** — A second LLM acts as a domain expert and scores the "
            "output on 5 PMF-specific criteria: Factual Accuracy, Regulatory Language, "
            "Site Specificity, Completeness, Structural Coherence. Score: 0–100.\n\n"
            "**Faithfulness** *(DeepEval)* — Extracts every factual claim from the "
            "generated text and checks each one against the retrieved source documents. "
            "Score = supported_claims / total_claims. Low score = hallucination.\n\n"
            "**Contextual Precision** *(DeepEval)* — Checks whether the most relevant "
            "retrieved context chunks are ranked first. Rank-weighted Average Precision.\n\n"
            "**Answer Relevancy** *(DeepEval)* — Generates questions the answer would "
            "address, then measures how closely they match the original instruction.\n\n"
            "**RAG Triad** — Harmonic mean of Faithfulness + Contextual Precision + "
            "Answer Relevancy (0–1 scale). Industry-standard composite for RAG quality.\n\n"
            "**Composite** — Weighted blend: Rule 20% + Judge 55% + RAG Triad 25%.\n\n"
            "---\n\n"
            "**Groundedness** *(Opik-style)* — 1 minus the hallucination score. Only "
            "penalises invented factual claims (names, numbers, certifications).\n\n"
            "**Answer Relevance** *(Opik-style)* — How well the output addresses the "
            "section instruction. One LLM call per section (0–1).\n\n"
            "**Regulatory Tone** *(PMF-specific)* — Language quality for EU GMP Annex 4 "
            "/ ICH Q10 regulatory submission. 1 = exemplary, 0 = inappropriate.\n\n"
            "**Opik Composite** — Mean of Groundedness + Answer Relevance + Regulatory Tone."
        )

    # ── 2. Grade + one-sentence overview ─────────────────────────────────
    _grade_val = _letter_grade(float(composite or rule_score or 0))
    _passed = float(composite or rule_score or 0) >= 65.0
    _site = run_meta.get("site_name", "")
    _health_clr = CLR_SUCCESS if _passed else CLR_DANGER
    _health_label = "Good Quality" if _grade_val in ("A", "B") else (
        "Acceptable" if _grade_val == "C" else "Needs Improvement"
    )
    st.markdown(
        f'<div style="background:{_health_clr}22;border-left:4px solid {_health_clr};'
        f'padding:12px 16px;border-radius:6px;margin:12px 0 16px 0;">'
        f'<span style="font-size:1.1rem;font-weight:700;color:{_health_clr};">'
        f'{_grade_badge_html(_grade_val)} &nbsp; {_health_label}</span><br>'
        f'<span style="color:#374151;">'
        f'Document for <strong>{_site or "unknown site"}</strong> — '
        f'composite score {float(composite or rule_score or 0):.1f}/100.'
        f'</span></div>',
        unsafe_allow_html=True,
    )

    # ── 3. Metric cards ───────────────────────────────────────────────────
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Rule Score", _fmt(rule_score),
              help="Structural checks: min length, required sections, keyword presence")
    c2.metric("Judge Score", _fmt(judge_score) if judge_score is not None else "—",
              help="LLM-as-Judge: 5-criterion rubric (factual accuracy, regulatory language, "
                   "site specificity, completeness, coherence). Score: 0–100.")
    c3.metric("Faithfulness", _fmt(mean_faith, 3) if mean_faith is not None else "—",
              help="DeepEval: fraction of generated claims grounded in source documents (0–1).")
    c4.metric("RAG Triad", _fmt(rag_triad, 3) if rag_triad is not None else "—",
              help="Harmonic mean of Faithfulness + Contextual Precision + Answer Relevancy (0–1).")
    c5.metric("Composite", _fmt(composite))

    # ── 4. Opik-style metric row ──────────────────────────────────────────
    _hallucination = ext.get("mean_hallucination_score")
    _answer_rel = ext.get("mean_answer_relevance_score")
    _reg_tone = ext.get("mean_regulatory_tone_score")
    _opik_composite = ext.get("mean_opik_composite")

    if any(v is not None for v in [_hallucination, _answer_rel, _reg_tone]):
        st.markdown(
            '<p style="font-size:0.78rem;color:#6b7280;margin:12px 0 4px 0;font-weight:600;">'
            'OPIK-STYLE METRICS</p>',
            unsafe_allow_html=True,
        )
        oc1, oc2, oc3, oc4 = st.columns(4)
        oc1.metric("Groundedness",
                   f"{(1.0 - _hallucination):.3f}" if _hallucination is not None else "—",
                   help="1 − Hallucination Score.")
        oc2.metric("Answer Relevance",
                   _fmt(_answer_rel, 3) if _answer_rel is not None else "—",
                   help="How well the output addresses the section instruction (0–1).")
        oc3.metric("Regulatory Tone",
                   _fmt(_reg_tone, 3) if _reg_tone is not None else "—",
                   help="GMP/ICH Q10 language quality (0–1).")
        oc4.metric("Opik Composite",
                   _fmt(_opik_composite, 3) if _opik_composite is not None else "—",
                   help="Mean of Groundedness + Answer Relevance + Regulatory Tone.")

    # ── 5. MLflow link ────────────────────────────────────────────────────
    _run_arts = payload.get("run_artifacts", payload)
    _mlflow_run_id = _run_arts.get("mlflow_run_id")
    _mlflow_url = _run_arts.get("mlflow_ui_url", "http://localhost:5000")
    if _mlflow_run_id:
        st.markdown(
            f'<div style="margin:8px 0 4px 0;">'
            f'<a href="{_mlflow_url}" target="_blank" style="font-size:0.85rem;color:#5340C0;">'
            f'🔗 Open in MLflow UI &nbsp;<span style="color:#9ca3af;font-size:0.75rem;">'
            f'(run: {_mlflow_run_id[:8]}…)</span></a></div>',
            unsafe_allow_html=True,
        )


# ═══════════════════════════════════════════════════════════════════════════
# DOWNLOAD REPORT (DOCX)
# ═══════════════════════════════════════════════════════════════════════════


def _render_download_report(
    payload: Dict[str, Any],
    eval_data: Dict[str, Any],
    ext: Dict[str, Any],
    grade: str,
    rule_score: Any,
) -> None:
    """Generate a formatted DOCX summary report and offer it for download."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        st.caption("Install python-docx to enable DOCX report download.")
        return

    import io

    if not st.button("Download Report (DOCX)", key="dl_report_btn"):
        return

    doc = DocxDocument()

    # ── Title ────────────────────────────────────────────────────────
    title = doc.add_heading("PMF Evaluation Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x53, 0x40, 0xC0)

    # ── Run metadata table ───────────────────────────────────────────
    doc.add_heading("Run Metadata", level=1)
    run_arts = payload.get("run_artifacts", payload)
    meta_items = [
        ("Timestamp", run_arts.get("timestamp", "—")),
        ("Site Name", run_arts.get("site_name", "—")),
        ("Model", run_arts.get("model_name", "—")),
        ("Template", run_arts.get("template_file", "—")),
        ("Grade", grade),
    ]
    tbl = doc.add_table(rows=len(meta_items), cols=2, style="Light Grid Accent 1")
    for i, (label, val) in enumerate(meta_items):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = str(val)

    # ── Metric scores table ──────────────────────────────────────────
    doc.add_heading("Metric Scores", level=1)
    _rag_val = ext.get("mean_rag_triad_score") or ext.get("mean_ragas")
    _hall = ext.get("mean_hallucination_score")
    metrics = [
        ("Rule Score", _fmt(rule_score)),
        ("Judge Score", _fmt(ext.get("mean_judge_normalized"))),
        ("Faithfulness (DeepEval)", _fmt(ext.get("mean_faithfulness"), 4)),
        ("RAG Triad Score (DeepEval)", _fmt(_rag_val, 4)),
        ("Composite", _fmt(ext.get("mean_composite", rule_score))),
        ("Groundedness (Opik)", _fmt((1.0 - _hall) if _hall is not None else None, 4)),
        ("Answer Relevance (Opik)", _fmt(ext.get("mean_answer_relevance_score"), 4)),
        ("Regulatory Tone (Opik)", _fmt(ext.get("mean_regulatory_tone_score"), 4)),
        ("Opik Composite", _fmt(ext.get("mean_opik_composite"), 4)),
        ("Retrieval Coverage", f"{eval_data.get('retrieval_coverage', 0)}%"),
        ("Framework", ext.get("framework", "rule-based")),
    ]
    tbl2 = doc.add_table(rows=len(metrics), cols=2, style="Light Grid Accent 1")
    for i, (label, val) in enumerate(metrics):
        tbl2.rows[i].cells[0].text = label
        tbl2.rows[i].cells[1].text = str(val)

    # ── Top 3 best and worst sections ────────────────────────────────
    sections = eval_data.get("sections", [])
    sorted_secs = sorted(sections, key=lambda s: s.get("score", 0), reverse=True)

    doc.add_heading("Top 3 Best Sections", level=1)
    for s in sorted_secs[:3]:
        doc.add_paragraph(
            f"{s.get('section_key', '?')} — Score: {s.get('score', 0)}",
            style="List Bullet",
        )

    doc.add_heading("Top 3 Worst Sections", level=1)
    for s in sorted_secs[-3:]:
        doc.add_paragraph(
            f"{s.get('section_key', '?')} — Score: {s.get('score', 0)}",
            style="List Bullet",
        )

    # ── Judge findings ───────────────────────────────────────────────
    doc.add_heading("Key Findings from LLM Judge", level=1)
    run_secs = run_arts.get("sections", [])
    all_critical: List[str] = []
    all_strengths: List[str] = []
    all_suggestions: List[str] = []
    for rs in run_secs:
        ext_sec = _section_extended(rs)
        judge = ext_sec.get("judge_scores") or {}
        if judge.get("judge_error"):
            continue
        key = rs.get("section_key", "?")
        for iss in judge.get("critical_issues", []):
            all_critical.append(f"[{key}] {iss}")
        for s in judge.get("strengths", [])[:1]:
            all_strengths.append(f"[{key}] {s}")
        for sg in judge.get("improvement_suggestions", [])[:1]:
            all_suggestions.append(f"[{key}] {sg}")

    if all_critical:
        doc.add_heading("Critical Issues", level=2)
        for item in all_critical[:10]:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("No critical issues identified.", style="List Bullet")

    if all_strengths:
        doc.add_heading("Top Strengths", level=2)
        for item in all_strengths[:5]:
            doc.add_paragraph(item, style="List Bullet")

    # ── Recommendations ──────────────────────────────────────────────
    doc.add_heading("Recommendations", level=1)
    if all_suggestions:
        for item in all_suggestions[:5]:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph(
            "Run the extended evaluation with the LLM judge enabled to "
            "generate specific improvement recommendations.",
        )

    # ── Footer ───────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"Generated by Healthark GenAI Evaluation Framework v1.0 "
        f"on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Serve as download ────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button(
        "Save Report",
        buf.getvalue(),
        file_name=f"eval_report_{run_arts.get('timestamp', 'unknown')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="dl_report_file",
    )


# ═══════════════════════════════════════════════════════════════════════════
# LIVE EVALUATION (CEO demo feature)
# ═══════════════════════════════════════════════════════════════════════════


def _render_live_evaluation() -> None:
    """Smart evaluation panel — rule-based runs automatically, reference metrics optional."""
    st.subheader("Quick Text Evaluation")
    st.caption(
        "Paste any generated PMF text to evaluate it instantly. "
        "**Rule-based checks run without any reference text.** "
        "Add a reference text on the right to also compute BLEU / ROUGE / BERTScore."
    )

    # ── Auto-populate from the last generated run ─────────────────────
    auto_gen = ""
    auto_sec = "EXECUTIVE SUMMARY"
    _runs = _load_runs()
    if _runs:
        _rf = _runs[0].get("run_file", "")
        if _rf and os.path.exists(_rf):
            try:
                _payload = _load_run_payload(_rf)
                for _sec in _payload.get("run_artifacts", {}).get("sections", []):
                    _gt = _sec.get("generated_text", "")
                    if _gt.strip() and "static" not in _sec.get("section_key", "").lower():
                        auto_gen = _gt[:2500]
                        auto_sec = _sec.get("section_key", auto_sec)
                        break
            except Exception:
                pass

    lc1, lc2 = st.columns(2)
    with lc1:
        gen_text = st.text_area(
            "Generated text",
            value=auto_gen,
            height=200,
            key="live_gen_txt",
            help="Auto-populated from your last run. Edit or paste any text to evaluate.",
        )
    with lc2:
        ref_text = st.text_area(
            "Reference text (optional)",
            height=200,
            key="live_ref_txt",
            help=(
                "Leave blank — rule-based and judge metrics work without it. "
                "Paste a ground-truth PMF section here to also compute BLEU / ROUGE / BERTScore."
            ),
        )

    live_section = st.text_input(
        "Section type", value=auto_sec, key="live_sec_key",
        help="E.g. EXECUTIVE SUMMARY, SITE DESCRIPTION, MANUFACTURING PROCESSES",
    )

    if not st.button("Evaluate Now", key="live_eval_now_btn", type="primary"):
        return

    if not gen_text.strip():
        st.warning("Paste some generated text first.")
        return

    has_ref = bool(ref_text.strip())

    # ── Phase 1: Rule-based checks (always — no reference needed) ─────
    rule_score = 0.0
    rule_checks: Dict[str, Any] = {}
    missing_kw: List[str] = []
    with st.spinner("Running rule-based quality checks..."):
        try:
            _rule_result = score_section(
                section_key=live_section,
                section_text=gen_text,
                rules=get_eval_rules(),
                context={},
            )
            rule_score = float(_rule_result.get("score", 0))
            rule_checks = _rule_result.get("checks", {})
            missing_kw = _rule_result.get("missing_keywords", [])
        except Exception as exc:
            logger.warning("Rule scoring failed in live eval: %s", exc)

    # ── Phase 2: Lexical metrics (BLEU / ROUGE) — needs reference ─────
    lex_result: Dict[str, Any] = {}
    if HAS_LEXICAL and has_ref:
        with st.spinner("Computing BLEU / ROUGE..."):
            try:
                lex_result = LexicalMetrics.compute_all_lexical(gen_text, ref_text)
            except Exception as exc:
                logger.warning("Lexical metrics failed in live eval: %s", exc)

    # ── Phase 3: BERTScore — needs reference + semantic toggle ────────
    bert_f1: Any = None
    sim_result: float = 0.0
    if st.session_state.get("toggle_semantic", False) and HAS_SEMANTIC and has_ref:
        with st.spinner("Computing BERTScore (this may take a few seconds)..."):
            try:
                sm = SemanticMetrics(model_type="distilbert-base-uncased")
                bs = sm.compute_bertscore([gen_text], [ref_text])
                bert_f1 = bs.get("bertscore_f1_mean")
                sim_result = sm.compute_semantic_similarity(gen_text, ref_text)
            except Exception as exc:
                logger.warning("BERTScore failed in live eval: %s", exc)

    # ── Phase 4: LLM Judge (if enabled, ~30-60s) ─────────────────────
    judge_result: Dict[str, Any] = {}
    if st.session_state.get("toggle_judge", False):
        with st.spinner("Running LLM Judge (30-60 s)..."):
            try:
                from src.eval.eval_judge import PMFJudge
                from dotenv import load_dotenv
                load_dotenv()
                # Prefer Azure OpenAI (already configured) over Anthropic
                _azure_key = os.getenv("AZURE_KEY", os.getenv("AZURE_OPENAI_API_KEY", ""))
                _azure_ep = os.getenv("AZURE_ENDPOINT", os.getenv("AZURE_OPENAI_ENDPOINT", ""))
                _azure_ver = os.getenv("AZURE_VERSION", os.getenv("AZURE_OPENAI_API_VERSION", "2024-06-01"))
                _azure_name = os.getenv("AZURE_NAME", "gpt-4o")

                if _azure_key and _azure_ep:
                    judge = PMFJudge(
                        provider="azure_openai",
                        model=_azure_name,
                        api_key=_azure_key,
                        azure_endpoint=_azure_ep,
                        azure_api_version=_azure_ver,
                        cache_enabled=False,
                    )
                else:
                    # Fallback to sidebar Anthropic key
                    api_key = st.session_state.get("sidebar_anthropic_key", "")
                    model = st.session_state.get("sidebar_judge_model", "claude-sonnet-4-6")
                    judge = PMFJudge(provider="anthropic", model=model, api_key=api_key, cache_enabled=False)

                judge_result = judge.score_section(
                    section_key=live_section,
                    section_instruction="Evaluate this PMF section for regulatory quality.",
                    retrieved_context="",
                    generated_output=gen_text,
                    site_name="",
                    reference_output=ref_text,
                )
            except Exception as exc:
                logger.warning("Judge failed in live eval: %s", exc)
                judge_result = {"judge_error": True, "error": str(exc)}

    # ── Phase 5: DeepEval RAG Triad (if enabled) ──────────────────────
    rag_result: Dict[str, Any] = {}
    if st.session_state.get("toggle_rag", False):
        with st.spinner("Running DeepEval RAG Triad (Faithfulness + Contextual Precision + Answer Relevancy)..."):
            try:
                from src.eval.eval_rag import RAGEvaluator
                from openai import AzureOpenAI
                from dotenv import load_dotenv
                load_dotenv()
                _azure_key = os.getenv("AZURE_KEY", os.getenv("AZURE_OPENAI_API_KEY", ""))
                _azure_ep = os.getenv("AZURE_ENDPOINT", os.getenv("AZURE_OPENAI_ENDPOINT", ""))
                _azure_ver = os.getenv("AZURE_VERSION", os.getenv("AZURE_OPENAI_API_VERSION", "2024-06-01"))
                _azure_name = os.getenv("AZURE_NAME", "gpt-4o")

                _client = AzureOpenAI(api_key=_azure_key, api_version=_azure_ver, azure_endpoint=_azure_ep)
                rag_ev = RAGEvaluator(llm_client=_client, model=_azure_name, cache_enabled=False)
                rag_result = rag_ev.evaluate_section(
                    section_key=live_section,
                    section_instruction=live_section,
                    retrieved_chunks=[ref_text] if ref_text.strip() else [],
                    generated_answer=gen_text,
                )
            except Exception as exc:
                logger.warning("RAG Triad eval failed in live eval: %s", exc)
                rag_result = {"error": str(exc)}

    # ── Results display ───────────────────────────────────────────────
    st.subheader("Evaluation Results")

    grade = _letter_grade(rule_score)
    badge_html = _grade_badge_html(grade)
    passed = rule_score >= 65.0
    status = "PASS" if passed else "NEEDS IMPROVEMENT"
    status_clr = CLR_SUCCESS if passed else CLR_WARNING
    st.markdown(
        f"{badge_html} &nbsp;"
        f'<span style="color:{status_clr};font-weight:700;">{status}</span>'
        f"&nbsp; Rule Score: <strong>{_fmt(rule_score)}/100</strong>",
        unsafe_allow_html=True,
    )

    st.write("")
    m1, m2, m3, m4, m5 = st.columns(5)
    m1.metric("Rule Score", _fmt(rule_score),
              help="Structural quality: min length, required sections, site name presence")
    m2.metric("BLEU", _fmt(lex_result.get("bleu"), 2) if lex_result else "N/A",
              help="Lexical overlap with reference (0–100). Needs reference text.")
    m3.metric("ROUGE-1 F1", _fmt(lex_result.get("rouge1_fmeasure"), 4) if lex_result else "N/A",
              help="Unigram overlap F1 with reference. Needs reference text.")
    _faith_live = rag_result.get("faithfulness")
    _rt_live = rag_result.get("rag_triad_score") or rag_result.get("ragas_score")
    m4.metric("Faithfulness",
              _fmt(_faith_live, 3) if _faith_live is not None else "N/A",
              help="DeepEval: claims grounded in source context. Enable 'DeepEval RAG Triad' in sidebar.")
    m5.metric("RAG Triad",
              _fmt(_rt_live, 3) if _rt_live is not None else "N/A",
              help="DeepEval composite: Faithfulness + Contextual Precision + Answer Relevancy")

    # Rule check breakdown
    if rule_checks:
        with st.expander("Rule Check Details", expanded=True):
            for check_name, check_passed in rule_checks.items():
                icon = "✓" if check_passed else "✗"
                clr = CLR_SUCCESS if check_passed else CLR_DANGER
                label = check_name.replace("_", " ").replace("passed", "").strip().title()
                st.markdown(
                    f'<span style="color:{clr};font-weight:600;">{icon} {label}</span>',
                    unsafe_allow_html=True,
                )
        if missing_kw:
            st.warning(f"Missing keywords: {', '.join(missing_kw)}")

    if not has_ref:
        st.info(
            "**BLEU / ROUGE / BERTScore** require a reference text — paste a ground-truth "
            "PMF section in the right box to compute them. "
            "The rule-based score above works without any reference."
        )

    if not st.session_state.get("toggle_semantic") and has_ref and HAS_SEMANTIC:
        st.caption(
            "BERTScore is available but disabled. "
            "Enable 'Run Semantic (BERTScore)' in the sidebar to compute it."
        )

    # ── Judge results ─────────────────────────────────────────────────
    if judge_result and not judge_result.get("judge_error"):
        st.subheader("LLM Judge Assessment")
        jc1, jc2, jc3, jc4, jc5 = st.columns(5)
        scores = judge_result.get("scores", {})
        jc1.metric("Factual Accuracy", scores.get("factual_accuracy", "—"))
        jc2.metric("Regulatory Lang.", scores.get("regulatory_language", "—"))
        jc3.metric("Site Specificity", scores.get("site_specificity", "—"))
        jc4.metric("Completeness", scores.get("completeness", "—"))
        jc5.metric("Coherence", scores.get("structural_coherence", "—"))

        norm = judge_result.get("normalized_score", 0)
        j_grade = _letter_grade(float(norm))
        st.markdown(
            f"**Judge Score:** {_fmt(norm)} / 100 &nbsp; {_grade_badge_html(j_grade)}",
            unsafe_allow_html=True,
        )

        with st.expander("Judge Details"):
            st.markdown(f"**Strengths:** {', '.join(judge_result.get('strengths', [])) or '—'}")
            st.markdown(f"**Weaknesses:** {', '.join(judge_result.get('weaknesses', [])) or '—'}")
            critical = judge_result.get("critical_issues", [])
            if critical:
                st.error(f"Critical Issues: {', '.join(critical)}")
            st.markdown(f"**Notes:** {judge_result.get('evaluation_notes', '—')}")

    elif judge_result.get("judge_error"):
        st.warning(
            f"Judge evaluation failed: {judge_result.get('error', 'unknown')}. "
            "Check Azure OpenAI credentials in .env file."
        )

    # ── DeepEval RAG Triad results ────────────────────────────────────
    if rag_result and not rag_result.get("error"):
        st.subheader("DeepEval RAG Triad")
        rc1, rc2, rc3, rc4 = st.columns(4)
        rc1.metric("Faithfulness", _fmt(rag_result.get("faithfulness"), 3),
                   help="Claims grounded in context (0–1)")
        rc2.metric("Ctx Precision", _fmt(rag_result.get("contextual_precision") or rag_result.get("context_precision"), 3),
                   help="Relevant contexts ranked first (0–1)")
        rc3.metric("Answer Relevancy", _fmt(rag_result.get("answer_relevancy"), 3),
                   help="Answer addresses the question (0–1)")
        _rt = rag_result.get("rag_triad_score") or rag_result.get("ragas_score")
        rc4.metric("RAG Triad", _fmt(_rt, 3), help="Harmonic mean of the three metrics (0–1)")

        with st.expander("RAG Triad Details"):
            claims = rag_result.get("faithfulness_claims", [])
            if claims:
                st.markdown(f"**Claims extracted:** {len(claims)}")
                for c in claims[:5]:
                    icon = "✓" if c.get("supported") else "✗"
                    clr = CLR_SUCCESS if c.get("supported") else CLR_DANGER
                    st.markdown(
                        f'<span style="color:{clr};">{icon}</span> {c.get("claim", "?")}',
                        unsafe_allow_html=True,
                    )
            gen_qs = rag_result.get("generated_questions", [])
            if gen_qs:
                st.markdown(f"**Generated questions (Answer Relevancy):** {', '.join(gen_qs)}")
    elif rag_result.get("error"):
        st.warning(f"DeepEval RAG Triad failed: {rag_result.get('error')}. Check Azure credentials.")

    # ── Full JSON (expandable) ────────────────────────────────────────
    with st.expander("Full Results (JSON)"):
        st.json({
            "rule_score": rule_score,
            "rule_checks": rule_checks,
            "lexical": lex_result if lex_result else "N/A — no reference text",
            "deepeval_rag_triad": rag_result if rag_result else None,
            "judge": judge_result if judge_result else None,
        })


# ═══════════════════════════════════════════════════════════════════════════
# TAB 2 — SECTION HEATMAP
# ═══════════════════════════════════════════════════════════════════════════

def _render_tab_heatmap(runs: List[Dict[str, Any]]) -> None:
    """Multi-metric heatmap (rows=sections, cols=metrics) with detail view."""

    if not runs:
        st.info("No runs available.")
        return

    # Always use the latest run — no selector
    run_file = runs[0].get("run_file", "")
    if not run_file or not os.path.exists(run_file):
        return

    payload = _load_run_payload(run_file)
    eval_data = payload.get("evaluation", {}).get("document_scores", {})
    sections = eval_data.get("sections", [])
    run_sections = payload.get("run_artifacts", payload).get("sections", [])

    if not sections:
        st.info("No section data available.")
        return

    # Build a lookup from section_key to the run_artifacts section
    run_sec_map: Dict[str, Dict[str, Any]] = {}
    for rs in run_sections:
        run_sec_map[rs.get("section_key", "")] = rs

    # ── Build heatmap DataFrame ──────────────────────────────────────────
    heatmap_rows: List[Dict[str, Any]] = []
    for s in sections:
        key = s.get("section_key", "?")
        ext = _section_extended(run_sec_map.get(key, {}))
        lex = ext.get("lexical_scores") or {}
        sem = ext.get("semantic_scores") or {}
        judge = ext.get("judge_scores") or {}
        rag = ext.get("rag_scores") or {}

        heatmap_rows.append({
            "Section": key[:35],
            "Rule Score": s.get("score", 0),
            "Judge Score": judge.get("normalized_score"),
            "Faithfulness": _safe_pct(rag.get("faithfulness")),
            "Ctx Precision": _safe_pct(rag.get("contextual_precision") or rag.get("context_precision")),
            "RAG Triad": _safe_pct(rag.get("rag_triad_score") or rag.get("ragas_score")),
        })

    df = pd.DataFrame(heatmap_rows).set_index("Section")

    # Replace None with NaN for Styler
    df = df.where(df.notna(), other=float("nan"))

    metric_cols = ["Rule Score", "Judge Score", "Faithfulness", "Ctx Precision", "RAG Triad"]

    st.subheader("Section Heatmap")
    st.caption("Green (>80) | Yellow (50-80) | Red (<50). Blank = metric not computed.")

    styled = df[metric_cols].style.background_gradient(
        cmap="RdYlGn", vmin=0, vmax=100, axis=None,
    ).format("{:.1f}", na_rep="—")

    st.dataframe(styled, use_container_width=True, height=min(len(df) * 40 + 60, 600))

    # ── Section detail drill-down ────────────────────────────────────────
    st.subheader("Section Detail")
    sec_names = [s.get("section_key", "?") for s in sections]
    sel_sec = st.selectbox("Select section", sec_names, key="tab2_sec_detail")

    rule_sec = next((s for s in sections if s.get("section_key") == sel_sec), {})
    run_sec = run_sec_map.get(sel_sec, {})
    ext_sec = _section_extended(run_sec)

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("**Section Instruction**")
        st.text(run_sec.get("prompt_text", "—")[:500])

        with st.expander("Retrieved Context"):
            paths = run_sec.get("retrieved_paths", [])
            if paths:
                st.write(f"{len(paths)} documents retrieved: {paths}")
            else:
                st.write("No retrieval data available.")

    with col_b:
        st.markdown("**Generated Text**")
        gen_text = run_sec.get("generated_text", "")
        st.text(gen_text[:800] if gen_text else "—")

    # Metric scores for this section
    st.markdown("**Metric Scores**")
    lex = ext_sec.get("lexical_scores") or {}
    sem = ext_sec.get("semantic_scores") or {}
    judge = ext_sec.get("judge_scores") or {}
    rag = ext_sec.get("rag_scores") or {}

    sc1, sc2, sc3, sc4, sc5 = st.columns(5)
    sc1.metric("Rule", _fmt(rule_sec.get("score")))
    sc2.metric("Judge Score", _fmt(judge.get("normalized_score")))
    sc3.metric("Faithfulness", _fmt(rag.get("faithfulness"), 4))
    _cp = rag.get("contextual_precision") or rag.get("context_precision")
    sc4.metric("Ctx Precision", _fmt(_cp, 4))
    _rt = rag.get("rag_triad_score") or rag.get("ragas_score")
    sc5.metric("RAG Triad", _fmt(_rt, 4))

    # Judge qualitative output
    if judge and not judge.get("judge_error"):
        with st.expander("Judge Output"):
            st.markdown(f"**Strengths:** {', '.join(judge.get('strengths', [])) or '—'}")
            st.markdown(f"**Weaknesses:** {', '.join(judge.get('weaknesses', [])) or '—'}")
            critical = judge.get("critical_issues", [])
            if critical:
                st.error(f"Critical Issues: {', '.join(critical)}")
            st.markdown(f"**Notes:** {judge.get('evaluation_notes', '—')}")


def _safe_pct(val: Any) -> Optional[float]:
    """Convert a 0-1 metric to 0-100, or return None."""
    if val is None:
        return None
    try:
        v = float(val)
        return round(v * 100, 2) if v <= 1.0 else round(v, 2)
    except (ValueError, TypeError):
        return None


# ═══════════════════════════════════════════════════════════════════════════
# TAB 3 — TREND ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════

def _render_tab_trends(runs: List[Dict[str, Any]]) -> None:
    """Time-series charts and regression alerts."""

    if len(runs) < 2:
        st.info("Need at least 2 runs for trend analysis.")
        return

    df = pd.DataFrame(runs)
    df["overall_score"] = pd.to_numeric(df.get("overall_score"), errors="coerce")
    df = df.sort_values("timestamp").reset_index(drop=True)

    # ── Regression alert ─────────────────────────────────────────────────
    latest = df.iloc[-1]["overall_score"]
    prev = df.iloc[-2]["overall_score"]
    if pd.notna(latest) and pd.notna(prev):
        drop = prev - latest
        if drop > 5:
            latest_ts = df.iloc[-1].get("timestamp", "?")
            st.markdown(
                f'<div style="background:{CLR_DANGER};color:white;padding:12px;'
                f'border-radius:8px;margin-bottom:16px;">'
                f'<strong>Regression detected:</strong> Rule score dropped from '
                f'{prev:.1f} to {latest:.1f} (delta {-drop:+.1f}) on {latest_ts}'
                f'</div>',
                unsafe_allow_html=True,
            )

    # ── Composite score over time ────────────────────────────────────────
    st.subheader("Score Trend Over Time")

    if HAS_PLOTLY:
        fig = px.line(df, x="timestamp", y="overall_score", markers=True,
                      labels={"overall_score": "Rule Score (%)", "timestamp": "Run"})
        fig.update_traces(line_color=CLR_PRIMARY)
        fig.update_layout(height=320, margin=dict(l=0, r=0, t=10, b=0))
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.line_chart(df.set_index("timestamp")["overall_score"])

    # ── Multi-metric trend ───────────────────────────────────────────────
    # Load extended data from each run payload to build multi-metric trend
    multi_rows: List[Dict[str, Any]] = []
    for _, row in df.iterrows():
        rf = row.get("run_file", "")
        if rf and os.path.exists(rf):
            p = _load_run_payload(rf)
            ext = _extract_extended(p)
            ed = p.get("evaluation", {}).get("document_scores", {})
            multi_rows.append({
                "timestamp": row["timestamp"],
                "Rule Score": ed.get("overall_score"),
                "BERTScore F1": _safe_pct(ext.get("mean_bertscore_f1")),
                "Judge Score": ext.get("mean_judge_normalized"),
                "Faithfulness": _safe_pct(ext.get("mean_faithfulness")),
                "RAGAS": _safe_pct(ext.get("mean_ragas")),
            })

    if multi_rows:
        mdf = pd.DataFrame(multi_rows).set_index("timestamp")
        # Only show columns that have at least one non-null value
        active_cols = [c for c in mdf.columns if mdf[c].notna().any()]
        if active_cols:
            st.subheader("Multi-Metric Trend")
            if HAS_PLOTLY:
                fig = go.Figure()
                colors = [CLR_PRIMARY, CLR_SUCCESS, "#3B82F6", CLR_WARNING, CLR_DANGER]
                for i, col in enumerate(active_cols):
                    fig.add_trace(go.Scatter(
                        x=mdf.index, y=mdf[col], name=col, mode="lines+markers",
                        line=dict(color=colors[i % len(colors)]),
                    ))
                fig.update_layout(height=350, margin=dict(l=0, r=0, t=10, b=0),
                                  yaxis_title="Score (0-100)")
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.line_chart(mdf[active_cols])

    # ── Section score distribution (histogram) ───────────────────────────
    st.subheader("Section Score Distribution (Latest Run)")
    latest_file = df.iloc[-1].get("run_file", "")
    if latest_file and os.path.exists(latest_file):
        lp = _load_run_payload(latest_file)
        secs = lp.get("evaluation", {}).get("document_scores", {}).get("sections", [])
        sec_scores = [s.get("score", 0) for s in secs]
        if sec_scores:
            if HAS_PLOTLY:
                fig = px.histogram(
                    x=sec_scores, nbins=10,
                    labels={"x": "Section Rule Score", "count": "Count"},
                    color_discrete_sequence=[CLR_PRIMARY],
                )
                fig.update_layout(height=300, margin=dict(l=0, r=0, t=10, b=0))
                st.plotly_chart(fig, use_container_width=True)
            else:
                hist_df = pd.DataFrame({"Score": sec_scores})
                st.bar_chart(hist_df["Score"].value_counts().sort_index())


# ═══════════════════════════════════════════════════════════════════════════
# TAB 4 — MODEL COMPARISON (unchanged from previous)
# ═══════════════════════════════════════════════════════════════════════════

def _render_tab_model_comparison() -> None:
    """Upload JSON results and compare models."""
    st.subheader("Model Comparison")

    uploaded = st.file_uploader(
        "Upload evaluation result JSON files (one per model run)",
        type=["json"], accept_multiple_files=True, key="mc_upload",
    )
    if not uploaded:
        st.info("Upload two or more JSON result files to compare models.")
        return

    model_data: List[Dict[str, Any]] = []
    for f in uploaded:
        try:
            data = json.loads(f.read().decode("utf-8"))
            data["_source_file"] = f.name
            model_data.append(data)
        except (json.JSONDecodeError, UnicodeDecodeError) as exc:
            st.warning(f"Could not parse {f.name}: {exc}")

    if len(model_data) < 2:
        st.warning("Upload at least 2 files to compare.")
        return

    rows: List[Dict[str, Any]] = []
    for d in model_data:
        label = d.get("_source_file", "unknown")
        judge = d.get("judge_scores") or {}
        rag = d.get("rag_scores") or {}
        lex = d.get("lexical_scores") or {}
        sem = d.get("semantic_scores") or {}
        rows.append({
            "Model / File": label,
            "Section": d.get("section_key", "—"),
            "Composite": d.get("composite_score", 0),
            "Grade": d.get("grade", "?"),
            "Rule Score": d.get("rule_score"),
            "BLEU": lex.get("bleu"),
            "BERTScore F1": sem.get("bertscore_f1_mean"),
            "Judge Norm.": judge.get("normalized_score"),
            "RAGAS": rag.get("ragas_score"),
        })

    cdf = pd.DataFrame(rows)
    st.dataframe(
        cdf.style.background_gradient(subset=["Composite"], cmap="RdYlGn",
                                       vmin=0, vmax=100),
        use_container_width=True, hide_index=True,
    )

    # Radar chart
    st.subheader("Metric Radar")
    radar_metrics = ["Rule Score", "BLEU", "BERTScore F1", "Judge Norm.", "RAGAS"]
    if HAS_PLOTLY:
        fig = go.Figure()
        for _, row in cdf.iterrows():
            vals = []
            for m in radar_metrics:
                v = row.get(m)
                if v is None:
                    vals.append(0)
                elif m in ("BERTScore F1", "RAGAS"):
                    vals.append(float(v) * 100)
                else:
                    vals.append(float(v))
            fig.add_trace(go.Scatterpolar(
                r=vals + [vals[0]], theta=radar_metrics + [radar_metrics[0]],
                fill="toself", name=str(row.get("Model / File", "")), opacity=0.6,
            ))
        fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                          height=400, margin=dict(l=40, r=40, t=20, b=20))
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.bar_chart(cdf[["Model / File"] + radar_metrics].set_index("Model / File"))

    # Cost analysis
    st.subheader("Cost Analysis")
    cost_entries: Dict[str, float] = {}
    cost_cols = st.columns(min(len(model_data), 4))
    for i, d in enumerate(model_data):
        label = d.get("_source_file", f"Model {i+1}")
        with cost_cols[i % len(cost_cols)]:
            cost = st.number_input(f"Cost ($) — {label}", min_value=0.0,
                                   value=0.0, step=0.01, key=f"cost_{i}")
            cost_entries[label] = cost

    if any(c > 0 for c in cost_entries.values()):
        st.subheader("Pareto Recommendation")
        pareto_rows = []
        for _, row in cdf.iterrows():
            label = row["Model / File"]
            cost = cost_entries.get(label, 0)
            comp = row.get("Composite", 0) or 0
            eff = round(comp / cost, 2) if cost > 0 else 0
            pareto_rows.append({"Model": label, "Composite": comp,
                                "Cost ($)": cost, "Score / $": eff})
        pdf = pd.DataFrame(pareto_rows).sort_values("Score / $", ascending=False)
        st.dataframe(pdf, use_container_width=True, hide_index=True)
        best = pdf.iloc[0]
        st.success(
            f"Recommended: **{best['Model']}** — "
            f"Composite {best['Composite']:.1f} at ${best['Cost ($)']:.2f} "
            f"({best['Score / $']:.1f} points per dollar)"
        )


# ═══════════════════════════════════════════════════════════════════════════
# TAB 5 — BENCHMARK MANAGEMENT (unchanged from previous)
# ═══════════════════════════════════════════════════════════════════════════

def _render_tab_benchmark() -> None:
    """Benchmark management interface."""
    if not HAS_BENCHMARK:
        st.info("Benchmark loader not available.")
        return

    st.subheader("Benchmark Management")
    loader = BenchmarkLoader(BENCHMARK_DIR)
    stats = loader.get_statistics()

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Cases", stats.get("total_cases", 0))
    c2.metric("With Reference", stats.get("cases_with_reference", 0))
    c3.metric("Expert Annotated", stats.get("cases_with_expert_scores", 0))
    c4.metric("Auto Scored", stats.get("cases_with_automated_scores", 0))

    # Filterable table
    st.subheader("Browse Cases")
    fc1, fc2, fc3 = st.columns(3)
    with fc1:
        ft = st.selectbox("Section Type",
                          ["All"] + list(stats.get("by_section_type", {}).keys()),
                          key="bm_ft")
    with fc2:
        fd = st.selectbox("Difficulty",
                          ["All"] + list(stats.get("by_difficulty", {}).keys()),
                          key="bm_fd")
    with fc3:
        fs = st.text_input("Search section key", "", key="bm_fs")

    filters: Dict[str, Any] = {}
    if ft != "All":
        filters["section_type"] = ft
    if fd != "All":
        filters["difficulty"] = fd
    if fs.strip():
        filters["section_key"] = fs.strip()

    cases = loader.load_cases(filters if filters else None)
    if cases:
        rows = [{
            "Case ID": c.get("case_id", ""),
            "Section": c.get("section_key", ""),
            "Site": c.get("site_name", ""),
            "Type": c.get("section_type", ""),
            "Difficulty": c.get("difficulty", ""),
            "Ref Len": len(c.get("reference_output", "")),
            "Tags": ", ".join(c.get("tags", [])),
        } for c in cases]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
    else:
        st.info("No cases match filters.")

    # Add case form
    st.subheader("Add New Case")
    with st.expander("New benchmark case form"):
        nsk = st.text_input("Section Key", key="bn_sk")
        nsite = st.text_input("Site Name", key="bn_site")
        ninstr = st.text_area("Section Instruction", height=80, key="bn_instr")
        nquery = st.text_input("Retrieval Query", key="bn_query")
        nref = st.text_area("Reference Output", height=120, key="bn_ref")
        ntype = st.selectbox("Type", ["text", "table", "image", "static"], key="bn_type")
        ndiff = st.selectbox("Difficulty", ["easy", "medium", "hard"], key="bn_diff")
        ntags = st.text_input("Tags (comma-separated)", key="bn_tags")
        if st.button("Add Case", key="bn_add"):
            if not nsk.strip() or not nref.strip():
                st.warning("Section Key and Reference Output are required.")
            else:
                new_case: Dict[str, Any] = {
                    "created_at": datetime.now().isoformat(),
                    "created_by": "human", "validated_by": None,
                    "site_name": nsite, "section_key": nsk,
                    "section_instruction": ninstr, "retrieval_query": nquery,
                    "source_documents": [], "retrieved_context": "",
                    "generated_output": {}, "reference_output": nref,
                    "expert_scores": {}, "automated_scores": {},
                    "tags": [t.strip() for t in ntags.split(",") if t.strip()],
                    "difficulty": ndiff, "section_type": ntype,
                }
                try:
                    cid = loader.add_case(new_case, validate=True)
                    st.success(f"Case **{cid}** added.")
                except ValueError as exc:
                    st.error(f"Validation failed: {exc}")

    # Export + Run
    st.subheader("Export & Run")
    ec1, ec2 = st.columns(2)
    with ec1:
        if st.button("Export to CSV", key="bn_csv"):
            csv_path = os.path.join(BENCHMARK_DIR, "benchmark_export.csv")
            loader.export_to_csv(csv_path)
            st.success(f"Exported to `{csv_path}`")
            with open(csv_path, "r", encoding="utf-8") as fh:
                st.download_button("Download CSV", fh.read(),
                                   file_name="benchmark_export.csv",
                                   mime="text/csv", key="bn_dl")
    with ec2:
        do_run = st.button("Run Full Benchmark", key="bn_run")
    if do_run:
        all_cases = loader.load_cases()
        if not all_cases:
            st.warning("No cases.")
            return
        try:
            from healthark_eval import EvalSuite
        except ImportError:
            st.error("healthark_eval not available.")
            return
        suite = EvalSuite(task="pmf", run_judge=False, run_rag=False,
                          run_semantic=False)
        prog = st.progress(0, text="Evaluating...")
        res: List[Dict[str, Any]] = []
        for i, case in enumerate(all_cases):
            ref = case.get("reference_output", "")
            r = suite.run(generated=ref, reference=ref,
                          section_key=case.get("section_key", ""),
                          section_instruction=case.get("section_instruction", ""),
                          site_name=case.get("site_name", ""))
            res.append({"Case ID": case.get("case_id"), "Section": case.get("section_key"),
                         "Composite": r.composite_score, "Grade": r.grade,
                         "Rule Score": r.rule_score})
            prog.progress((i + 1) / len(all_cases),
                          text=f"Evaluated {i+1}/{len(all_cases)}...")
        prog.empty()
        st.success(f"Benchmark complete: {len(res)} cases.")
        rdf = pd.DataFrame(res)
        st.dataframe(rdf, use_container_width=True, hide_index=True)
        st.metric("Mean Composite", f"{rdf['Composite'].mean():.1f}")


# ═══════════════════════════════════════════════════════════════════════════
# TAB 6 — PERFORMANCE (latency · failures · improvements)
# ═══════════════════════════════════════════════════════════════════════════

def _render_tab_performance(runs: List[Dict[str, Any]]) -> None:
    """Latency breakdown and issues for the current run."""

    if not runs:
        st.info("No runs available. Generate a PMF document first.")
        return

    # Always use the latest run — no selector
    payload = _load_run_payload(runs[0].get("run_file", ""))
    run_arts = payload.get("run_artifacts", payload)
    perf = run_arts.get("performance_report", {})

    st.divider()

    # ── Summary banner ────────────────────────────────────────────────
    summary = perf.get("summary_plain", "") or perf.get("summary_technical", "")
    if not summary:
        st.info("No performance data available for this run. Performance tracking is captured from the next document generation onwards.")
        return

    st.info(summary)

    # ── Overall timing cards ──────────────────────────────────────────
    ot = perf.get("overall_timing", {})
    total_s = (ot.get("total_pipeline_ms") or 0) / 1000
    gen_s   = (ot.get("total_generation_ms") or 0) / 1000
    ret_s   = (ot.get("total_retrieval_ms") or 0) / 1000
    eval_s  = (ot.get("total_eval_ms") or 0) / 1000
    avg_s   = (ot.get("avg_section_ms") or 0) / 1000

    st.subheader("Overall Timing")
    tc1, tc2, tc3, tc4, tc5 = st.columns(5)
    tc1.metric("Total Time",     f"{total_s:.1f}s" if total_s else "—")
    tc2.metric("LLM Generation", f"{gen_s:.1f}s"   if gen_s  else "—")
    tc3.metric("Retrieval",      f"{ret_s:.1f}s"   if ret_s  else "—")
    tc4.metric("Evaluation",     f"{eval_s:.1f}s"  if eval_s else "—")
    tc5.metric("Avg / Section",  f"{avg_s:.1f}s"   if avg_s  else "—")

    # ── Time breakdown donut chart ────────────────────────────────────
    if HAS_PLOTLY and gen_s + ret_s + eval_s > 0:
        other_s = max(0, total_s - gen_s - ret_s - eval_s)
        fig_donut = go.Figure(go.Pie(
            labels=["LLM Generation", "Retrieval", "Evaluation", "Other"],
            values=[gen_s, ret_s, eval_s, other_s],
            hole=0.55,
            marker_colors=["#5340C0", "#1D9E75", "#BA7517", "#d1d5db"],
            textinfo="label+percent",
            hovertemplate="%{label}: %{value:.1f}s (%{percent})<extra></extra>",
        ))
        fig_donut.update_layout(
            margin=dict(t=30, b=10, l=10, r=10), height=260,
            showlegend=False, title_text="Time Breakdown", title_x=0.5,
        )
        st.plotly_chart(fig_donut, use_container_width=True)

    st.divider()

    # ── Per-section latency bar chart (total time only) ───────────────
    st.subheader("Section Latency")
    section_timings = perf.get("section_timings", [])
    if section_timings:
        df_timing = pd.DataFrame([
            {
                "Section":        t["section_key"][:45],
                "Retrieval (s)":  round((t.get("retrieval_ms") or 0) / 1000, 2),
                "Generation (s)": round((t.get("generation_ms") or 0) / 1000, 2),
                "Evaluation (s)": round((t.get("eval_ms") or 0) / 1000, 2),
                "Total (s)":      round((t.get("total_ms") or 0) / 1000, 2),
            }
            for t in section_timings
        ]).sort_values("Total (s)", ascending=False)

        if HAS_PLOTLY:
            fig_bar = px.bar(
                df_timing, x="Total (s)", y="Section", orientation="h",
                color="Total (s)",
                color_continuous_scale=["#1D9E75", "#BA7517", "#D85A30"],
                labels={"Total (s)": "Time (seconds)", "Section": ""},
                title="Time per Section (slowest first)",
                height=max(300, len(df_timing) * 28),
            )
            fig_bar.update_coloraxes(showscale=False)
            fig_bar.update_layout(margin=dict(t=40, b=20, l=10, r=20), yaxis_autorange="reversed")
            st.plotly_chart(fig_bar, use_container_width=True)
        else:
            st.dataframe(df_timing[["Section", "Retrieval (s)", "Generation (s)", "Total (s)"]],
                         use_container_width=True, hide_index=True)

        slowest = ot.get("slowest_section")
        if slowest:
            st.caption(f"Slowest section: **{slowest}** ({(ot.get('slowest_section_ms') or 0)/1000:.1f}s)")
    else:
        st.info("Section timing data not available for this run.")

    st.divider()

    # ── Issues detected ───────────────────────────────────────────────
    st.subheader("Issues Detected")
    failures = perf.get("failures", [])
    if failures:
        SEVERITY_ICON  = {"critical": "🔴", "warning": "🟡", "info": "🔵"}
        SEVERITY_ORDER = {"critical": 0, "warning": 1, "info": 2}
        df_fail = pd.DataFrame([
            {
                "Severity": SEVERITY_ICON.get(f["severity"], "") + " " + f["severity"].upper(),
                "Section":  f["section_key"][:45],
                "Type":     f["failure_type"].replace("_", " ").title(),
                "Details":  f.get("plain_english", ""),
                "Metric":   f"{f['metric_value']:.3f}" if f.get("metric_value") is not None else "—",
                "_ord":     SEVERITY_ORDER.get(f["severity"], 9),
            }
            for f in failures
        ]).sort_values("_ord").drop(columns=["_ord"])
        st.dataframe(df_fail, use_container_width=True, hide_index=True,
                     column_config={"Details": st.column_config.TextColumn("Details", width="large")})
        n_crit = sum(1 for f in failures if f["severity"] == "critical")
        n_warn = sum(1 for f in failures if f["severity"] == "warning")
        n_info = sum(1 for f in failures if f["severity"] == "info")
        st.caption(f"🔴 {n_crit} critical &nbsp; 🟡 {n_warn} warnings &nbsp; 🔵 {n_info} informational")
    else:
        st.success("No issues detected. All sections generated and evaluated successfully.")


def _render_performance_legend() -> None:
    with st.expander("How to read this tab"):
        st.markdown(
            "**Total Time** — how long the entire document generation + evaluation took.\n\n"
            "**LLM Generation** — time the AI spent writing each section. "
            "This is usually the largest chunk.\n\n"
            "**Retrieval** — time spent searching your uploaded documents for relevant content.\n\n"
            "**Evaluation** — time spent running quality checks (DeepEval + Opik metrics).\n\n"
            "**🔴 Critical issues** — sections that could not be generated at all. Must be fixed.\n\n"
            "**🟡 Warnings** — sections with quality concerns (possible inaccuracies, low scores). "
            "Should be reviewed before submission.\n\n"
            "**🔵 Informational** — minor observations (e.g. slightly informal language). "
            "Optional to address.\n\n"
            "**🔥 High priority improvements** — actions that will most improve document quality.\n\n"
            "**📌 Medium priority** — worthwhile improvements for the next run.\n\n"
            "**💡 Low priority** — optimisations for speed or efficiency."
        )


# ═══════════════════════════════════════════════════════════════════════════
# RETRIEVAL EVALUATION HELPERS
# ═══════════════════════════════════════════════════════════════════════════

import re as _re

def _suggest_retrieval_query(section_key: str) -> str:
    """Auto-generate a focused retrieval query from a section title."""
    cleaned = _re.sub(r'^\d+[\d.]*\s*', '', section_key).lower()
    stop = {
        'the','a','an','and','or','of','in','on','at','for','to','is','are',
        'was','were','any','other','with','by','as','be','has','have','that',
        'this','which','carried','out','site','name','exact',
    }
    words = [w for w in _re.findall(r'\b[a-z]+\b', cleaned) if w not in stop and len(w) > 2]
    return ' '.join(words[:6]) if words else cleaned[:40]


def _diagnose_section(sec: Dict[str, Any]) -> Dict[str, Any]:
    """Return a structured diagnosis dict for one section's retrieval quality."""
    key       = sec.get("section_key", "?")
    is_static = sec.get("is_static", False)
    query     = (sec.get("retrieval_query") or "").strip()
    paths     = sec.get("retrieved_paths") or []
    rag       = (sec.get("extended_eval") or {}).get("rag_scores") or {}
    has_eval  = bool(rag)

    faith   = rag.get("faithfulness")
    ctx_p   = rag.get("contextual_precision") or rag.get("context_precision")
    ans_r   = rag.get("answer_relevancy")
    triad   = rag.get("rag_triad_score") or rag.get("ragas_score")
    ctx_rel = rag.get("context_relevance") or []
    claims  = rag.get("faithfulness_claims") or []

    if is_static:
        return {"key": key, "status": "static", "severity": "none"}

    if not query:
        return {
            "key": key, "status": "no_query", "severity": "critical",
            "query": query, "paths": paths, "rag": rag,
            "suggested_query": _suggest_retrieval_query(key),
        }

    if not paths:
        return {
            "key": key, "status": "no_results", "severity": "critical",
            "query": query, "paths": paths, "rag": rag,
            "suggested_query": _suggest_retrieval_query(key),
        }

    if not has_eval:
        return {
            "key": key, "status": "no_eval", "severity": "info",
            "query": query, "paths": paths, "rag": rag,
        }

    # Evaluate precision
    n_irrelevant = sum(1 for r in ctx_rel if not r.get("relevant", True))
    low_precision = ctx_p is not None and float(ctx_p) < 0.4

    # Evaluate faithfulness
    unsupported = [c for c in claims if not c.get("supported")]
    low_faith = faith is not None and float(faith) < 0.65

    if low_precision and n_irrelevant > 0:
        return {
            "key": key, "status": "irrelevant_chunks", "severity": "high",
            "query": query, "paths": paths, "rag": rag,
            "n_irrelevant": n_irrelevant, "n_total_chunks": len(paths),
            "ctx_relevance": ctx_rel,
            "suggested_query": _suggest_retrieval_query(key),
            "faith": faith, "ctx_p": ctx_p, "ans_r": ans_r, "triad": triad,
        }

    if low_faith and unsupported:
        return {
            "key": key, "status": "unsupported_claims", "severity": "high",
            "query": query, "paths": paths, "rag": rag,
            "unsupported": unsupported[:5],
            "faith": faith, "ctx_p": ctx_p, "ans_r": ans_r, "triad": triad,
        }

    # Partial — something is off but not clearly one root cause
    overall_ok = (
        (faith is None or float(faith) >= 0.65) and
        (ctx_p is None or float(ctx_p) >= 0.4) and
        (ans_r is None or float(ans_r) >= 0.5)
    )
    if overall_ok:
        return {
            "key": key, "status": "ok", "severity": "ok",
            "query": query, "paths": paths, "rag": rag,
            "faith": faith, "ctx_p": ctx_p, "ans_r": ans_r, "triad": triad,
        }

    return {
        "key": key, "status": "partial", "severity": "medium",
        "query": query, "paths": paths, "rag": rag,
        "faith": faith, "ctx_p": ctx_p, "ans_r": ans_r, "triad": triad,
        "suggested_query": _suggest_retrieval_query(key),
    }


def _render_retrieval_overview(sections: List[Dict[str, Any]], site_name: str, timestamp: str) -> None:
    """Overview sub-tab: heatmap table, aggregate cards, section drill-down."""
    st.caption(
        f"Run: **{timestamp}** | Site: **{site_name or '—'}**  "
        "Measures how relevant the retrieved document chunks are for each section's query."
    )

    with st.expander("What do these metrics mean?"):
        st.markdown(
            "**Faithfulness** — Fraction of generated claims grounded in the retrieved context. "
            "Low = hallucination.\n\n"
            "**Contextual Precision** — Were the most relevant chunks ranked first? "
            "Uses rank-weighted Average Precision.\n\n"
            "**Answer Relevancy** — Does the output address the original section instruction?\n\n"
            "**RAG Triad** — Harmonic mean of all three. Any single zero collapses it to 0."
        )

    rows = []
    for sec in sections:
        key = sec.get("section_key", "?")
        rag = (sec.get("extended_eval") or {}).get("rag_scores") or {}
        rows.append({
            "Section": key[:40],
            "Static": "Yes" if sec.get("is_static") else "No",
            "Retrieved Docs": len(sec.get("retrieved_paths") or []),
            "Faithfulness": _safe_pct(rag.get("faithfulness")),
            "Ctx Precision": _safe_pct(rag.get("contextual_precision") or rag.get("context_precision")),
            "Ans Relevancy": _safe_pct(rag.get("answer_relevancy")),
            "RAG Triad":     _safe_pct(rag.get("rag_triad_score") or rag.get("ragas_score")),
        })

    df = pd.DataFrame(rows)
    rag_cols = ["Faithfulness", "Ctx Precision", "Ans Relevancy", "RAG Triad"]
    has_data = df[rag_cols].notna().any().any()

    if has_data:
        styled = (
            df.set_index("Section")
            .style.background_gradient(cmap="RdYlGn", vmin=0, vmax=100, subset=rag_cols, axis=None)
            .format("{:.1f}", subset=rag_cols, na_rep="—")
        )
        st.dataframe(styled, use_container_width=True, height=min(len(df) * 40 + 60, 500))

        st.subheader("Aggregate")
        agg = st.columns(4)
        for i, col in enumerate(rag_cols):
            vals = df[col].dropna()
            agg[i].metric(f"Mean {col}", f"{vals.mean():.1f}" if not vals.empty else "—")

        if HAS_PLOTLY:
            triad_df = df[["Section", "RAG Triad"]].dropna(subset=["RAG Triad"]).sort_values("RAG Triad")
            if not triad_df.empty:
                fig = px.bar(
                    triad_df, x="RAG Triad", y="Section", orientation="h",
                    color="RAG Triad",
                    color_continuous_scale=["#D85A30", "#BA7517", "#1D9E75"],
                    range_color=[0, 100],
                    labels={"RAG Triad": "RAG Triad (0–100)", "Section": ""},
                    height=max(280, len(triad_df) * 30),
                )
                fig.update_coloraxes(showscale=False)
                fig.update_layout(margin=dict(t=20, b=10, l=10, r=20), yaxis_autorange="reversed")
                st.plotly_chart(fig, use_container_width=True)
    else:
        st.dataframe(df, use_container_width=True)
        st.info("RAG scores not yet computed — run a document generation first.")

    # ── Section drill-down ────────────────────────────────────────────────
    st.subheader("Section Detail")
    sel = st.selectbox("Select section", [s.get("section_key", "?") for s in sections], key="ret_overview_sel")
    sec_data = next((s for s in sections if s.get("section_key") == sel), {})
    rag = (sec_data.get("extended_eval") or {}).get("rag_scores") or {}

    dc1, dc2 = st.columns(2)
    with dc1:
        st.markdown("**Retrieval Query**")
        q = sec_data.get("retrieval_query", "")
        st.text(q if q else "— (no retrieval query)")
        st.markdown("**Retrieved Documents**")
        paths = sec_data.get("retrieved_paths") or []
        for p in paths:
            st.markdown(f"- `{os.path.basename(str(p))}`")
        if not paths:
            st.write("No retrieved paths recorded.")

    with dc2:
        st.markdown("**Scores**")
        faith = rag.get("faithfulness")
        ctx_p = rag.get("contextual_precision") or rag.get("context_precision")
        ans_r = rag.get("answer_relevancy")
        triad = rag.get("rag_triad_score") or rag.get("ragas_score")
        m1, m2 = st.columns(2)
        m1.metric("Faithfulness",   f"{float(faith):.3f}" if faith  is not None else "—")
        m2.metric("Ctx Precision",  f"{float(ctx_p):.3f}" if ctx_p  is not None else "—")
        m3, m4 = st.columns(2)
        m3.metric("Ans Relevancy",  f"{float(ans_r):.3f}" if ans_r  is not None else "—")
        m4.metric("RAG Triad",      f"{float(triad):.3f}" if triad  is not None else "—")

        claims = rag.get("faithfulness_claims") or []
        if claims:
            with st.expander(f"Faithfulness Claims ({len(claims)})"):
                for c in claims[:10]:
                    icon = "✓" if c.get("supported") else "✗"
                    clr  = CLR_SUCCESS if c.get("supported") else CLR_DANGER
                    st.markdown(
                        f'<span style="color:{clr};">{icon}</span> {c.get("claim","?")}',
                        unsafe_allow_html=True,
                    )
        gen_qs = rag.get("generated_questions") or []
        if gen_qs:
            with st.expander(f"Answer Relevancy Questions ({len(gen_qs)})"):
                for gq in gen_qs:
                    st.markdown(f"- {gq}")


def _render_retrieval_plan(sections: List[Dict[str, Any]], template_name: str) -> None:
    """Recommended Plan sub-tab: per-section diagnosis with exact fixes."""

    diagnoses = [_diagnose_section(s) for s in sections]
    active = [d for d in diagnoses if d["status"] not in ("static", "ok")]

    sev_order = {"critical": 0, "high": 1, "medium": 2, "info": 3, "none": 4, "ok": 5}
    active.sort(key=lambda d: sev_order.get(d["severity"], 9))

    # ── Summary counts ────────────────────────────────────────────────────
    n_crit   = sum(1 for d in active if d["severity"] == "critical")
    n_high   = sum(1 for d in active if d["severity"] == "high")
    n_medium = sum(1 for d in active if d["severity"] == "medium")
    n_ok     = sum(1 for d in diagnoses if d["status"] == "ok")
    n_static = sum(1 for d in diagnoses if d["status"] == "static")

    top_clr = CLR_DANGER if n_crit or n_high else CLR_WARNING if n_medium else CLR_SUCCESS
    st.markdown(
        f'<div style="background:{top_clr}22;border-left:4px solid {top_clr};'
        f'padding:12px 18px;border-radius:8px;margin-bottom:20px;">'
        f'<strong style="font-size:1rem;color:{top_clr};">Retrieval Health Check</strong><br>'
        f'<span style="color:#374151;font-size:0.9rem;">'
        f'{"🔴 " + str(n_crit) + " no data &nbsp;" if n_crit else ""}'
        f'{"🟠 " + str(n_high) + " quality issues &nbsp;" if n_high else ""}'
        f'{"🟡 " + str(n_medium) + " partial &nbsp;" if n_medium else ""}'
        f'{"✅ " + str(n_ok) + " passing &nbsp;" if n_ok else ""}'
        f'{"⬜ " + str(n_static) + " static (skipped)" if n_static else ""}'
        f'</span></div>',
        unsafe_allow_html=True,
    )

    if not active:
        st.success("All evaluated sections have acceptable retrieval quality.")
        return

    SEV_BADGE = {
        "critical": f'<span style="background:{CLR_DANGER};color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">NO DATA</span>',
        "high":     f'<span style="background:#EA580C;color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">HIGH</span>',
        "medium":   f'<span style="background:{CLR_WARNING};color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">MEDIUM</span>',
        "info":     f'<span style="background:#3B82F6;color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">INFO</span>',
    }

    for d in active:
        sev      = d["severity"]
        badge    = SEV_BADGE.get(sev, "")
        sec_name = d["key"]
        expanded = sev in ("critical", "high")

        with st.expander(f'{sec_name[:55]}', expanded=expanded):
            st.markdown(badge, unsafe_allow_html=True)
            st.markdown("")

            # ── STATUS: no retrieval query ────────────────────────────
            if d["status"] == "no_query":
                st.markdown("**Root cause:** This section has no `@!` retrieval query in the template. "
                            "The system searched all uploaded documents at once with no specific focus — "
                            "effectively no targeted retrieval occurred.")
                st.markdown("**Impact:** The LLM received a broad, unfiltered context dump. "
                            "Relevant facts are diluted by unrelated content, increasing hallucination risk.")
                st.markdown("**Exact fix — 3 steps:**")
                suggested = d["suggested_query"]
                st.markdown(
                    f"1. Open **`{template_name}`** in Word.\n"
                    f"2. Locate the section block for **`{sec_name}`**.\n"
                    f"3. After the section prompt text, add the `@!` separator followed by a search phrase. "
                    f"Based on this section's title, a starting query would be:\n\n"
                    f"   > `@! {suggested}`\n\n"
                    f"   Refine it by adding site-specific terms — process names, equipment, "
                    f"regulatory keywords — that appear in your uploaded documents."
                )

            # ── STATUS: query returned no results ────────────────────
            elif d["status"] == "no_results":
                q = d["query"]
                st.markdown(f"**Root cause:** Query `{q}` found **0 matching documents** in the "
                            "vector database. The FAISS index searched all uploaded documents but "
                            "none scored above the similarity threshold for this query.")
                st.markdown(
                    "**Why this happens** (in order of likelihood):\n"
                    "1. The uploaded documents don't contain content relevant to this section topic.\n"
                    "2. The query uses terminology that differs from the document language "
                    "(e.g. query says 'address' but document says 'location' or 'premises').\n"
                    "3. The documents are very short or image-heavy PDFs where text extraction fails."
                )
                suggested = d["suggested_query"]
                st.markdown("**Exact fix — choose one:**")
                st.markdown(
                    f"**Option A — Broaden the query** (try synonym-rich terms):\n\n"
                    f"   Change `@! {q}` → `@! {suggested} facility premises location`\n\n"
                    f"**Option B — Add a targeted document:**\n"
                    f"   Upload a document (SOP, regulatory filing, site master file) that "
                    f"explicitly covers *{sec_name}*. Re-run after adding it to the ZIP.\n\n"
                    f"**Option C — Check PDF text extraction:**\n"
                    f"   Open the uploaded PDFs in a text editor or `pdfplumber`. If the text is "
                    f"empty or garbled, the PDF is scanned (image-based). Use OCR: "
                    f"`pytesseract` or Adobe Acrobat → Export as searchable PDF."
                )

            # ── STATUS: irrelevant chunks retrieved ───────────────────
            elif d["status"] == "irrelevant_chunks":
                q         = d["query"]
                n_irr     = d.get("n_irrelevant", 0)
                n_total   = d.get("n_total_chunks", 0)
                ctx_p_val = d.get("ctx_p")
                faith_val = d.get("faith")
                ctx_rel   = d.get("ctx_relevance") or []

                st.markdown(
                    f"**Root cause:** {n_irr} of {n_total} retrieved chunks were marked as "
                    f"**not relevant** by the LLM evaluator. Contextual Precision = "
                    f"`{float(ctx_p_val):.2f}` — relevant documents are present but ranked "
                    f"below irrelevant ones (or absent entirely)."
                )

                if ctx_rel:
                    rel_lines = ""
                    for i, r in enumerate(ctx_rel[:5]):
                        icon = "✓" if r.get("relevant") else "✗"
                        reason = r.get("reason", "")
                        rel_lines += f"\n- Chunk {i+1}: {icon} {reason}"
                    st.markdown(f"**Chunk relevance breakdown:**{rel_lines}")

                st.markdown("**Exact fix — 3 steps:**")
                suggested = d.get("suggested_query", _suggest_retrieval_query(sec_name))
                st.markdown(
                    f"1. **Tighten the query** — current: `{q}`\n"
                    f"   Replace with more discriminative terms that only match the relevant document:\n"
                    f"   > `@! {suggested} {sec_name.split()[0] if sec_name.split() else ''}`.strip()\n\n"
                    f"2. **Split broad documents** — if one PDF covers many unrelated topics, "
                    f"split it into section-specific files. FAISS will then retrieve only the "
                    f"relevant file instead of returning chunks from multiple topics.\n\n"
                    f"3. **Increase `top_k` selectivity** — in `Vector_db.py → search()`, "
                    f"reduce `top_k=5` to `top_k=3` for sections where precision matters more "
                    f"than recall."
                )

            # ── STATUS: claims unsupported ────────────────────────────
            elif d["status"] == "unsupported_claims":
                faith_val   = d.get("faith", 0)
                unsupported = d.get("unsupported") or []
                st.markdown(
                    f"**Root cause:** Faithfulness = `{float(faith_val):.2f}` — the LLM generated "
                    f"**{len(unsupported)} claim(s)** that could not be verified in the retrieved documents."
                )
                if unsupported:
                    st.markdown("**Specific unsupported claims (these are the hallucinations):**")
                    for c in unsupported:
                        st.markdown(
                            f'<span style="color:{CLR_DANGER};">✗</span> {c.get("claim","?")}',
                            unsafe_allow_html=True,
                        )
                st.markdown("**Exact fix:**")
                st.markdown(
                    "1. For each ✗ claim above, identify what **factual information** it references "
                    "(a date, name, address, certification, quantity).\n"
                    "2. Find or create a document that explicitly states that fact.\n"
                    "3. Add the document to the ZIP and re-run. The claim will be supported once "
                    "FAISS retrieves the document containing it.\n\n"
                    "If no such document exists yet, add this instruction to the section prompt:\n"
                    "> *'If a fact is not present in the provided context, write: "
                    "[To be confirmed by site]. Do not infer or assume.'*"
                )

            # ── STATUS: partial / mixed ───────────────────────────────
            elif d["status"] == "partial":
                faith_val = d.get("faith")
                ctx_val   = d.get("ctx_p")
                ans_val   = d.get("ans_r")
                issues = []
                if faith_val is not None and float(faith_val) < 0.65:
                    issues.append(f"Faithfulness = `{float(faith_val):.2f}` (target ≥ 0.65) — some claims ungrounded")
                if ctx_val is not None and float(ctx_val) < 0.4:
                    issues.append(f"Ctx Precision = `{float(ctx_val):.2f}` (target ≥ 0.40) — irrelevant chunks in top results")
                if ans_val is not None and float(ans_val) < 0.5:
                    issues.append(f"Ans Relevancy = `{float(ans_val):.2f}` (target ≥ 0.50) — output doesn't fully address the instruction")
                for issue in issues:
                    st.markdown(f"- {issue}")
                suggested = d.get("suggested_query", _suggest_retrieval_query(sec_name))
                q = d.get("query", "")
                st.markdown(
                    "**Fix:** Tighten the retrieval query and verify document coverage:\n"
                    f"- Current query: `{q or '(none)'}`\n"
                    f"- Suggested: `@! {suggested}`\n"
                    "- Check that at least one uploaded document directly addresses this section topic."
                )

            # ── STATUS: eval not run ──────────────────────────────────
            elif d["status"] == "no_eval":
                st.info(
                    f"Documents were retrieved ({len(d.get('paths') or [])} files) but "
                    "DeepEval scores were not computed for this section. "
                    "This happens when the section was skipped (empty generated text) or "
                    "the extended evaluation failed. Re-run after a successful generation."
                )


def _render_tab_rag_evaluation(runs: List[Dict[str, Any]]) -> None:
    """Retrieval Evaluation — two sub-tabs: Overview and Recommended Plan."""

    if not runs:
        st.info("No evaluation runs found yet. Generate a PMF document first.")
        return

    run_file = runs[0].get("run_file", "")
    if not run_file or not os.path.exists(run_file):
        st.warning("Run file not found.")
        return

    payload      = _load_run_payload(run_file)
    run_arts     = payload.get("run_artifacts", payload)
    sections     = run_arts.get("sections", [])
    template_name = os.path.basename(run_arts.get("template_file", "template.docx"))

    if not sections:
        st.info("No section data available. Generate a PMF document first.")
        return

    site_name  = run_arts.get("site_name", "")
    timestamp  = run_arts.get("timestamp", "")

    ov_tab, plan_tab = st.tabs(["Overview", "Recommended Plan"])

    with ov_tab:
        _render_retrieval_overview(sections, site_name, timestamp)

    with plan_tab:
        _render_retrieval_plan(sections, template_name)


# ═══════════════════════════════════════════════════════════════════════════
# LLM IMPROVEMENT PLAN TAB
# ═══════════════════════════════════════════════════════════════════════════

def _metric_card(label: str, value: str, target: str, status: str) -> str:
    """Return HTML for a metric status card (current → target)."""
    clr = {"ok": CLR_SUCCESS, "warn": CLR_WARNING, "bad": CLR_DANGER}.get(status, "#6b7280")
    return (
        f'<div style="border:1px solid {clr};border-radius:8px;padding:10px 14px;'
        f'background:{clr}11;min-width:120px;">'
        f'<div style="font-size:0.72rem;color:#6b7280;font-weight:600;text-transform:uppercase;">{label}</div>'
        f'<div style="font-size:1.25rem;font-weight:700;color:{clr};">{value}</div>'
        f'<div style="font-size:0.72rem;color:#9ca3af;">target {target}</div>'
        f'</div>'
    )


def _render_tab_improvements(runs: List[Dict[str, Any]]) -> None:
    """LLM-specific improvement plan — technical, structured, section-aware."""

    if not runs:
        st.info("No evaluation runs found yet. Generate a PMF document first.")
        return

    run_file = runs[0].get("run_file", "")
    if not run_file or not os.path.exists(run_file):
        st.warning("Run file not found.")
        return

    payload       = _load_run_payload(run_file)
    run_arts      = payload.get("run_artifacts", payload)
    eval_data     = payload.get("evaluation", {}).get("document_scores", {})
    ext           = _extract_extended(payload)
    sections      = run_arts.get("sections", [])
    eval_sections = eval_data.get("sections", [])
    template_name = os.path.basename(run_arts.get("template_file", "template.docx"))
    site_name     = run_arts.get("site_name", "")

    rule_score    = float(eval_data.get("overall_score") or 0)
    composite     = float(ext.get("mean_composite") or rule_score or 0)
    judge_score   = ext.get("mean_judge_normalized")
    mean_faith    = ext.get("mean_faithfulness")
    mean_prec     = ext.get("mean_contextual_precision") or ext.get("mean_context_precision")
    mean_relev    = ext.get("mean_answer_relevancy") or ext.get("mean_answer_relevance_score")
    mean_hall     = ext.get("mean_hallucination_score")
    mean_reg_tone = ext.get("mean_regulatory_tone_score")

    # ── Build eval_sections lookup ────────────────────────────────────────
    eval_sec_lkp: Dict[str, Dict] = {s.get("section_key", ""): s for s in eval_sections}

    # ── Per-section diagnosis ─────────────────────────────────────────────
    SEV_ORDER = {"critical": 0, "high": 1, "medium": 2, "info": 3}

    sec_diagnoses: List[Dict[str, Any]] = []
    for sec in sections:
        if sec.get("is_static"):
            continue
        sk       = sec.get("section_key", "?")
        eval_sec = eval_sec_lkp.get(sk, {})

        rag   = (sec.get("extended_eval") or {}).get("rag_scores")   or {}
        judge = (sec.get("extended_eval") or {}).get("judge_scores") or {}

        faith      = rag.get("faithfulness")
        ans_rel    = rag.get("answer_relevancy")
        claims     = rag.get("faithfulness_claims") or []
        unsup      = [c for c in claims if not c.get("supported", True)]

        j_score    = judge.get("normalized_score")
        j_issues   = judge.get("critical_issues") or []
        j_feedback = judge.get("feedback", "")
        reg_tone   = judge.get("regulatory_tone_score")
        hall       = judge.get("hallucination_score")

        missing_kws = eval_sec.get("missing_keywords") or []
        char_len    = eval_sec.get("char_len", 0)
        min_chars   = eval_sec.get("required_min_chars", 0)
        too_short   = (min_chars > 0 and char_len < min_chars)

        issues: List[Dict] = []
        severity = "info"

        def _raise(new_sev):
            nonlocal severity
            if SEV_ORDER.get(new_sev, 9) < SEV_ORDER.get(severity, 9):
                severity = new_sev

        if j_score is not None:
            if float(j_score) < 60:
                issues.append({"type": "judge", "score": j_score, "j_issues": j_issues, "feedback": j_feedback})
                _raise("critical")
            elif float(j_score) < 75:
                issues.append({"type": "judge", "score": j_score, "j_issues": j_issues, "feedback": j_feedback})
                _raise("high")

        if faith is not None and float(faith) < 0.7:
            issues.append({"type": "faithfulness", "score": faith, "unsupported": unsup})
            _raise("high")

        if hall is not None and float(hall) > 0.35:
            issues.append({"type": "hallucination", "score": hall})
            _raise("high")

        if ans_rel is not None and float(ans_rel) < 0.6:
            issues.append({"type": "answer_relevancy", "score": ans_rel})
            _raise("medium")

        if reg_tone is not None and float(reg_tone) < 0.65:
            issues.append({"type": "reg_tone", "score": reg_tone})
            _raise("medium")

        if missing_kws:
            issues.append({"type": "missing_keywords", "keywords": missing_kws})
            _raise("medium")

        if too_short:
            issues.append({"type": "too_short", "char_len": char_len, "min_chars": min_chars})
            _raise("medium")

        if issues:
            sec_diagnoses.append({"key": sk, "severity": severity, "issues": issues})

    sec_diagnoses.sort(key=lambda d: SEV_ORDER.get(d["severity"], 9))

    # ── Global issues (not section-specific) ──────────────────────────────
    global_issues: List[Dict] = []
    missing_secs = eval_data.get("missing_required_sections", [])
    if missing_secs:
        global_issues.append({"type": "missing_sections", "sections": missing_secs})
    if mean_reg_tone is not None and float(mean_reg_tone) < 0.65:
        global_issues.append({"type": "global_reg_tone", "score": mean_reg_tone})
    if mean_hall is not None and float(mean_hall) > 0.35:
        global_issues.append({"type": "global_hallucination", "score": mean_hall})

    # ── Render ────────────────────────────────────────────────────────────
    if not sec_diagnoses and not global_issues:
        st.success(
            f"Composite score: **{composite:.1f}/100** — all metrics are within acceptable "
            "thresholds. No specific improvements required for this run."
        )
        return

    n_crit   = sum(1 for d in sec_diagnoses if d["severity"] == "critical")
    n_high_s = sum(1 for d in sec_diagnoses if d["severity"] == "high")
    n_med_s  = sum(1 for d in sec_diagnoses if d["severity"] == "medium")

    top_clr = CLR_DANGER if n_crit or n_high_s else CLR_WARNING if n_med_s else CLR_SUCCESS
    st.markdown(
        f'<div style="background:{top_clr}22;border-left:4px solid {top_clr};'
        f'padding:12px 18px;border-radius:8px;margin-bottom:20px;">'
        f'<strong style="font-size:1rem;color:{top_clr};">LLM Generation Health — {composite:.1f}/100</strong><br>'
        f'<span style="color:#374151;font-size:0.9rem;">'
        f'{"🔴 " + str(n_crit) + " critical &nbsp;" if n_crit else ""}'
        f'{"🟠 " + str(n_high_s) + " high &nbsp;" if n_high_s else ""}'
        f'{"🟡 " + str(n_med_s) + " medium &nbsp;" if n_med_s else ""}'
        f'{"✅ no section issues" if not sec_diagnoses else ""}'
        f'</span></div>',
        unsafe_allow_html=True,
    )

    SEV_BADGE = {
        "critical": f'<span style="background:{CLR_DANGER};color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">CRITICAL</span>',
        "high":     f'<span style="background:#EA580C;color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">HIGH</span>',
        "medium":   f'<span style="background:{CLR_WARNING};color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">MEDIUM</span>',
        "info":     f'<span style="background:#3B82F6;color:white;padding:2px 10px;border-radius:4px;font-size:0.78rem;font-weight:700;">INFO</span>',
    }

    # ── Global issues ──────────────────────────────────────────────────────
    if global_issues:
        st.markdown("#### Template-Level Issues")
        for gi in global_issues:

            if gi["type"] == "missing_sections":
                with st.expander("🔴 **Missing Required Sections — CRITICAL**", expanded=True):
                    st.markdown(SEV_BADGE["critical"], unsafe_allow_html=True)
                    st.markdown("")
                    st.markdown(
                        f"**Root cause:** {len(gi['sections'])} section(s) required by the PMF "
                        "evaluation rules are completely absent from the generated document."
                    )
                    st.markdown("**Exact fix:**")
                    st.markdown(
                        f"1. Open **`{template_name}`** in Word.\n"
                        "2. Add a new section block for each missing section listed below. "
                        "Each block must have a heading and a prompt instruction.\n"
                        "3. After the instruction add an `@!` retrieval query so the system "
                        "fetches relevant documents. Example format:\n\n"
                        "> `Describe the [topic]. @! [search phrase for this topic]`\n\n"
                        "**Missing sections:**\n"
                        + "\n".join(f"- `{s}`" for s in gi["sections"])
                    )

            elif gi["type"] == "global_reg_tone":
                with st.expander(
                    f'🟡 **Global Regulatory Tone = {float(gi["score"]):.2f} — MEDIUM**',
                    expanded=False,
                ):
                    st.markdown(SEV_BADGE["medium"], unsafe_allow_html=True)
                    st.markdown("")
                    st.markdown(
                        f"**Root cause:** Mean Regulatory Tone = `{float(gi['score']):.2f}` "
                        "(target ≥ 0.65). Language across sections is too informal or imprecise "
                        "for a regulatory submission."
                    )
                    st.markdown("**Exact fix — add this block to the template preamble (before Section 1):**")
                    st.code(
                        "You are a regulatory affairs specialist writing a Plant Master File "
                        "for submission to a competent authority under EU GMP Annex 4 / ICH Q10.\n"
                        "All sections must:\n"
                        "- Use passive voice and formal register "
                        "(e.g. 'is maintained', 'are performed', 'has been validated')\n"
                        "- Use precise technical terminology — avoid colloquial language\n"
                        "- State all facts with specificity: dates, version numbers, frequencies\n"
                        "- Avoid hedging phrases (may, might, could) unless genuine regulatory "
                        "uncertainty exists"
                    )

            elif gi["type"] == "global_hallucination":
                with st.expander(
                    f'🟠 **Global Hallucination = {float(gi["score"]):.2f} — HIGH**',
                    expanded=True,
                ):
                    st.markdown(SEV_BADGE["high"], unsafe_allow_html=True)
                    st.markdown("")
                    st.markdown(
                        f"**Root cause:** Mean Hallucination Score = `{float(gi['score']):.2f}` "
                        f"(Groundedness = `{1.0 - float(gi['score']):.2f}`, target ≤ 0.20). "
                        "The model is generating factual statements — addresses, certification "
                        "numbers, names, dates — not present in the retrieved documents."
                    )
                    st.markdown("**Exact fix — add this constraint to every section prompt:**")
                    st.code(
                        "Use only facts explicitly stated in the provided context. "
                        "Do not infer, extrapolate, or use general knowledge. "
                        "If a required fact is absent from the context, write exactly: "
                        "[To be completed by site]."
                    )
                    st.markdown(
                        "**Additionally:**\n"
                        "1. Upload site-specific documents (GMP certificates, manufacturing "
                        "authorisations, SOPs) — generic documents force hallucination.\n"
                        "2. In `.env`, set `AZURE_OPENAI_TEMPERATURE` ≤ 0.3. Higher values "
                        "increase fabrication risk."
                    )

    # ── Per-section issues ─────────────────────────────────────────────────
    if sec_diagnoses:
        if global_issues:
            st.markdown("---")
        st.markdown("#### Per-Section Generation Issues")

        for d in sec_diagnoses:
            sk       = d["key"]
            sev      = d["severity"]
            badge    = SEV_BADGE.get(sev, "")
            expanded = sev in ("critical", "high")

            with st.expander(f'{sk[:65]}', expanded=expanded):
                st.markdown(badge, unsafe_allow_html=True)
                st.markdown("")

                for issue in d["issues"]:
                    itype = issue["type"]

                    # ── Judge score ────────────────────────────────────
                    if itype == "judge":
                        j_sc  = float(issue["score"])
                        j_iss = issue.get("j_issues") or []
                        j_fb  = issue.get("feedback", "")
                        st.markdown(f"**Judge Score = `{j_sc:.0f}/100`** (target ≥ 75)")
                        if j_iss:
                            st.markdown("**Issues identified by the LLM judge:**")
                            for i in j_iss[:4]:
                                st.markdown(f"- {i}")
                        if j_fb:
                            st.markdown(
                                f"**Judge feedback:** *{j_fb[:320]}"
                                f"{'…' if len(j_fb) > 320 else ''}*"
                            )
                        st.markdown("**Exact prompt fix — add to this section's instruction:**")
                        st.code(
                            "Write this section in formal EU GMP Annex 4 regulatory submission "
                            "language. Be specific: reference actual site processes, equipment "
                            "identifiers, and personnel roles by name. Avoid vague statements. "
                            "State compliance status explicitly (e.g. 'compliant with', "
                            "'validated per', 'approved by')."
                        )
                        if j_sc < 60:
                            st.markdown(
                                "Score is critically low (`< 60`). Also verify:\n"
                                "1. Are source documents for this section in the ZIP? "
                                "Check **Retrieval Evaluation → Overview** for retrieved paths.\n"
                                "2. If no documents were retrieved, this section is generated "
                                "purely from general knowledge. Upload a relevant SOP or "
                                "procedure document and add an `@!` query to the template."
                            )
                        st.markdown("---")

                    # ── Faithfulness / unsupported claims ──────────────
                    elif itype == "faithfulness":
                        faith_val = float(issue["score"])
                        unsup     = issue.get("unsupported") or []
                        st.markdown(
                            f"**Faithfulness = `{faith_val:.2f}`** (target ≥ 0.70) — "
                            f"{len(unsup)} unsupported claim(s) found"
                        )
                        if unsup:
                            st.markdown("**Hallucinated claims (not grounded in retrieved documents):**")
                            for c in unsup[:5]:
                                st.markdown(
                                    f'<span style="color:{CLR_DANGER};">✗</span>&nbsp;'
                                    f'{c.get("claim", "?")}',
                                    unsafe_allow_html=True,
                                )
                        st.markdown("**Exact fix:**")
                        st.markdown(
                            "**Option A — Constrain via prompt** (add to this section's instruction):"
                        )
                        st.code(
                            "Only use facts explicitly stated in the provided context. "
                            "If a specific fact is not present, write: [To be confirmed by site]. "
                            "Do not infer or assume."
                        )
                        st.markdown(
                            "**Option B — Upload a supporting document:** Each ✗ claim above "
                            "references a specific fact (name, date, certification, address). "
                            "Find or create a document that explicitly contains that fact, "
                            "add it to the ZIP, and re-run."
                        )
                        st.markdown("---")

                    # ── Hallucination (Opik) ───────────────────────────
                    elif itype == "hallucination":
                        hall_val = float(issue["score"])
                        st.markdown(
                            f"**Hallucination Score = `{hall_val:.2f}`** "
                            f"(Groundedness = `{1.0 - hall_val:.2f}`, target hallucination ≤ 0.20)"
                        )
                        st.markdown("**Exact fix — add to this section's prompt:**")
                        st.code(
                            "Use only information present in the provided context. "
                            "Do not infer or generate facts from general knowledge. "
                            "Missing facts must be written as: [To be completed by site]."
                        )
                        st.markdown("---")

                    # ── Answer Relevancy ───────────────────────────────
                    elif itype == "answer_relevancy":
                        rel_val = float(issue["score"])
                        st.markdown(
                            f"**Answer Relevancy = `{rel_val:.2f}`** (target ≥ 0.60) — "
                            "generated text does not fully address the section instruction"
                        )
                        st.markdown(
                            "**Root cause:** The section prompt is too vague or too broad. "
                            "The LLM produces related content but not a direct, complete "
                            "response to the actual question asked."
                        )
                        st.markdown("**Exact fix — rewrite the section instruction to be explicit:**")
                        st.markdown(
                            "> ❌ *'Describe the manufacturing site.'*\n\n"
                            "> ✅ *'Describe the [Site Name] manufacturing facility. "
                            "Your response must include: (1) physical address and GMP "
                            "classification, (2) regulatory approval status and manufacturing "
                            "authorisation number, (3) types of medicinal products manufactured, "
                            "(4) responsible persons and their qualifications.'*\n\n"
                            "Adding a numbered list of required sub-points forces the model to "
                            "address each one, which directly raises Answer Relevancy."
                        )
                        st.markdown("---")

                    # ── Regulatory tone ────────────────────────────────
                    elif itype == "reg_tone":
                        tone_val = float(issue["score"])
                        st.markdown(
                            f"**Regulatory Tone = `{tone_val:.2f}`** (target ≥ 0.65) — "
                            "language in this section is informal or imprecise"
                        )
                        st.markdown("**Exact fix — add to this section's instruction:**")
                        st.code(
                            "Write in formal regulatory submission language consistent with "
                            "EU GMP Annex 4. Use passive constructions "
                            "(e.g. 'is performed', 'are maintained', 'has been validated'). "
                            "State all facts with specificity: include dates, version numbers, "
                            "and frequencies."
                        )
                        st.markdown("---")

                    # ── Missing keywords ───────────────────────────────
                    elif itype == "missing_keywords":
                        kws = issue.get("keywords") or []
                        st.markdown(
                            f"**Missing required keywords:** `{'`, `'.join(kws[:6])}`"
                        )
                        st.markdown(
                            "**Root cause:** The prompt doesn't explicitly ask the model to "
                            "cover these topics, or the source documents don't contain them."
                        )
                        st.markdown("**Exact fix — add to this section's instruction:**")
                        kw_str = ", ".join(kws[:6])
                        st.code(
                            f"Your response must explicitly address the following topics: "
                            f"{kw_str}. "
                            f"If any topic is not covered in the provided context, include: "
                            f"[Information on {kws[0] if kws else 'this topic'} to be provided "
                            f"by site]."
                        )
                        st.markdown("---")

                    # ── Section too short ──────────────────────────────
                    elif itype == "too_short":
                        char_len  = issue.get("char_len", 0)
                        min_chars = issue.get("min_chars", 0)
                        st.markdown(
                            f"**Section length = `{char_len}` chars** "
                            f"(required ≥ `{min_chars}`) — content is too brief"
                        )
                        st.markdown("**Exact fix — add to this section's instruction:**")
                        st.code(
                            "Provide a comprehensive and detailed response covering all "
                            "relevant aspects. Do not summarise — expand each point with "
                            "specific details, examples, and context from the provided documents."
                        )
                        st.markdown(
                            "If the output remains short after this prompt change, the source "
                            "documents may not contain enough detail. Upload additional "
                            "relevant documents for this section topic."
                        )
                        st.markdown("---")


# ═══════════════════════════════════════════════════════════════════════════
# MAIN RENDER
# ═══════════════════════════════════════════════════════════════════════════

def render_eval_dashboard() -> None:
    """Render the two-tab evaluation dashboard (Retrieval Evaluation | LLM Evaluation)."""
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)
    st.markdown(
        '<div class="dashboard-header">'
        'Healthark GenAI Evaluation Framework v1.0'
        '</div>',
        unsafe_allow_html=True,
    )
    st.caption("PMF Document Generator — LLM Evaluation & Benchmarking")

    runs = _load_runs()

    outer1, outer2 = st.tabs(["Retrieval Evaluation", "LLM Evaluation"])

    with outer1:
        _render_tab_rag_evaluation(runs)

    with outer2:
        t1, t2, t3, t4 = st.tabs(["Run Overview", "Section Heatmap", "Performance", "Recommended Plan"])
        with t1:
            _render_tab_overview(runs)
        with t2:
            _render_tab_heatmap(runs)
        with t3:
            _render_tab_performance(runs)
        with t4:
            _render_tab_improvements(runs)


# ═══════════════════════════════════════════════════════════════════════════
# STANDALONE
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__" or "streamlit" in getattr(
    sys.modules.get("__main__"), "__module__", ""
):
    st.set_page_config(
        page_title="PMF Evaluation Dashboard",
        page_icon="chart_with_upwards_trend",
        layout="wide",
    )
    render_eval_dashboard()
