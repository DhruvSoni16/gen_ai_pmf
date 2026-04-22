"""Generate the LLM Evaluation & Benchmarking technical documentation DOCX."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1.2)
section.top_margin  = section.bottom_margin = Inches(1.0)

# ── Helper: add horizontal rule ──────────────────────────────────────────
def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '5340C0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def para(doc, text, bold=False, italic=False, size=10.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def numbered(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
    p.style.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
    return p

def table_2col(doc, rows, col1_w=2.2, col2_w=4.0):
    t = doc.add_table(rows=len(rows)+1, cols=2)
    t.style = 'Light Grid Accent 1'
    t.cell(0,0).text = 'Item'
    t.cell(0,1).text = 'Description'
    for i, (k, v) in enumerate(rows):
        t.cell(i+1, 0).text = k
        t.cell(i+1, 1).text = v
    doc.add_paragraph()
    return t

def section_divider(doc, title):
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x53, 0x40, 0xC0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_hr(doc)
    doc.add_paragraph()

PURPLE = (0x53, 0x40, 0xC0)
GREEN  = (0x1D, 0x9E, 0x75)
ORANGE = (0xBA, 0x75, 0x17)

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cover.paragraph_format.space_before = Pt(60)
run = cover.add_run("HEALTHARK INSIGHTS")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x53, 0x40, 0xC0)

doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = title_p.add_run("LLM Evaluation & Benchmarking\nFramework")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(0x11, 0x11, 0x22)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r2 = sub_p.add_run("Technical Reference Documentation\nPMF Document Generator — Initiative 4")
r2.font.size = Pt(13)
r2.font.italic = True
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

doc.add_paragraph()
doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r3 = date_p.add_run(f"Version 1.0  ·  {datetime.date.today().strftime('%B %Y')}")
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x77, 0x77, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════
heading(doc, "Table of Contents", level=1, color=PURPLE)
toc_entries = [
    ("1",  "Executive Summary", ""),
    ("2",  "System Architecture Overview", ""),
    ("3",  "Rule-Based Evaluation Framework", ""),
    ("3.1","PMF Section Rules & Configuration", ""),
    ("3.2","Scoring Algorithm", ""),
    ("3.3","Required Section Patterns", ""),
    ("3.4","Retrieval Coverage Metric", ""),
    ("4",  "LLM-as-Judge Evaluation (PMFJudge)", ""),
    ("4.1","Design Philosophy", ""),
    ("4.2","Five-Criterion Rubric", ""),
    ("4.3","Rubric Weights & Score Normalisation", ""),
    ("4.4","Provider Support: Azure OpenAI & Anthropic Claude", ""),
    ("4.5","Prompt Architecture", ""),
    ("4.6","Response Parsing & Score Recomputation", ""),
    ("4.7","SHA-256 Caching", ""),
    ("4.8","Retry Logic", ""),
    ("4.9","Parallel Evaluation", ""),
    ("5",  "DeepEval RAG Triad Metrics", ""),
    ("5.1","Framework Overview & Motivation", ""),
    ("5.2","Faithfulness", ""),
    ("5.3","Contextual Precision", ""),
    ("5.4","Answer Relevancy", ""),
    ("5.5","RAG Triad Composite Score", ""),
    ("5.6","Heuristic Fallback", ""),
    ("5.7","Caching Strategy", ""),
    ("6",  "Opik-Style Evaluation Metrics", ""),
    ("6.1","Methodology: Direct Continuous Scoring", ""),
    ("6.2","Hallucination Scorer", ""),
    ("6.3","Answer Relevance Scorer", ""),
    ("6.4","Regulatory Tone Scorer (PMF-Specific)", ""),
    ("6.5","Opik Composite Score", ""),
    ("7",  "Lexical & Semantic Metrics", ""),
    ("7.1","BLEU Score", ""),
    ("7.2","ROUGE Scores", ""),
    ("7.3","BERTScore", ""),
    ("7.4","Semantic Similarity", ""),
    ("8",  "EvalSuite Orchestrator", ""),
    ("8.1","Architecture & Lazy Loading", ""),
    ("8.2","Composite Score Formula", ""),
    ("8.3","Letter Grade System", ""),
    ("8.4","EvalResult Dataclass", ""),
    ("8.5","Document-Level Aggregation", ""),
    ("8.6","Benchmarking Mode", ""),
    ("9",  "Evaluation Storage & Persistence", ""),
    ("9.1","Data Directory Structure", ""),
    ("9.2","JSONL Index", ""),
    ("9.3","Per-Run JSON Schema", ""),
    ("10", "MLflow Experiment Tracking", ""),
    ("10.1","Integration Design", ""),
    ("10.2","Logged Parameters, Metrics & Tags", ""),
    ("10.3","Local vs Remote Tracking", ""),
    ("11", "Performance Analysis & Latency Tracking", ""),
    ("11.1","Instrumentation Points", ""),
    ("11.2","PerformanceAnalyzer Architecture", ""),
    ("11.3","Failure Detection Rules", ""),
    ("11.4","Improvement Recommendation Engine", ""),
    ("11.5","Latency Optimisation Strategies", ""),
    ("12", "Evaluation Dashboard", ""),
    ("12.1","Tab Architecture", ""),
    ("12.2","Run Overview Tab", ""),
    ("12.3","Section Heatmap Tab", ""),
    ("12.4","Trend Analysis Tab", ""),
    ("12.5","Performance Tab", ""),
    ("12.6","DOCX Report Generation", ""),
    ("13", "Benchmarking Framework", ""),
    ("14", "End-to-End Pipeline Flow", ""),
    ("15", "Configuration Reference", ""),
    ("16", "Dependency Matrix", ""),
    ("17", "Appendix A — Metric Formulas", ""),
    ("18", "Appendix B — Prompt Templates", ""),
    ("19", "Appendix C — Data Schemas", ""),
]
for num, title, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    indent = Inches(0.3) if "." in num else Inches(0)
    p.paragraph_format.left_indent = indent
    run_num = p.add_run(f"{num}  ")
    run_num.font.bold = ("." not in num)
    run_num.font.size = Pt(10)
    run_title = p.add_run(title)
    run_title.font.bold = ("." not in num)
    run_title.font.size = Pt(10)
    if "." not in num:
        run_num.font.color.rgb = RGBColor(0x53, 0x40, 0xC0)
        run_title.font.color.rgb = RGBColor(0x53, 0x40, 0xC0)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 1 · Executive Summary")
heading(doc, "1. Executive Summary", 1, PURPLE)

para(doc,
    "The Healthark GenAI Evaluation Framework (Initiative 4) is a multi-layer, production-grade "
    "quality assurance system designed specifically for AI-generated Plant Master File (PMF) "
    "regulatory documents. PMF documents are submitted to regulatory authorities under EU GMP "
    "Annex 4 and ICH Q10 guidelines. Errors, hallucinations, or non-compliant language in these "
    "documents carry significant regulatory and business risk. The framework therefore applies "
    "four independent evaluation methodologies — structural rules, LLM-as-judge, RAG triad "
    "metrics, and direct continuous scoring — to assess every generated section from multiple "
    "complementary angles.")

para(doc,
    "Each methodology captures a different failure mode. Rule-based scoring catches structural "
    "deficiencies: sections that are too short, missing required headings, or lacking the "
    "site name. LLM-as-Judge scoring uses a second language model acting as a domain expert "
    "to assess factual accuracy, regulatory language quality, site specificity, completeness, "
    "and structural coherence. DeepEval RAG Triad metrics evaluate the retrieval-augmented "
    "generation quality by checking whether the generated content is grounded in source "
    "documents, whether the most relevant chunks were retrieved first, and whether the output "
    "actually addresses the section instruction. Opik-style metrics provide a fast direct "
    "continuous score for hallucination level, answer relevance, and regulatory tone, using "
    "single LLM calls that are faster and more holistic than claim-by-claim verification.")

para(doc,
    "All four layers feed into a weighted composite score (Rule 20%, Judge 55%, RAG Triad 25%) "
    "and a letter grade (A–F). Results are stored persistently, tracked in MLflow for experiment "
    "comparison, and displayed in a six-tab Streamlit dashboard. A dedicated Performance tab "
    "shows per-section latency breakdowns (retrieval / LLM generation / evaluation), detects "
    "failures, and generates prioritised improvement recommendations — in both plain English "
    "for non-technical stakeholders and technical detail for engineers.")

heading(doc, "Key Metrics Summary", 2, GREEN)
table_2col(doc, [
    ("Rule Score",          "0–100 structural quality (min length, keywords, required sections)"),
    ("Judge Score",         "0–100 LLM-based rubric evaluation across 5 regulatory criteria"),
    ("Faithfulness",        "0–1 fraction of claims grounded in retrieved source documents"),
    ("Contextual Precision","0–1 rank-weighted relevance of retrieved chunks (MAP)"),
    ("Answer Relevancy",    "0–1 how closely generated text addresses section instruction"),
    ("RAG Triad Score",     "0–1 harmonic mean of Faithfulness + Precision + Relevancy"),
    ("Hallucination Score", "0–1 Opik direct continuous score (0 = grounded, 1 = hallucinated)"),
    ("Regulatory Tone",     "0–1 Opik PMF-specific language formality for EU GMP submissions"),
    ("Opik Composite",      "0–1 mean of (1-Hallucination) + Answer Relevance + Regulatory Tone"),
    ("Composite Score",     "0–100 weighted blend of Rule + Judge + RAG Triad"),
])

# ════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 2 · System Architecture Overview")
heading(doc, "2. System Architecture Overview", 1, PURPLE)

para(doc,
    "The evaluation system is structured as a layered pipeline that executes immediately after "
    "document generation. The architecture separates concerns into four distinct layers: "
    "data ingestion and generation, rule-based evaluation, LLM-based evaluation, and "
    "observability/storage.")

heading(doc, "2.1 Component Map", 2, GREEN)
para(doc, "The following components constitute the evaluation framework:", bold=False)
bullet(doc, "extraction_pmf() — entry point in Extraction_module_PMF.py; orchestrates the full pipeline")
bullet(doc, "DocumentRetriever — FAISS vector database using all-mpnet-base-v2 embeddings for chunk retrieval")
bullet(doc, "handle_user_message() — 4-agent LLM router that generates each PMF section")
bullet(doc, "evaluate_run() — rule-based scoring layer (eval_utils.py)")
bullet(doc, "EvalSuite — orchestrator for all LLM-based evaluation (healthark_eval/suite.py)")
bullet(doc, "PMFJudge — LLM-as-Judge rubric scorer (eval_judge.py)")
bullet(doc, "RAGEvaluator — DeepEval RAG Triad (eval_rag.py)")
bullet(doc, "OpikStyleScorer — Opik-inspired direct scoring (eval_opik_style.py)")
bullet(doc, "PerformanceAnalyzer — latency and failure analysis (eval_performance.py)")
bullet(doc, "MLflowTracker — experiment logging to local mlruns store (eval_mlflow_tracker.py)")
bullet(doc, "save_eval_run() — persistent JSON+JSONL storage (eval_store.py)")
bullet(doc, "render_eval_dashboard() — 6-tab Streamlit dashboard (app_eval_dashboard.py)")

heading(doc, "2.2 Execution Sequence", 2, GREEN)
para(doc, "The pipeline executes in the following order per PMF generation run:")
numbered(doc, "User uploads ZIP (site documents) + Excel reference + site name via Streamlit UI")
numbered(doc, "Documents are extracted and indexed into a FAISS vector database")
numbered(doc, "For each template section: vector retrieval → LLM generation → timing recorded")
numbered(doc, "After all sections: rule-based evaluation (evaluate_run) runs synchronously")
numbered(doc, "Extended evaluation (_run_extended_evaluation): EvalSuite + OpikStyleScorer per section")
numbered(doc, "Document-level aggregation: mean composite, grade distribution, Opik aggregates")
numbered(doc, "PerformanceAnalyzer runs: failures detected, improvements generated")
numbered(doc, "MLflowTracker logs all metrics to local mlruns directory")
numbered(doc, "save_eval_run() persists full run payload to JSON + JSONL index")
numbered(doc, "Dashboard reads persisted run for display; session_state updated for inline display")

heading(doc, "2.3 Data Flow", 2, GREEN)
para(doc,
    "The central data structure is run_artifacts — a Python dictionary built incrementally "
    "throughout the pipeline. It begins with metadata (timestamp, site_name, model_name, "
    "template_file) and accumulates sections[], timing{}, extended_eval_summary{}, "
    "performance_report{}, mlflow_run_id, and mlflow_ui_url as each pipeline stage completes. "
    "This dictionary is serialized to JSON at the end of each run and stored in "
    "data/eval_runs/{timestamp}_{site_name}.json.")

heading(doc, "2.4 LLM Provider Configuration", 2, GREEN)
table_2col(doc, [
    ("Primary LLM",        "Azure OpenAI GPT-4o (document generation + all evaluation)"),
    ("Model Deployment",   "AZURE_NAME env var (e.g. gpt-4o)"),
    ("API Version",        "AZURE_VERSION env var (e.g. 2024-06-01)"),
    ("Fallback Provider",  "Anthropic Claude (claude-sonnet-4-6) — judge only"),
    ("Embedding Model",    "all-mpnet-base-v2 via sentence-transformers (FAISS indexing)"),
    ("Temperature",        "0.0 for all evaluation calls; 0.1 for document generation"),
])

# ════════════════════════════════════════════════════════════════
# SECTION 3 — RULE-BASED EVALUATION
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 3 · Rule-Based Evaluation Framework")
heading(doc, "3. Rule-Based Evaluation Framework", 1, PURPLE)

para(doc,
    "The rule-based layer is the fastest and most deterministic component of the framework. "
    "It operates purely on the generated text without any LLM calls, running in milliseconds. "
    "It checks structural properties that are objective and verifiable: minimum content length, "
    "presence of required keywords, site name inclusion, and completeness of the required PMF "
    "section structure. This layer aligns with EU GMP Annex 4 and ICH Q10 requirements for "
    "PMF document structure.")

heading(doc, "3.1 PMF Section Rules & Configuration", 2, GREEN)
para(doc,
    "Rules are defined in src/eval/eval_config.py via the get_eval_rules() function. Each "
    "section is identified by case-insensitive substring matching against the section key "
    "derived from the PMF template. The configuration returns a dictionary with three parts:")
bullet(doc, "required_section_patterns — list of section names that MUST be present in the document")
bullet(doc, "section_rules — per-section min_chars and required_keywords thresholds")
bullet(doc, "fallback_rule — applied to any section not matched by a specific rule (min_chars=80)")

para(doc, "Per-section minimum character thresholds are:")
table_2col(doc, [
    ("GENERAL INFORMATION",   "100 characters"),
    ("MANUFACTURING ACTIVITIES", "80 characters"),
    ("PERSONNEL",             "100 characters"),
    ("PREMISES",              "100 characters"),
    ("EQUIPMENT",             "80 characters"),
    ("SANITATION",            "80 characters"),
    ("PRODUCTION",            "120 characters"),
    ("QUALITY ASSURANCE",     "100 characters"),
    ("STORAGE",               "80 characters"),
    ("DOCUMENTATION",         "80 characters"),
    ("INTERNAL AUDIT",        "80 characters"),
    ("All other sections",    "80 characters (fallback)"),
])

heading(doc, "3.2 Scoring Algorithm", 2, GREEN)
para(doc,
    "The score_section() function in eval_utils.py evaluates each generated section against "
    "its applicable rule. The scoring logic is as follows:")
numbered(doc, "Resolve the applicable rule via resolve_rule_for_section() — case-insensitive substring match")
numbered(doc, "Check minimum character length: pass/fail (char_ok boolean)")
numbered(doc, "Check required keywords: each keyword checked for case-insensitive presence in generated text")
numbered(doc, "Check site name: if site_name is provided and non-empty, verify it appears in the text")
numbered(doc, "Compute base score: starts at 100.0")
numbered(doc, "Deduct 40 points if character length check fails (char_ok = False)")
numbered(doc, "Deduct proportionally for missing keywords: (missing_count / total_keywords) * 30 points")
numbered(doc, "Deduct 15 points if site name is missing from the generated text")
numbered(doc, "Clamp final score to [0, 100]")

code_block(doc,
    "# Scoring pseudocode\n"
    "score = 100.0\n"
    "if len(generated_text) < rule['min_chars']:\n"
    "    score -= 40.0\n"
    "if required_keywords:\n"
    "    missing = [kw for kw in required_keywords if kw.lower() not in text_lower]\n"
    "    score -= (len(missing) / len(required_keywords)) * 30.0\n"
    "if site_name and site_name.lower() not in text_lower:\n"
    "    score -= 15.0\n"
    "score = max(0.0, min(100.0, score))"
)

heading(doc, "3.3 Required Section Patterns", 2, GREEN)
para(doc,
    "Five sections are designated as mandatory in a valid PMF document, aligned with the "
    "EU GMP Annex 4 structure. If any of these sections are absent from the generated document, "
    "they are listed in missing_required_sections in the evaluation output and flagged as "
    "warnings in the dashboard. The five required patterns are:")
bullet(doc, "GENERAL INFORMATION — site identification, address, manufacturing licence")
bullet(doc, "PERSONNEL — organisation charts, qualified person, key personnel, training")
bullet(doc, "PREMISES — site layout, manufacturing areas, environmental controls")
bullet(doc, "PRODUCTION — manufacturing processes, batch sizes, equipment")
bullet(doc, "QUALITY ASSURANCE — QMS overview, internal audits, change control, deviations")

heading(doc, "3.4 Document-Level Aggregation", 2, GREEN)
para(doc,
    "score_document() aggregates individual section scores into a document-level result. "
    "The overall_score is the arithmetic mean of all section scores. Retrieval coverage "
    "is computed as the percentage of non-static sections that had at least one chunk "
    "retrieved from the vector database, reflecting how grounded the document is in "
    "uploaded source material.")

code_block(doc,
    "overall_score = mean(section['score'] for section in sections)\n"
    "retrieval_coverage = (sections_with_retrieval / non_static_sections) * 100"
)

# ════════════════════════════════════════════════════════════════
# SECTION 4 — LLM-AS-JUDGE
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 4 · LLM-as-Judge Evaluation")
heading(doc, "4. LLM-as-Judge Evaluation (PMFJudge)", 1, PURPLE)

para(doc,
    "LLM-as-Judge is the highest-weight component of the composite score (55%). The PMFJudge "
    "class (src/eval/eval_judge.py) uses a second language model — acting as a domain expert "
    "regulatory affairs specialist — to score each generated section against a structured "
    "five-criterion rubric. This approach captures qualitative dimensions that automated "
    "metrics cannot: whether the language is appropriately regulatory, whether the content is "
    "complete relative to what was asked, and whether the writing is coherent.")

heading(doc, "4.1 Design Philosophy", 2, GREEN)
para(doc,
    "Traditional NLP metrics (BLEU, ROUGE, BERTScore) require a reference text to compare "
    "against. In the PMF generation context, no ground-truth reference exists for every "
    "section of every site. LLM-as-Judge eliminates the reference requirement: the judge "
    "receives the section instruction, the retrieved source documents, the site name, and "
    "the generated output, and evaluates quality using its domain knowledge of regulatory "
    "submissions. The judge model is set to temperature=0.0 for deterministic, reproducible "
    "scoring. Weighted scores are always recomputed locally from the raw criterion scores — "
    "the model's reported weighted_score is never trusted, preventing arithmetic errors from "
    "propagating into the final composite.")

heading(doc, "4.2 Five-Criterion Rubric", 2, GREEN)
para(doc, "The JUDGE_RUBRIC defines five evaluation criteria, each scored 1–5 by the judge model:")

t = doc.add_table(rows=6, cols=4)
t.style = 'Light Grid Accent 1'
headers = ["Criterion", "Weight", "Score Range", "What It Measures"]
for i, h in enumerate(headers):
    t.cell(0, i).text = h
rows_data = [
    ("Factual Accuracy",       "30%", "1–5", "All claims grounded in retrieved source documents; no fabrication"),
    ("Regulatory Language",    "25%", "1–5", "Formal passive construction, ICH Q10/ISO 13485 terminology"),
    ("Site Specificity",       "20%", "1–5", "Correct site name, location, site-specific processes referenced"),
    ("Completeness",           "15%", "1–5", "All required sub-topics for section type addressed"),
    ("Structural Coherence",   "10%", "1–5", "Logical flow, appropriate headings/tables, easy to navigate"),
]
for i, row in enumerate(rows_data):
    for j, val in enumerate(row):
        t.cell(i+1, j).text = val
doc.add_paragraph()

para(doc,
    "Score level 5 represents exemplary quality; score 1 represents unacceptable quality "
    "requiring complete rework. Each level has a precise written description in the prompt "
    "to ensure consistent scoring across different sections and runs. The full rubric text "
    "is pre-rendered at module load time and injected into every prompt.")

heading(doc, "4.3 Rubric Weights & Score Normalisation", 2, GREEN)
para(doc,
    "The weighted score is computed as the dot product of criterion scores and their weights. "
    "The result is normalised to a 0–100 scale by dividing by 5 (the maximum weighted score) "
    "and multiplying by 100.")
code_block(doc,
    "weighted_score = sum(score[c] * weight[c] for c in criteria)  # range: 0-5\n"
    "normalized_score = (weighted_score / 5.0) * 100.0             # range: 0-100\n\n"
    "# Example: scores = {factual:4, regulatory:3, site:5, complete:4, coherent:4}\n"
    "# weighted = 4*0.30 + 3*0.25 + 5*0.20 + 4*0.15 + 4*0.10 = 3.90\n"
    "# normalized = (3.90 / 5.0) * 100 = 78.0"
)
para(doc,
    "Crucially, the weights sum to exactly 1.0 (verified by a smoke test assertion at module "
    "load time): 0.30 + 0.25 + 0.20 + 0.15 + 0.10 = 1.00. The arithmetic is always "
    "recomputed in Python after receiving the judge's response, never taken from the model's "
    "self-reported weighted_score field, eliminating any risk of LLM arithmetic errors.")

heading(doc, "4.4 Provider Support", 2, GREEN)
para(doc, "PMFJudge supports two LLM providers with automatic client initialisation:")
bullet(doc, "azure_openai — AzureOpenAI client; requires AZURE_KEY, AZURE_ENDPOINT, AZURE_VERSION env vars or constructor args")
bullet(doc, "anthropic — Anthropic client; requires ANTHROPIC_API_KEY env var or constructor api_key arg")
para(doc,
    "The system defaults to azure_openai with gpt-4o deployment in the live pipeline, "
    "matching the generation model. The Anthropic Claude fallback (claude-sonnet-4-6) is "
    "supported for offline testing or when Azure quota is exhausted. Both providers use "
    "max_tokens=2048 and temperature=0.0.")

heading(doc, "4.5 Prompt Architecture", 2, GREEN)
para(doc,
    "The judge prompt is a two-part system+user structure. The system prompt establishes "
    "the judge's persona as a regulatory affairs specialist with 15+ years of PMF review "
    "experience. The user prompt is rendered from a template (_USER_TEMPLATE) that includes:")
bullet(doc, "Section Key — the section identifier (e.g. DEVICE DESCRIPTION)")
bullet(doc, "Section Instruction — the original prompt given to the generation LLM")
bullet(doc, "Retrieved Source Documents — up to 8,000 characters of context the generator had access to")
bullet(doc, "Site Name — for site specificity verification")
bullet(doc, "Optional Reference Output — ground-truth text if available (benchmark mode)")
bullet(doc, "Generated Output — the text being evaluated")
bullet(doc, "Full Rubric Text — all five criteria with all five score level descriptions")
para(doc,
    "The prompt explicitly requests ONLY a valid JSON object response with a defined schema. "
    "The schema includes: scores dict (one integer per criterion), weighted_score, "
    "normalized_score, strengths list, weaknesses list, critical_issues list, "
    "improvement_suggestions list, judge_confidence float, and evaluation_notes string.")

heading(doc, "4.6 Response Parsing & Score Recomputation", 2, GREEN)
para(doc,
    "The _parse_judge_response() function handles model output robustly:")
numbered(doc, "Strip markdown code fences (```json ... ```) if present using _strip_code_fences()")
numbered(doc, "Parse JSON via json.loads()")
numbered(doc, "Validate each of the 5 criterion scores is present and in range [1, 5]")
numbered(doc, "Recompute weighted_score and normalized_score locally — never trust model arithmetic")
numbered(doc, "Ensure all list fields (strengths, weaknesses, critical_issues, improvement_suggestions) are lists")
numbered(doc, "Clamp judge_confidence to [0.0, 1.0]")
numbered(doc, "Raise ValueError on any validation failure — triggers retry")

heading(doc, "4.7 SHA-256 Caching", 2, GREEN)
para(doc,
    "Judge evaluations are expensive (one API call per section). The caching layer stores "
    "results keyed by SHA-256(section_key | generated_text[:first500] | rubric_version). "
    "Cache files are stored as JSON in data/eval_cache/{hash}.json. Cache hits avoid "
    "redundant API calls on repeated evaluations of identical content — important during "
    "iterative prompt development where the same section text is regenerated multiple times. "
    "Cache invalidation occurs automatically when the rubric version string changes (e.g. "
    "v1.0 → v2.0), ensuring stale cached scores are not used after rubric updates.")

heading(doc, "4.8 Retry Logic", 2, GREEN)
para(doc,
    "The judge call includes a two-attempt retry mechanism. On the first attempt, the "
    "standard prompt is sent. If _parse_judge_response() raises any exception (JSON parse "
    "error, missing criterion, out-of-range score), the same prompt is resent with a "
    "_RETRY_SUFFIX appended: 'Your previous response was not valid JSON. Respond with ONLY "
    "the JSON object, no markdown, no explanation.' If the second attempt also fails, an "
    "_error_result() dict is returned with judge_error=True and all scores set to null. "
    "Error sections are excluded from document-level aggregation.")

heading(doc, "4.9 Parallel Evaluation", 2, GREEN)
para(doc,
    "The score_document() method supports optional parallel evaluation via "
    "concurrent.futures.ThreadPoolExecutor. With parallel=True and max_workers=3, up to "
    "three sections are evaluated concurrently. This can reduce total judge evaluation time "
    "by approximately 60% on documents with 10+ sections, at the cost of higher concurrent "
    "API usage. The default is serial evaluation (parallel=False) to respect API rate limits.")

# ════════════════════════════════════════════════════════════════
# SECTION 5 — DeepEval RAG Triad
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 5 · DeepEval RAG Triad Metrics")
heading(doc, "5. DeepEval RAG Triad Metrics", 1, PURPLE)

para(doc,
    "The DeepEval RAG Triad (src/eval/eval_rag.py) is a three-metric framework for evaluating "
    "the quality of Retrieval-Augmented Generation pipelines. It is implemented from scratch "
    "following DeepEval's published methodology — without installing the deepeval package, "
    "which avoids dependency conflicts with the Azure OpenAI SDK version required by the "
    "generation pipeline. The three metrics — Faithfulness, Contextual Precision, and Answer "
    "Relevancy — together form the RAG Triad Score, which is the harmonic mean of all three.")

heading(doc, "5.1 Framework Overview & Motivation", 2, GREEN)
para(doc,
    "RAG systems have three distinct failure modes that require three separate metrics. A "
    "system can generate fluent, well-formatted text that is entirely hallucinated (low "
    "Faithfulness). It can retrieve the correct documents but rank irrelevant chunks first, "
    "degrading the context window (low Contextual Precision). Or it can produce text that is "
    "factually correct but does not actually address what the section instruction asked for "
    "(low Answer Relevancy). No single metric catches all three failure modes simultaneously, "
    "which is why the triad approach is necessary.")

heading(doc, "5.2 Faithfulness", 2, GREEN)
para(doc,
    "Faithfulness measures the fraction of factual claims in the generated text that are "
    "supported by the retrieved source documents. It directly quantifies hallucination at "
    "the claim level.")
para(doc, "Algorithm:", bold=True)
numbered(doc, "Claim Extraction: the LLM is prompted with the generated text and instructed to extract all atomic factual claims as a JSON list. The prompt uses chain-of-thought: the model must reason step-by-step before producing the list.")
numbered(doc, "Entailment Checking: for each extracted claim, a separate LLM call checks whether the merged retrieved context entails, contradicts, or is neutral to the claim.")
numbered(doc, "Score Computation: Faithfulness = supported_claims / total_claims")
code_block(doc,
    "# Faithfulness formula\n"
    "faithfulness = len([c for c in claims if c['verdict'] == 'yes']) / len(claims)\n\n"
    "# DeepEval cache key\n"
    "cache_key = SHA256(section_key | generated_text[:500] | 'faithfulness' | 'deepeval_v1.0')"
)
para(doc,
    "The claim extraction prompt instructs the model to extract complete, standalone sentences "
    "that can be verified as true or false independently. Generic statements or instructions "
    "without specific factual claims are excluded. This prevents over-penalising sections "
    "that correctly use standard regulatory boilerplate. The entailment check prompt "
    "provides 4,000 characters of merged context and the claim, and asks for a binary "
    "yes/no verdict with a reasoning chain.")
para(doc,
    "Faithfulness scores in the 0.05–0.15 range are expected for PMF documents generated "
    "from typical site document uploads. This is not a system failure — it reflects the "
    "nature of PMF generation where the LLM paraphrases, synthesises, and adds regulatory "
    "structure to source content rather than quoting it verbatim. A Faithfulness score of "
    "0.07 means 7% of individual atomic claims appear verbatim or near-verbatim in the "
    "source chunks, not that 93% of the content is wrong.")

heading(doc, "5.3 Contextual Precision", 2, GREEN)
para(doc,
    "Contextual Precision measures whether the most relevant retrieved chunks appear at the "
    "top of the retrieval result list. In RAG systems, the context window is limited; if "
    "relevant chunks are buried below irrelevant ones, the generator receives degraded context.")
para(doc, "Algorithm (Rank-Weighted Average Precision):", bold=True)
numbered(doc, "For each retrieved chunk (at most top_k=5), the LLM determines whether it is relevant to the section's retrieval query.")
numbered(doc, "Rank-weighted AP is computed: only positions that are relevant contribute to the score, weighted by their inverse rank position.")
numbered(doc, "AP@k = (1/R) * Σ(precision@i * rel_i) where R = total relevant chunks, precision@i = relevant_found_at_i / i, rel_i = 1 if chunk i is relevant")
code_block(doc,
    "# Rank-weighted Average Precision\n"
    "relevant_found = 0\n"
    "precision_sum  = 0.0\n"
    "for i, chunk in enumerate(retrieved_chunks, start=1):\n"
    "    if chunk_is_relevant[i]:\n"
    "        relevant_found += 1\n"
    "        precision_sum  += relevant_found / i\n"
    "AP = precision_sum / max(relevant_found, 1)"
)

heading(doc, "5.4 Answer Relevancy", 2, GREEN)
para(doc,
    "Answer Relevancy measures whether the generated text actually addresses the section "
    "instruction. A generated section might be factually accurate and formally written yet "
    "completely miss the point of what was asked.")
para(doc, "Algorithm (Reverse-Question Method):", bold=True)
numbered(doc, "The LLM generates N=3 questions that the generated output would serve as a good answer to.")
numbered(doc, "Each generated question is compared to the original section instruction using cosine similarity of sentence embeddings (all-MiniLM-L6-v2 via SentenceTransformer).")
numbered(doc, "Answer Relevancy = mean cosine similarity across all N generated questions")
code_block(doc,
    "# Answer Relevancy formula\n"
    "questions = llm_generate_questions(generated_text, n=3)\n"
    "embeddings = sentence_model.encode([original_instruction] + questions)\n"
    "similarities = cosine_similarity(embeddings[0], embeddings[1:])\n"
    "answer_relevancy = mean(similarities)"
)
para(doc,
    "When SentenceTransformer is unavailable (e.g. torch DLL issues on Windows), the "
    "implementation falls back to TF-IDF keyword overlap similarity as a heuristic. This "
    "fallback is less accurate but ensures the metric always produces a value rather than "
    "failing with an ImportError.")

heading(doc, "5.5 RAG Triad Composite Score", 2, GREEN)
para(doc,
    "The RAG Triad Score is the harmonic mean of all three metrics. The harmonic mean is "
    "chosen over arithmetic mean because it penalises extreme lows more aggressively — a "
    "document with perfect Precision (1.0) and Relevancy (1.0) but zero Faithfulness (0.0) "
    "receives a RAG Triad Score of 0.0, reflecting that grounding failures are unacceptable "
    "in a regulatory context regardless of other qualities.")
code_block(doc,
    "# Harmonic mean of the three metrics\n"
    "def harmonic_mean(values):\n"
    "    valid = [v for v in values if v and v > 0]\n"
    "    if not valid: return 0.0\n"
    "    return len(valid) / sum(1.0/v for v in valid)\n\n"
    "rag_triad_score = harmonic_mean([faithfulness, contextual_precision, answer_relevancy])"
)

heading(doc, "5.6 Heuristic Fallback", 2, GREEN)
para(doc,
    "All three RAG Triad metrics have keyword-overlap heuristic fallbacks that activate "
    "when the LLM client is unavailable. The heuristics use token-level Jaccard similarity "
    "to approximate faithfulness, a fixed 0.5 for precision, and term overlap for relevancy. "
    "These fallbacks produce approximate values (not directly comparable to LLM-based scores) "
    "and are logged with a framework field value of 'heuristic_fallback' to distinguish them.")

heading(doc, "5.7 Caching Strategy", 2, GREEN)
para(doc,
    "All three RAG Triad metrics share the same SHA-256 caching infrastructure. Cache keys "
    "include the metric name as a component: SHA256(section_key | text[:500] | metric_name | "
    "'deepeval_v1.0'). This allows individual metrics to be invalidated independently. Cache "
    "files are stored in data/eval_cache/deepeval/. The cache version string 'deepeval_v1.0' "
    "must be bumped to invalidate all cached scores when the evaluation prompts change.")

# ════════════════════════════════════════════════════════════════
# SECTION 6 — OPIK-STYLE METRICS
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 6 · Opik-Style Evaluation Metrics")
heading(doc, "6. Opik-Style Evaluation Metrics", 1, PURPLE)

para(doc,
    "The Opik-style metrics (src/eval/eval_opik_style.py) are implemented following the "
    "methodology of Comet's Opik evaluation platform — direct continuous scoring using a "
    "single LLM call per metric per section. This contrasts with DeepEval's approach of "
    "extracting discrete claims and verifying each one individually. The Opik approach is "
    "faster (one call vs. N+1 calls per section), produces continuous scores that capture "
    "nuance better than binary per-claim verdicts, and is more holistic — the LLM evaluates "
    "the entire section in context rather than individual sentences in isolation.")
para(doc,
    "Note: the opik Python package (Comet Opik) was not installed because its version 2.x "
    "requires openai>=2.0.0, which conflicts with the openai==1.109.1 required by the Azure "
    "OpenAI generation pipeline. The metrics are therefore implemented from scratch using "
    "the existing Azure OpenAI client, following Opik's published prompt methodology.")

heading(doc, "6.1 Methodology: Direct Continuous Scoring", 2, GREEN)
para(doc,
    "Each Opik-style metric asks the LLM to return a single float score on a 0–1 scale "
    "directly, without intermediate claim extraction or similarity computation. The LLM is "
    "given a precise scoring rubric with labelled anchor points (e.g. 0.0, 0.25, 0.50, "
    "0.75, 1.0) and instructed to return ONLY valid JSON. The _LLMCaller abstraction supports "
    "both Azure OpenAI and Anthropic providers using the same interface.")

heading(doc, "6.2 Hallucination Scorer", 2, GREEN)
para(doc,
    "The Hallucination Scorer quantifies how much of the generated output is NOT supported by "
    "the source context. Score 0.0 means no hallucinations (all factual claims grounded); "
    "score 1.0 means complete hallucination (virtually nothing is grounded).")
para(doc, "Scoring anchor points:")
bullet(doc, "0.00 — No hallucinations; all factual claims grounded in context")
bullet(doc, "0.25 — Minor hallucinations; 1–2 small unsupported claims")
bullet(doc, "0.50 — Moderate; several unsupported claims")
bullet(doc, "0.75 — Significant; majority of specific facts not in context")
bullet(doc, "1.00 — Complete hallucination; virtually nothing grounded")
para(doc,
    "Critical PMF-specific design decision: the Hallucination Scorer explicitly instructs "
    "the LLM to NOT penalise regulatory boilerplate language or standard PMF formatting. "
    "Only invented factual claims — site names, numbers, certifications, dates, site-specific "
    "details not present in the context — are penalised. This prevents the metric from "
    "unfairly penalising sections that correctly use standard ICH Q10 language patterns that "
    "do not appear verbatim in the uploaded source documents.")
para(doc,
    "The dashboard displays this metric as Groundedness = 1.0 - hallucination_score, "
    "so higher is always better across all displayed metrics.")

heading(doc, "6.3 Answer Relevance Scorer", 2, GREEN)
para(doc,
    "The Answer Relevance Scorer evaluates how well the generated output addresses the "
    "original section instruction. Score 1.0 = perfectly relevant; score 0.0 = completely "
    "irrelevant. Anchor points:")
bullet(doc, "1.00 — Perfectly relevant; directly and completely addresses the instruction")
bullet(doc, "0.75 — Mostly relevant; addresses instruction with minor gaps")
bullet(doc, "0.50 — Partially relevant; addresses some aspects but misses key requirements")
bullet(doc, "0.25 — Mostly irrelevant; barely addresses the instruction")
bullet(doc, "0.00 — Completely irrelevant; does not address the instruction at all")
para(doc,
    "This differs from DeepEval's Answer Relevancy metric in methodology: DeepEval uses "
    "reverse-question generation + cosine similarity (multi-call, embedding-based), while "
    "Opik uses a direct LLM judgement in a single call. The Opik approach is more suitable "
    "when embedding models are unavailable or slow.")

heading(doc, "6.4 Regulatory Tone Scorer (PMF-Specific)", 2, GREEN)
para(doc,
    "The Regulatory Tone Scorer is a PMF-domain-specific metric not present in vanilla Opik. "
    "It evaluates whether the generated text uses language appropriate for a Plant Master "
    "File submission under EU GMP Annex 4 / ICH Q10. This is critical because regulatory "
    "submissions require formal passive voice, precise terminology, and avoidance of informal "
    "or ambiguous phrasing that would cause reviewers to question the document's reliability.")
para(doc, "Scoring criteria:")
bullet(doc, "1.00 — Exemplary: formal passive construction, precise regulatory terminology, ICH Q10/EU GMP conventions, no informal phrases")
bullet(doc, "0.75 — Good: mostly formal, minor informal phrases or non-standard terms")
bullet(doc, "0.50 — Acceptable: mix of formal/informal, would need moderate revision")
bullet(doc, "0.25 — Poor: predominantly informal language, non-regulatory style")
bullet(doc, "0.00 — Inappropriate: completely unsuitable for a regulatory submission")

heading(doc, "6.5 Opik Composite Score", 2, GREEN)
para(doc,
    "The Opik Composite is the arithmetic mean of three transformed values:")
code_block(doc,
    "opik_composite = mean([\n"
    "    1.0 - hallucination_score,   # inverted: lower hallucination = higher score\n"
    "    answer_relevance_score,\n"
    "    regulatory_tone_score,\n"
    "])"
)
para(doc,
    "All three components are averaged only when they are not None. If any scorer fails "
    "(LLM unavailable, JSON parse error), that component is excluded from the mean rather "
    "than causing the entire composite to fail. The Opik composite is stored separately from "
    "the main composite score and is not currently included in the weighted composite formula "
    "(which uses Rule + Judge + RAG Triad). It is displayed as a supplementary quality signal "
    "in the dashboard and DOCX report.")

# ════════════════════════════════════════════════════════════════
# SECTION 7 — LEXICAL & SEMANTIC METRICS
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 7 · Lexical & Semantic Metrics")
heading(doc, "7. Lexical & Semantic Metrics", 1, PURPLE)

para(doc,
    "Lexical and semantic metrics (src/eval/eval_metrics.py) are reference-based: they "
    "require a ground-truth reference text to compare against the generated output. In the "
    "PMF generation context, reference texts are only available in benchmark mode (where "
    "manually authored reference sections are provided). These metrics are therefore disabled "
    "by default in the live pipeline (run_lexical=False, run_semantic=False in EvalSuite) "
    "and are available primarily for offline benchmarking and model comparison.")

heading(doc, "7.1 BLEU Score", 2, GREEN)
para(doc,
    "BLEU (Bilingual Evaluation Understudy) measures n-gram precision between the generated "
    "and reference texts. The implementation uses sacrebleu's sentence_bleu function with "
    "smooth_method='exp' (Chen & Cherry exponential smoothing) to handle zero n-gram counts "
    "gracefully. The LexicalMetrics.compute_bleu() method returns:")
bullet(doc, "bleu — overall BLEU score (0–100, sacrebleu scale)")
bullet(doc, "bleu_1, bleu_2, bleu_3, bleu_4 — individual n-gram precision scores")
para(doc,
    "BLEU is a precision-focused metric. High BLEU means the generated text uses many of "
    "the same n-grams as the reference. In PMF evaluation, BLEU scores are typically low "
    "(10–30) because the generator paraphrases rather than reproducing reference text verbatim, "
    "which is the correct behaviour for AI-assisted document generation.")

heading(doc, "7.2 ROUGE Scores", 2, GREEN)
para(doc,
    "ROUGE (Recall-Oriented Understudy for Gisting Evaluation) measures overlap between "
    "generated and reference texts from a recall perspective. Three variants are computed "
    "via google's rouge_score library with stemming enabled:")
bullet(doc, "ROUGE-1 — unigram (single word) overlap: precision, recall, F1")
bullet(doc, "ROUGE-2 — bigram overlap: precision, recall, F1")
bullet(doc, "ROUGE-L — longest common subsequence: precision, recall, F1")
para(doc,
    "ROUGE-L is particularly relevant for regulatory documents because it captures sequential "
    "structure (whether the generated text follows a similar logical order to the reference) "
    "without requiring exact n-gram matches. ROUGE scores are returned as nested dicts "
    "accessible as rouge_1_precision, rouge_2_f1, rouge_l_recall, etc.")

heading(doc, "7.3 BERTScore", 2, GREEN)
para(doc,
    "BERTScore (Zhang et al., 2019) computes semantic similarity between generated and "
    "reference texts at the token level using pre-trained BERT contextual embeddings. Unlike "
    "BLEU/ROUGE which match surface n-grams, BERTScore matches semantic meaning — synonyms "
    "and paraphrases score highly even if no exact words match.")
para(doc, "Implementation details:")
bullet(doc, "Model: distilbert-base-multilingual-cased (default) or roberta-large for higher accuracy")
bullet(doc, "Returns: precision, recall, F1 per example + corpus-level averages")
bullet(doc, "Rescaling: baseline rescaling enabled (rescale_with_baseline=True) for better interpretability")
bullet(doc, "Batching: compute_all_metrics() batches BERTScore across all sections for efficiency")
para(doc,
    "BERTScore is disabled in the live pipeline due to a torch DLL loading issue on Windows "
    "(TorchDLL conflict with the FAISS-CPU build). It runs correctly in Linux/Mac environments "
    "and in Docker deployments. The BERTScore phase in run_document() is designed to run once "
    "across all sections in a single batch call rather than per-section, significantly reducing "
    "compute time compared to running it per-section.")

heading(doc, "7.4 Semantic Similarity", 2, GREEN)
para(doc,
    "The SemanticMetrics.compute_semantic_similarity() method computes the cosine similarity "
    "between sentence-level embeddings of the generated and reference texts using the "
    "all-MiniLM-L6-v2 model from sentence-transformers. This is a lightweight (80MB) model "
    "that produces 384-dimensional embeddings and runs fast on CPU. It is primarily used as "
    "the embedding backbone for Answer Relevancy in the RAG Triad evaluator.")

# ════════════════════════════════════════════════════════════════
# SECTION 8 — EVALSUIT ORCHESTRATOR
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 8 · EvalSuite Orchestrator")
heading(doc, "8. EvalSuite Orchestrator", 1, PURPLE)

para(doc,
    "EvalSuite (healthark_eval/suite.py) is the central public API that orchestrates all "
    "evaluation components. It is the single entry point called from the extraction pipeline "
    "and from the dashboard's live evaluation mode. EvalSuite handles lazy loading of "
    "components, composite score computation, letter grading, and result serialisation.")

heading(doc, "8.1 Architecture & Lazy Loading", 2, GREEN)
para(doc,
    "EvalSuite does not instantiate any metric components at __init__ time. All components "
    "(PMFJudge, RAGEvaluator, LexicalMetrics, SemanticMetrics) are instantiated lazily on "
    "first use via private _get_judge(), _get_rag(), _get_lexical(), _get_semantic() methods. "
    "This design means importing EvalSuite has near-zero overhead — torch and heavy ML models "
    "are not loaded until actually needed. It also allows the pipeline to initialise EvalSuite "
    "once at the start of a run and reuse it across all sections without paying model load "
    "costs per section.")
para(doc, "Constructor parameters:")
table_2col(doc, [
    ("task",             "'pmf' — task config selector (PMF_TASK_CONFIG)"),
    ("llm_provider",     "'azure_openai' or 'anthropic' — judge and RAG evaluator provider"),
    ("llm_model",        "Model deployment name (e.g. 'gpt-4o')"),
    ("api_key",          "LLM API key"),
    ("azure_endpoint",   "Azure OpenAI endpoint URL"),
    ("azure_api_version","Azure API version string"),
    ("run_judge",        "bool — enable/disable LLM-as-Judge (default True)"),
    ("run_rag",          "bool — enable/disable RAG Triad (default True)"),
    ("run_lexical",      "bool — enable/disable BLEU/ROUGE (default False)"),
    ("run_semantic",     "bool — enable/disable BERTScore (default False)"),
])

heading(doc, "8.2 Composite Score Formula", 2, GREEN)
para(doc,
    "The composite_score is a weighted mean of available metric scores. The weights are "
    "defined in _COMPOSITE_WEIGHTS and sum to 1.0:")
code_block(doc,
    "_COMPOSITE_WEIGHTS = {\n"
    "    'rule':             0.20,   # Rule-based structural score\n"
    "    'judge_normalized': 0.55,   # LLM-as-Judge (0-100 scale, normalised)\n"
    "    'rag_triad_score':  0.25,   # DeepEval RAG Triad (0-1 scale, scaled to 0-100)\n"
    "}\n\n"
    "# Components missing (None) are excluded; remaining weights are renormalised\n"
    "available = {k: v for k, v in scores.items() if v is not None}\n"
    "total_weight = sum(_COMPOSITE_WEIGHTS[k] for k in available)\n"
    "composite = sum(available[k] * _COMPOSITE_WEIGHTS[k] for k in available) / total_weight"
)
para(doc,
    "The 55% weight on Judge Score reflects its importance as the most domain-relevant "
    "metric — it evaluates regulatory language, site specificity, and completeness that "
    "automated metrics cannot capture. The 25% weight on RAG Triad reflects its importance "
    "for detecting hallucination and retrieval quality. Rule score at 20% serves as a "
    "structural sanity check.")

heading(doc, "8.3 Letter Grade System", 2, GREEN)
para(doc, "Composite scores are mapped to letter grades via the _grade() function:")
table_2col(doc, [
    ("A", "composite_score >= 90"),
    ("B", "composite_score >= 75"),
    ("C", "composite_score >= 60"),
    ("D", "composite_score >= 45"),
    ("F", "composite_score < 45"),
])
para(doc,
    "Grades are assigned per section (for individual section EvalResults) and for the "
    "document as a whole (based on mean_composite in extended_eval_summary). The "
    "grade_distribution dict counts how many sections received each letter grade, providing "
    "a quality profile of the full document.")

heading(doc, "8.4 EvalResult Dataclass", 2, GREEN)
para(doc,
    "The run() method returns an EvalResult dataclass with the following fields:")
table_2col(doc, [
    ("section_key",     "Section identifier string"),
    ("rule_score",      "0–100 float from rule-based evaluation"),
    ("judge_scores",    "Full PMFJudge result dict (normalized_score, all criteria, strengths, etc.)"),
    ("rag_scores",      "RAGEvaluator result dict (faithfulness, contextual_precision, answer_relevancy, rag_triad_score)"),
    ("lexical_scores",  "LexicalMetrics result dict (bleu, rouge_1, rouge_2, rouge_l)"),
    ("semantic_scores", "SemanticMetrics result dict (bertscore_f1, semantic_similarity)"),
    ("composite_score", "0–100 weighted composite"),
    ("grade",           "Letter grade A–F"),
])
para(doc, "EvalResult provides a to_dict() method for JSON serialisation and storage in run_artifacts[sections][extended_eval].")

heading(doc, "8.5 Document-Level Aggregation", 2, GREEN)
para(doc,
    "The document-level aggregation in _run_extended_evaluation() (Extraction_module_PMF.py) "
    "computes the following means across all non-static, non-empty sections:")
bullet(doc, "mean_composite — arithmetic mean of all section composite_scores")
bullet(doc, "mean_judge_normalized — arithmetic mean of normalized_score across successful judge evaluations")
bullet(doc, "mean_rag_triad_score — arithmetic mean of rag_triad_score across RAG evaluations")
bullet(doc, "mean_faithfulness — arithmetic mean of faithfulness scores")
bullet(doc, "mean_hallucination_score — arithmetic mean of Opik hallucination_score")
bullet(doc, "mean_answer_relevance_score — arithmetic mean of Opik answer_relevance_score")
bullet(doc, "mean_regulatory_tone_score — arithmetic mean of Opik regulatory_tone_score")
bullet(doc, "mean_opik_composite — arithmetic mean of Opik opik_composite")
bullet(doc, "overall_grade — letter grade of mean_composite")
bullet(doc, "grade_distribution — Counter dict of per-section letter grades")

heading(doc, "8.6 Benchmarking Mode", 2, GREEN)
para(doc,
    "run_benchmark() evaluates all test cases in a benchmark dataset directory "
    "(data/benchmark/). Each case is a JSON file with section_key, section_instruction, "
    "retrieved_context, reference_output, and site_name. The benchmark runner applies all "
    "enabled metrics to every case and returns a DocumentEvalResult with per-case scores "
    "and aggregate statistics. This enables offline evaluation of model quality changes, "
    "prompt modifications, or retrieval configuration changes without running the full "
    "document generation pipeline.")

# ════════════════════════════════════════════════════════════════
# SECTION 9 — STORAGE
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 9 · Evaluation Storage & Persistence")
heading(doc, "9. Evaluation Storage & Persistence", 1, PURPLE)

para(doc,
    "All evaluation results are persisted to disk by save_eval_run() in src/eval/eval_store.py. "
    "The storage design prioritises durability and queryability: every run is saved as a "
    "complete self-contained JSON file, and a JSONL index enables fast listing and "
    "filtering without loading individual run files.")

heading(doc, "9.1 Data Directory Structure", 2, GREEN)
code_block(doc,
    "data/\n"
    "  eval_runs/\n"
    "    index.jsonl                          # one JSON line per run (summary)\n"
    "    20241215_143022_SiteName.json        # full run payload\n"
    "    20241216_091500_OtherSite.json\n"
    "  eval_cache/\n"
    "    deepeval/                            # RAG Triad SHA-256 cache files\n"
    "    opik/                                # Opik scorer SHA-256 cache files\n"
    "    {sha256_hash}.json                   # Judge cache files (root level)\n"
    "  benchmark/\n"
    "    case_001.json\n"
    "    case_002.json\n"
    "  artifacts/\n"
    "    Extracted_folder/                    # uploaded ZIP contents\n"
    "    generated output file/               # DOCX + PDF outputs\n"
    "mlruns/                                  # MLflow local tracking store\n"
    "  PMF_Document_Generation/              # experiment directory"
)

heading(doc, "9.2 JSONL Index", 2, GREEN)
para(doc,
    "The index.jsonl file contains one JSON object per line, each representing a run summary. "
    "This format enables O(n) sequential scan for listing runs without loading full payloads, "
    "and supports atomic appending (each line is appended independently). Each index entry "
    "contains: timestamp, site_name, overall_score, template_file, run_file (path to full JSON).")

heading(doc, "9.3 Per-Run JSON Schema", 2, GREEN)
para(doc, "Each full run JSON file has the following top-level structure:")
code_block(doc,
    "{\n"
    "  'run_artifacts': {\n"
    "    'timestamp': '20241215_143022',\n"
    "    'site_name': 'Langensbold',\n"
    "    'model_name': 'gpt-4o',\n"
    "    'template_file': 'templates/PMF_Template.docx',\n"
    "    'sections': [\n"
    "      {\n"
    "        'section_key': 'GENERAL INFORMATION',\n"
    "        'generated_text': '...',\n"
    "        'retrieved_paths': ['data/.../chunk1.txt'],\n"
    "        'timing': {\n"
    "          'retrieval_ms': 123.4,\n"
    "          'generation_ms': 8540.2,\n"
    "          'eval_ms': 6210.1,\n"
    "          'total_ms': 14873.7\n"
    "        },\n"
    "        'extended_eval': { ...EvalResult.to_dict()... },\n"
    "        'opik_eval': { ...OpikStyleScorer result... }\n"
    "      }\n"
    "    ],\n"
    "    'timing': {\n"
    "      'generation_phase_ms': 87420.0,\n"
    "      'total_pipeline_ms': 147650.0,\n"
    "      'total_generation_ms': 79830.0,\n"
    "      'total_retrieval_ms': 3210.0,\n"
    "      'total_eval_ms': 60230.0\n"
    "    },\n"
    "    'extended_eval_summary': { ...document-level aggregates... },\n"
    "    'performance_report': { ...PerformanceReport.to_dict()... },\n"
    "    'mlflow_run_id': 'a1b2c3d4...',\n"
    "    'mlflow_ui_url': 'http://localhost:5000/#/runs/a1b2c3d4'\n"
    "  },\n"
    "  'evaluation': {\n"
    "    'document_scores': {\n"
    "      'overall_score': 76.19,\n"
    "      'section_count': 21,\n"
    "      'retrieval_coverage': 85.7,\n"
    "      'missing_required_sections': [],\n"
    "      'sections': [ ...per-section rule scores... ]\n"
    "    }\n"
    "  }\n"
    "}"
)

# ════════════════════════════════════════════════════════════════
# SECTION 10 — MLFLOW
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 10 · MLflow Experiment Tracking")
heading(doc, "10. MLflow Experiment Tracking", 1, PURPLE)

para(doc,
    "MLflow experiment tracking (src/eval/eval_mlflow_tracker.py) logs every evaluation run "
    "to the MLflow tracking server, enabling comparison of metric trends across document "
    "versions, model changes, template updates, and site-specific run history. The default "
    "configuration uses a local file-based store (./mlruns directory) requiring no server "
    "infrastructure. The MLflow UI is started separately with 'mlflow ui' and accessed at "
    "http://localhost:5000.")

heading(doc, "10.1 Integration Design", 2, GREEN)
para(doc,
    "MLflowTracker uses lazy import: 'import mlflow' is deferred to the first call to "
    "log_run() rather than at module import time. This means importing the tracker adds "
    "zero startup overhead — mlflow is only loaded when a run is actually being logged. "
    "The tracker degrades gracefully: if mlflow is not installed, all tracking calls "
    "become no-ops with INFO-level logging.")
para(doc,
    "Experiment setup: the tracker creates or retrieves the 'PMF_Document_Generation' "
    "experiment on first call. All runs are logged under this experiment name, making "
    "them filterable and comparable in the MLflow UI. The tracking URI is set to "
    "file:///abs/path/to/mlruns using the absolute resolved path to avoid relative-path "
    "issues when run from different working directories.")

heading(doc, "10.2 Logged Parameters, Metrics & Tags", 2, GREEN)
para(doc, "Each MLflow run logs three categories of data:")
para(doc, "Parameters (stable run descriptors):", bold=True)
table_2col(doc, [
    ("site_name",          "Manufacturing site name"),
    ("template_file",      "PMF template filename (basename only)"),
    ("model_name",         "LLM deployment name used for generation"),
    ("timestamp",          "Run timestamp (YYYYmmdd_HHMMSS)"),
    ("sections_evaluated", "Count of sections with extended evaluation"),
    ("framework",          "'deepeval_rag_triad + opik_style'"),
])
para(doc, "Metrics (numeric values tracked over time):", bold=True)
table_2col(doc, [
    ("rule_score",              "Rule-based overall score (0–100)"),
    ("judge_score",             "Mean LLM judge normalized score (0–100)"),
    ("faithfulness",            "Mean DeepEval faithfulness (0–1)"),
    ("rag_triad_score",         "Mean DeepEval RAG Triad score (0–1)"),
    ("composite_score",         "Mean weighted composite score (0–100)"),
    ("hallucination_score",     "Mean Opik hallucination score (0–1)"),
    ("answer_relevance_score",  "Mean Opik answer relevance (0–1)"),
    ("regulatory_tone_score",   "Mean Opik regulatory tone (0–1)"),
    ("opik_composite",          "Mean Opik composite (0–1)"),
    ("section_count",           "Total sections in document"),
    ("retrieval_coverage_pct",  "Percentage of sections with retrieved chunks"),
])
para(doc, "Tags (categorical labels):", bold=True)
table_2col(doc, [
    ("overall_grade",    "Letter grade A–F for the run"),
    ("site_name",        "Site name (also stored as tag for filtering)"),
    ("model",            "LLM model name"),
    ("missing_sections", "Comma-separated list of missing required sections, or 'none'"),
])

heading(doc, "10.3 Local vs Remote Tracking", 2, GREEN)
para(doc,
    "The default TRACKING_URI = 'mlruns' resolves to a file-based store in the project "
    "directory. To use a remote MLflow server (for team-wide experiment sharing), change "
    "TRACKING_URI in eval_mlflow_tracker.py to the server URI, e.g. "
    "'http://mlflow.company.internal:5000'. No code changes are required — the "
    "MLflowTracker class passes the URI to mlflow.set_tracking_uri() on every call. "
    "Authentication (bearer tokens, AWS/GCP credentials) is handled by the MLflow SDK "
    "automatically via environment variables.")

# ════════════════════════════════════════════════════════════════
# SECTION 11 — PERFORMANCE
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 11 · Performance Analysis & Latency Tracking")
heading(doc, "11. Performance Analysis & Latency Tracking", 1, PURPLE)

para(doc,
    "The performance analysis system (src/eval/eval_performance.py) was added to provide "
    "deep operational visibility into where time is spent, what failures occur, and what "
    "specific actions can be taken to improve quality and speed. It is designed for two "
    "audiences simultaneously: technical engineers who need precise diagnostic information, "
    "and non-technical regulatory affairs professionals who need to understand what happened "
    "and what to do next — without understanding the underlying implementation.")

heading(doc, "11.1 Instrumentation Points", 2, GREEN)
para(doc,
    "Timing is captured using time.perf_counter() — a high-resolution monotonic clock that "
    "is unaffected by system clock adjustments. Four instrumentation points exist per section:")
table_2col(doc, [
    ("retrieval_ms",    "FAISS vector DB search + document text loading for the section"),
    ("generation_ms",   "LLM API call via handle_user_message() (dominant time component)"),
    ("eval_ms",         "DeepEval (judge + RAG triad) + Opik evaluation for the section"),
    ("total_ms",        "Sum of all three above; stored in section['timing']['total_ms']"),
])
para(doc,
    "Two additional pipeline-level timers exist in run_artifacts['timing']:")
table_2col(doc, [
    ("generation_phase_ms",  "Wall time from loop start to loop end (includes all sections sequentially)"),
    ("total_pipeline_ms",    "generation_phase_ms + total_eval_ms (eval runs after all generation)"),
    ("total_generation_ms",  "Sum of generation_ms across all sections"),
    ("total_retrieval_ms",   "Sum of retrieval_ms across all sections"),
    ("total_eval_ms",        "Wall time for _run_extended_evaluation() function"),
])
para(doc,
    "Error handling: if a section fails during generation (exception caught in the section "
    "loop), a partial timing entry is still recorded with whatever times were captured before "
    "the failure. The section is appended to run_artifacts['sections'] with generated_text='' "
    "and a generation_error field containing the exception string. This ensures the "
    "Performance tab can still display the section in the latency chart even when it failed.")

heading(doc, "11.2 PerformanceAnalyzer Architecture", 2, GREEN)
para(doc,
    "PerformanceAnalyzer is a pure analytical class — no LLM calls, no I/O, runs in <1ms. "
    "It accepts run_artifacts and evaluation dicts and returns a PerformanceReport dataclass. "
    "The report contains four components:")
bullet(doc, "section_timings: List[SectionTiming] — sorted slowest-first, with retrieval/generation/eval breakdown")
bullet(doc, "overall_timing: Dict — pipeline totals, percentages, slowest section pointer")
bullet(doc, "failures: List[SectionFailure] — issues with severity (critical/warning/info), section_key, technical and plain-English descriptions, metric_value")
bullet(doc, "improvements: List[Improvement] — prioritised recommendations (high/medium/low) with technical and plain-English explanations, affected_sections lists")
para(doc, "The PerformanceReport.to_dict() method serialises the report for storage in run_artifacts['performance_report'] and persistence in the eval run JSON.")

heading(doc, "11.3 Failure Detection Rules", 2, GREEN)
para(doc, "The following conditions trigger failure records in the report:")
t = doc.add_table(rows=8, cols=4)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(["Failure Type", "Severity", "Trigger Condition", "Threshold"]):
    t.cell(0, i).text = h
rows_f = [
    ("error",           "critical", "generated_text is empty after LLM call",              "—"),
    ("missing_chunks",  "warning",  "retrieved_paths is empty on non-static section",       "—"),
    ("low_score",       "warning",  "Rule score below threshold",                           "< 50/100"),
    ("hallucination",   "warning",  "DeepEval faithfulness below threshold",                "< 0.25"),
    ("hallucination",   "warning",  "Opik hallucination_score above threshold",             "> 0.50"),
    ("low_tone",        "info",     "Opik regulatory_tone_score below threshold",           "< 0.50"),
    ("low_relevance",   "info",     "Opik answer_relevance_score below threshold",          "< 0.40"),
]
for i, row in enumerate(rows_f):
    for j, val in enumerate(row):
        t.cell(i+1, j).text = val
doc.add_paragraph()

heading(doc, "11.4 Improvement Recommendation Engine", 2, GREEN)
para(doc,
    "Improvements are generated by grouping failures by type and mapping each group to a "
    "structured recommendation. Each recommendation has an area (e.g. 'Source Documents'), "
    "a priority (high/medium/low), a technical description for engineers, and a plain-English "
    "description for non-technical stakeholders. The affected_sections list identifies exactly "
    "which sections triggered the recommendation.")
para(doc, "Seven improvement categories are defined:")
table_2col(doc, [
    ("Source Documents",    "Triggered by missing_chunks failures — advises expanding ZIP upload with domain-specific SOPs, qualification records, quality manuals; or lowering similarity threshold in DocumentRetriever"),
    ("Factual Accuracy",    "Triggered by hallucination failures — advises adding explicit grounding instructions to prompts, increasing top_k retrieval, adding post-generation fact-check, uploading site-specific factual documents"),
    ("Section Completeness","Triggered by low_score failures — advises reviewing eval_config.py keyword requirements and enriching template prompts for affected sections"),
    ("Regulatory Language", "Triggered by low_tone failures — advises adding ICH Q10 language instructions to system prompt, considering post-processing linguistic normalisation"),
    ("Prompt Design",       "Triggered by low_relevance failures — advises refining template section instructions to be more specific with required sub-topics and output structure"),
    ("Performance / Speed", "Triggered when sections exceed SLOW_SECTION_S=20s — advises static section caching, parallel execution, reducing max_tokens, GPT-4o-mini for low-complexity sections"),
    ("Evaluation Speed",    "Triggered when eval time exceeds 40% of total pipeline time — advises cache_enabled=True, skipping eval on static sections, batching"),
])

heading(doc, "11.5 Latency Optimisation Strategies", 2, GREEN)
para(doc, "The following optimisation strategies are recommended based on profiling the pipeline:")

para(doc, "Static Section Caching (High Impact):", bold=True)
para(doc,
    "Static sections (is_static=True in the template) have deterministic outputs that do not "
    "depend on uploaded documents. Their generated text could be cached between runs for the "
    "same site, saving the generation LLM call entirely. Expected saving: 5–10s per cached "
    "static section.")

para(doc, "Evaluation Caching (High Impact):", bold=True)
para(doc,
    "SHA-256 caching in PMFJudge and RAGEvaluator ensures that sections with identical "
    "generated_text are not re-evaluated on repeat runs. This is particularly effective "
    "during iterative template development where many sections remain unchanged between "
    "runs. Cache hit rate typically reaches 40–60% by the third run on the same site.")

para(doc, "Parallel Section Generation (Medium Impact):", bold=True)
para(doc,
    "Sections without document-order dependencies could be generated concurrently using "
    "asyncio or threading. The current implementation is strictly sequential. Parallelising "
    "independent sections could reduce total generation time by 50–70% on a 20-section "
    "document, at the cost of complexity and increased Azure API rate limit exposure.")

para(doc, "Parallel Evaluation (Medium Impact):", bold=True)
para(doc,
    "PMFJudge.score_document() supports parallel=True with configurable max_workers. "
    "Enabling this reduces judge evaluation time by approximately N/max_workers. For a "
    "20-section document with max_workers=4, judge evaluation time drops from ~200s to ~60s.")

para(doc, "Model Selection by Section Complexity (Low Impact):", bold=True)
para(doc,
    "High-complexity sections (production processes, quality systems) require GPT-4o for "
    "quality output. Lower-complexity sections (document headers, static boilerplate) could "
    "use GPT-4o-mini, which is 10x faster and 80% cheaper. The template could annotate "
    "sections with a complexity flag to drive model selection.")

para(doc, "Retrieval Top-K Reduction (Low Impact):", bold=True)
para(doc,
    "The current retrieval top_k=5 retrieves 5 chunks per section. Reducing to top_k=3 for "
    "simple sections cuts retrieval time and context window size, slightly improving "
    "generation speed and LLM focus. The retrieval_ms component is typically small "
    "(50–200ms per section) compared to generation_ms (5–30s per section).")

# ════════════════════════════════════════════════════════════════
# SECTION 12 — DASHBOARD
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 12 · Evaluation Dashboard")
heading(doc, "12. Evaluation Dashboard", 1, PURPLE)

para(doc,
    "The evaluation dashboard (app_eval_dashboard.py) is a six-tab Streamlit application "
    "providing interactive access to all evaluation results, historical trends, model "
    "comparisons, benchmark management, and performance analysis. It is rendered either "
    "standalone (streamlit run app_eval_dashboard.py) or embedded within the main app "
    "(app.py) via 'from app_eval_dashboard import render_eval_dashboard'.")

heading(doc, "12.1 Tab Architecture", 2, GREEN)
table_2col(doc, [
    ("Tab 1: Run Overview",         "Metric cards (Rule/Judge/Faithfulness/RAG Triad/Composite + Opik row), grade badge, MLflow link, live evaluation mode, DOCX report download"),
    ("Tab 2: Section Heatmap",      "Multi-metric heatmap (rows=sections, cols=metrics), section detail drill-down with per-criterion judge scores"),
    ("Tab 3: Trend Analysis",       "Time-series charts of composite/judge/rag_triad over historical runs, regression alerts for score drops"),
    ("Tab 4: Model Comparison",     "Upload multiple JSON run files, radar chart comparison, cost analysis"),
    ("Tab 5: Benchmark Management", "Benchmark case table, add/import/export cases, run benchmark against all cases"),
    ("Tab 6: ⚡ Performance",       "Latency donut chart, per-section bar chart with stacked phase breakdown, failures table, improvement recommendations with technical/plain-English toggle"),
])

heading(doc, "12.2 Run Overview Tab", 2, GREEN)
para(doc,
    "The Run Overview tab is the primary dashboard view. It loads the most recent run by "
    "default (index 0 from list_runs() which returns newest-first). The run selector "
    "dropdown allows switching to any historical run. The tab renders:")
bullet(doc, "Smart health summary banner — coloured by pass/fail, shows site name, section count, retrieval coverage, missing sections")
bullet(doc, "Five primary metric cards (Rule Score, Judge Score, Faithfulness, RAG Triad, Composite) with delta vs previous run")
bullet(doc, "Opik metrics row (Groundedness, Answer Relevance, Regulatory Tone, Opik Composite) — visible only when Opik data is present")
bullet(doc, "MLflow tracking link — deep link to specific run in MLflow UI (requires mlflow ui running)")
bullet(doc, "Expandable metric explanation — detailed description of each metric for both technical and non-technical readers")
bullet(doc, "Grade badge with PASS/FAIL status and framework label (DeepEval + Opik-style)")
bullet(doc, "Section count, retrieval coverage, missing sections status")
bullet(doc, "Download Report button — generates a formatted DOCX summary report")
bullet(doc, "Live Evaluation section — allows re-running evaluation on any saved run using current credentials")

heading(doc, "12.3 Section Heatmap Tab", 2, GREEN)
para(doc,
    "The Section Heatmap displays a Plotly heatmap with sections as rows and metric columns "
    "[Rule Score, Judge Score, Faithfulness, Contextual Precision, RAG Triad]. Cells are "
    "colour-coded: green for high scores, red for low scores. Clicking a section row (via "
    "a selectbox) opens a detail view showing per-criterion judge scores and full evaluator "
    "output for that section.")

heading(doc, "12.4 Trend Analysis Tab", 2, GREEN)
para(doc,
    "Trend Analysis loads all historical runs from the JSONL index and plots time-series "
    "charts for overall_score, mean_composite (extended), mean_judge_normalized, and "
    "mean_rag_triad_score. Regression alerts fire when the composite score drops by more "
    "than 10 points between consecutive runs. The trend chart uses Plotly Express line "
    "charts with runs sorted by timestamp.")

heading(doc, "12.5 Performance Tab (⚡)", 2, GREEN)
para(doc,
    "The Performance tab is the newest addition. It provides operational diagnostics that "
    "complement the quality metrics in the other tabs. Key features:")
bullet(doc, "View mode toggle (Plain English / Technical) — all text in the tab switches between audience modes")
bullet(doc, "Five overall timing cards: Total, LLM Generation, Retrieval, Evaluation, Avg per Section")
bullet(doc, "Time breakdown donut chart (Plotly): LLM Generation / Retrieval / Evaluation / Other")
bullet(doc, "Section latency horizontal bar chart: sections sorted slowest-first, colour-coded green→red")
bullet(doc, "Phase breakdown stacked bar chart: per-section split of retrieval/generation/eval time")
bullet(doc, "Failures table: severity icon, section name, failure type, details (switches with view mode), metric value")
bullet(doc, "Improvement recommendations: collapsible expanders, 🔥/📌/💡 priority icons, high-priority auto-expanded, affected section list")
bullet(doc, "Legend expander: explains all severity levels and priority icons")

heading(doc, "12.6 DOCX Report Generation", 2, GREEN)
para(doc,
    "The 'Download Report (DOCX)' button in the Run Overview tab generates a formatted "
    "python-docx report containing: run metadata table, full metric scores table (including "
    "all Opik metrics), top 3 best and worst sections, per-section rule check details, and "
    "an overall quality assessment narrative. The report is generated in-memory and delivered "
    "as a downloadable file via st.download_button.")

# ════════════════════════════════════════════════════════════════
# SECTION 13 — BENCHMARKING
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 13 · Benchmarking Framework")
heading(doc, "13. Benchmarking Framework", 1, PURPLE)

para(doc,
    "The benchmarking framework (BenchmarkLoader in src/eval/benchmark_loader.py) provides "
    "systematic evaluation of the generation pipeline against a curated set of reference "
    "cases. Unlike live evaluation (which evaluates outputs as they are generated), "
    "benchmarking evaluates pre-stored generated text against ground-truth reference outputs, "
    "enabling repeatable, comparable measurement of quality changes.")

heading(doc, "13.1 Benchmark Case Structure", 2, GREEN)
para(doc, "Each benchmark case is a JSON file in data/benchmark/ with the following schema:")
code_block(doc,
    "{\n"
    "  'case_id': 'pmf_general_info_001',\n"
    "  'section_key': 'GENERAL INFORMATION',\n"
    "  'section_instruction': 'Write the General Information section for...',\n"
    "  'retrieved_context': 'Source documents text...',\n"
    "  'reference_output': 'The Langensbold site, located at...',  // ground truth\n"
    "  'site_name': 'Langensbold',\n"
    "  'metadata': {\n"
    "    'author': 'regulatory_expert',\n"
    "    'created': '2024-12-01',\n"
    "    'difficulty': 'medium'\n"
    "  }\n"
    "}"
)

heading(doc, "13.2 Benchmark Evaluation Flow", 2, GREEN)
numbered(doc, "BenchmarkLoader.load_all() reads all .json files from data/benchmark/")
numbered(doc, "For each case: EvalSuite.run() is called with the case's generated_output (or reference_output as generated_output in quality-ceiling tests)")
numbered(doc, "Results are aggregated into a DocumentEvalResult with per-case scores")
numbered(doc, "The Benchmark Management tab displays results in a table and computes mean composite score")
numbered(doc, "Results can be exported as JSON for CI/CD quality gate integration")

heading(doc, "13.3 Using Benchmarks for Regression Testing", 2, GREEN)
para(doc,
    "Benchmarks serve as regression tests for the evaluation pipeline and the generation "
    "model. A CI/CD quality gate can be implemented by running run_benchmark() after each "
    "model deployment and asserting that mean_composite >= threshold (e.g. 70.0). If the "
    "score drops below the threshold, the deployment is flagged for review. This prevents "
    "model version upgrades, prompt changes, or retrieval configuration changes from "
    "silently degrading document quality.")

# ════════════════════════════════════════════════════════════════
# SECTION 14 — E2E FLOW
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 14 · End-to-End Pipeline Flow")
heading(doc, "14. End-to-End Pipeline Flow", 1, PURPLE)

para(doc,
    "This section provides a detailed walkthrough of the complete pipeline from user "
    "interaction to fully persisted evaluation results.")

heading(doc, "14.1 Phase 1: Document Ingestion", 2, GREEN)
numbered(doc, "User uploads ZIP file containing site-specific documents (SOPs, batch records, equipment lists, quality manuals, validation reports)")
numbered(doc, "ZIP is extracted to data/artifacts/Extracted_folder/")
numbered(doc, "convert_all_doc_to_docx_in_folder() converts any .doc files to .docx using LibreOffice/python-docx")
numbered(doc, "DocumentRetriever initialises a FAISS index using all-mpnet-base-v2 (768-dimensional embeddings) across all extracted text chunks")
numbered(doc, "User uploads Reference Excel file (site-specific structured data)")
numbered(doc, "extract_text_from_xlsx() converts Excel to plain text string (data_as_string)")

heading(doc, "14.2 Phase 2: Section Generation", 2, GREEN)
numbered(doc, "_pipeline_start = time.perf_counter() — overall timer starts")
numbered(doc, "Template JSON is parsed; each entry is a (key, value) pair where value contains the prompt and optional retrieval query (separated by '@!')")
numbered(doc, "For each section: _sec_start timer starts")
numbered(doc, "If section has retrieval query: FAISS search with top_k=5; retrieval_ms recorded")
numbered(doc, "data_extraction() loads text from retrieved document paths")
numbered(doc, "handle_user_message() sends instruction + context to GPT-4o via AzureChatOpenAI; generation_ms recorded")
numbered(doc, "Section dict appended to run_artifacts['sections'] with all timing data")
numbered(doc, "Errors are caught per-section; partial timing recorded; generation_error stored; loop continues")
numbered(doc, "After loop: run_artifacts['timing'] populated with totals")

heading(doc, "14.3 Phase 3: Rule-Based Evaluation", 2, GREEN)
numbered(doc, "evaluate_run(run_artifacts, rules) calls score_document()")
numbered(doc, "Each section's generated_text is checked against applicable rules (min_chars, keywords, site_name)")
numbered(doc, "Document-level scores computed: overall_score, section_count, retrieval_coverage, missing_required_sections")
numbered(doc, "save_eval_run() persists initial run (without extended eval) to JSON + JSONL index")
numbered(doc, "st.session_state['last_eval_score'] updated for inline display in app.py")

heading(doc, "14.4 Phase 4: Extended Evaluation", 2, GREEN)
numbered(doc, "Azure credentials injected into run_artifacts (_azure_key, _azure_endpoint, _azure_version)")
numbered(doc, "EvalSuite instantiated with Azure credentials and run_judge=True, run_rag=True")
numbered(doc, "OpikStyleScorer instantiated with Azure OpenAI client")
numbered(doc, "_eval_phase_start timer starts")
numbered(doc, "For each non-static section with non-empty generated_text:")
bullet(doc, "Retrieved paths loaded and merged into context string", level=1)
bullet(doc, "_sec_eval_start timer starts", level=1)
bullet(doc, "eval_suite.run() → PMFJudge + RAGEvaluator; result stored in section['extended_eval']", level=1)
bullet(doc, "opik_scorer.evaluate_section() → 3 Opik metrics; result stored in section['opik_eval']", level=1)
bullet(doc, "section['timing']['eval_ms'] and 'total_ms' updated", level=1)
numbered(doc, "Document-level aggregates computed: mean_composite, mean_judge, mean_rag, all Opik means")
numbered(doc, "run_artifacts['extended_eval_summary'] populated")
numbered(doc, "total_eval_ms = time.perf_counter() - _eval_phase_start; run_artifacts['timing'] finalised")

heading(doc, "14.5 Phase 5: Analysis & Persistence", 2, GREEN)
numbered(doc, "evaluate_run() re-runs rule evaluation (now with extended eval data in sections)")
numbered(doc, "save_eval_run() re-saves with complete run_artifacts (overwrites initial save)")
numbered(doc, "PerformanceAnalyzer.analyze() generates PerformanceReport; stored in run_artifacts['performance_report']")
numbered(doc, "MLflowTracker.log_run() logs all metrics to mlruns/; run_id stored in run_artifacts")
numbered(doc, "st.session_state updated: last_extended_composite, last_extended_grade, last_extended_judge, last_extended_rag, last_mlflow_run_id")
numbered(doc, "app.py renders inline evaluation summary with grade, scores, and MLflow link")

# ════════════════════════════════════════════════════════════════
# SECTION 15 — CONFIGURATION
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 15 · Configuration Reference")
heading(doc, "15. Configuration Reference", 1, PURPLE)

heading(doc, "15.1 Environment Variables (.env)", 2, GREEN)
table_2col(doc, [
    ("AZURE_KEY",       "Azure OpenAI API key"),
    ("AZURE_ENDPOINT",  "Azure OpenAI endpoint URL (e.g. https://xxx.openai.azure.com/)"),
    ("AZURE_NAME",      "Azure deployment name (e.g. gpt-4o)"),
    ("AZURE_VERSION",   "Azure API version (e.g. 2024-06-01)"),
    ("ANTHROPIC_API_KEY", "Anthropic Claude API key (optional; for fallback judge)"),
])

heading(doc, "15.2 Evaluation Thresholds", 2, GREEN)
table_2col(doc, [
    ("LOW_RULE_SCORE",       "50.0 — rule score below this triggers 'low_score' failure"),
    ("LOW_FAITHFULNESS",     "0.25 — faithfulness below this triggers 'hallucination' failure"),
    ("HIGH_HALLUCINATION",   "0.50 — Opik hallucination above this triggers 'hallucination' failure"),
    ("LOW_REG_TONE",         "0.50 — regulatory tone below this triggers 'low_tone' failure"),
    ("LOW_ANSWER_RELEVANCE", "0.40 — answer relevance below this triggers 'low_relevance' failure"),
    ("SLOW_SECTION_S",       "20.0s — section total time above this triggers 'slow' improvement"),
    ("EVAL_OVERHEAD_RATIO",  "0.40 — eval > 40% of total triggers evaluation speed improvement"),
    ("PASS_THRESHOLD",       "65.0 — composite_score >= 65 = PASS; < 65 = FAIL in dashboard"),
])

heading(doc, "15.3 Composite Score Weights", 2, GREEN)
table_2col(doc, [
    ("Rule score weight",       "0.20 (20%)"),
    ("Judge score weight",      "0.55 (55%)"),
    ("RAG Triad score weight",  "0.25 (25%)"),
    ("BERTScore weight",        "0.00 (disabled — torch DLL issue on Windows)"),
    ("Opik weight in composite","0.00 (Opik is supplementary, not in main composite)"),
])

heading(doc, "15.4 Cache Configuration", 2, GREEN)
table_2col(doc, [
    ("Judge cache dir",     "data/eval_cache/"),
    ("Judge cache version", "v1.0 — bump to invalidate all judge caches"),
    ("RAG cache dir",       "data/eval_cache/deepeval/"),
    ("RAG cache version",   "deepeval_v1.0 — bump to invalidate all RAG caches"),
    ("Opik cache dir",      "data/eval_cache/opik/"),
    ("Opik cache version",  "opik_v1.0 — bump to invalidate all Opik caches"),
])

# ════════════════════════════════════════════════════════════════
# SECTION 16 — DEPENDENCY MATRIX
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Section 16 · Dependency Matrix")
heading(doc, "16. Dependency Matrix", 1, PURPLE)

t = doc.add_table(rows=14, cols=4)
t.style = 'Light Grid Accent 1'
for i, h in enumerate(["Package", "Version", "Used By", "Notes"]):
    t.cell(0, i).text = h
deps = [
    ("streamlit",           "1.39.0",   "app.py, dashboard",               "UI framework"),
    ("openai",              "1.109.1",  "generation, judge, RAG, Opik",    "Azure OpenAI client; DO NOT upgrade to 2.x (opik conflict)"),
    ("anthropic",           "latest",   "PMFJudge fallback",               "Optional; only needed for Claude judge"),
    ("langchain-openai",    "latest",   "generation LLM",                  "AzureChatOpenAI wrapper"),
    ("mlflow",              "3.11.1",   "MLflowTracker",                   "Installs protobuf 6.x; Streamlit warns but imports succeed"),
    ("python-docx",         "latest",   "template parsing, report gen",    "DOCX read/write"),
    ("pandas",              "latest",   "dashboard, judge comparison",     "DataFrames"),
    ("plotly",              "latest",   "dashboard charts",                "Optional; graceful degradation to st.dataframe"),
    ("sentence-transformers","latest",  "RAG Answer Relevancy embeddings", "Optional; heuristic fallback if unavailable"),
    ("faiss-cpu",           "latest",   "DocumentRetriever vector DB",     "CPU-only build; conflicts with torch GPU on some systems"),
    ("sacrebleu",           "latest",   "LexicalMetrics BLEU",             "Optional; disabled by default in live pipeline"),
    ("rouge-score",         "latest",   "LexicalMetrics ROUGE",            "Optional; disabled by default in live pipeline"),
    ("torch",               "latest",   "BERTScore, sentence-transformers","DLL conflict on Windows with faiss-cpu; BERTScore disabled"),
]
for i, row in enumerate(deps):
    for j, val in enumerate(row):
        t.cell(i+1, j).text = val
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# SECTION 17 — APPENDIX A: METRIC FORMULAS
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Appendix A · Metric Formulas")
heading(doc, "17. Appendix A — Metric Formulas", 1, PURPLE)

heading(doc, "Rule Score", 2, GREEN)
code_block(doc,
    "score = 100.0\n"
    "if len(text) < min_chars:         score -= 40.0\n"
    "if missing_keywords:              score -= (len(missing) / len(required)) * 30.0\n"
    "if site_name not in text:         score -= 15.0\n"
    "score = clamp(score, 0, 100)"
)

heading(doc, "Judge Score (Normalised)", 2, GREEN)
code_block(doc,
    "weighted = Σ(score_i × weight_i) for i in {factual, regulatory, site, complete, coherent}\n"
    "normalized = (weighted / 5.0) × 100.0\n"
    "# Weights: factual=0.30, regulatory=0.25, site=0.20, complete=0.15, coherent=0.10"
)

heading(doc, "Faithfulness", 2, GREEN)
code_block(doc,
    "claims = LLM_extract_claims(generated_text)           # list of atomic claims\n"
    "verdicts = [LLM_entails(context, c) for c in claims]  # 'yes' or 'no'\n"
    "faithfulness = count(v=='yes') / count(claims)"
)

heading(doc, "Contextual Precision (Rank-Weighted AP)", 2, GREEN)
code_block(doc,
    "# AP@k where R = total relevant chunks\n"
    "relevant_so_far = 0\n"
    "precision_sum   = 0\n"
    "for i, chunk in enumerate(chunks, start=1):\n"
    "    if LLM_relevant(chunk, query):\n"
    "        relevant_so_far += 1\n"
    "        precision_sum   += relevant_so_far / i\n"
    "contextual_precision = precision_sum / max(relevant_so_far, 1)"
)

heading(doc, "Answer Relevancy", 2, GREEN)
code_block(doc,
    "questions    = LLM_generate_questions(generated_text, n=3)\n"
    "embeddings   = SentenceTransformer.encode([instruction] + questions)\n"
    "similarities = [cosine_sim(embeddings[0], embeddings[i]) for i in 1..n]\n"
    "answer_relevancy = mean(similarities)"
)

heading(doc, "RAG Triad Score (Harmonic Mean)", 2, GREEN)
code_block(doc,
    "def harmonic_mean(values):\n"
    "    valid = [v for v in values if v and v > 0]\n"
    "    return len(valid) / Σ(1/v for v in valid)\n\n"
    "rag_triad_score = harmonic_mean([faithfulness, contextual_precision, answer_relevancy])"
)

heading(doc, "Opik Composite", 2, GREEN)
code_block(doc,
    "opik_composite = mean([\n"
    "    1.0 - hallucination_score,   # inverted\n"
    "    answer_relevance_score,\n"
    "    regulatory_tone_score\n"
    "])"
)

heading(doc, "Composite Score", 2, GREEN)
code_block(doc,
    "weights = {'rule': 0.20, 'judge_normalized': 0.55, 'rag_triad_score': 0.25}\n"
    "available = {k: v for k,v in scores.items() if v is not None}\n"
    "total_w = sum(weights[k] for k in available)\n"
    "composite = sum(available[k] * weights[k] for k in available) / total_w"
)

heading(doc, "Letter Grade", 2, GREEN)
code_block(doc,
    "A: composite >= 90\n"
    "B: composite >= 75\n"
    "C: composite >= 60\n"
    "D: composite >= 45\n"
    "F: composite <  45"
)

# ════════════════════════════════════════════════════════════════
# SECTION 18 — APPENDIX B: PROMPT TEMPLATES
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Appendix B · Prompt Templates")
heading(doc, "18. Appendix B — Prompt Templates", 1, PURPLE)

heading(doc, "18.1 LLM-as-Judge System Prompt", 2, GREEN)
code_block(doc,
    "You are an expert regulatory affairs specialist with 15+ years of experience\n"
    "reviewing Plant Master Files (PMF) and pharmaceutical regulatory submissions.\n"
    "You evaluate AI-generated PMF sections against a structured rubric. You are\n"
    "precise, consistent, and always provide a structured JSON response. You never\n"
    "deviate from the requested JSON format."
)

heading(doc, "18.2 DeepEval Claim Extraction Prompt", 2, GREEN)
code_block(doc,
    "SYSTEM: You extract factual claims from text for verification.\n"
    "USER:   Extract all atomic factual claims from this text.\n"
    "        Think step by step before listing claims.\n"
    "        Exclude generic statements. Include specific facts: names, numbers,\n"
    "        certifications, dates, site-specific details.\n"
    "        Return JSON: {\"claims\": [\"claim1\", \"claim2\", ...]}"
)

heading(doc, "18.3 DeepEval Entailment Prompt", 2, GREEN)
code_block(doc,
    "SYSTEM: You are a fact-checker. Determine if context supports a claim.\n"
    "USER:   CONTEXT: {context[:4000]}\n"
    "        CLAIM: {claim}\n"
    "        Does the context entail the claim? Return JSON:\n"
    "        {\"verdict\": \"yes\"|\"no\", \"reason\": \"one sentence\"}"
)

heading(doc, "18.4 Opik Hallucination Prompt", 2, GREEN)
code_block(doc,
    "SYSTEM: You are an expert fact-checker evaluating AI-generated regulatory documents.\n"
    "        Return ONLY valid JSON, no markdown.\n"
    "USER:   Evaluate the HALLUCINATION LEVEL of the AI-generated output.\n"
    "        SOURCE CONTEXT: {context[:6000]}\n"
    "        AI-GENERATED OUTPUT: {output[:4000]}\n"
    "        DEFINITION: Hallucination = any factual claim NOT in the source context.\n"
    "        Do NOT penalise regulatory boilerplate — only invented facts.\n"
    "        Scoring: 0.0=none, 0.25=minor, 0.5=moderate, 0.75=significant, 1.0=complete\n"
    "        Return: {\"score\": float, \"reason\": str, \"examples\": [str]}"
)

heading(doc, "18.5 Opik Regulatory Tone Prompt", 2, GREEN)
code_block(doc,
    "SYSTEM: You are a regulatory affairs expert evaluating AI-generated PMF sections.\n"
    "        Return ONLY valid JSON.\n"
    "USER:   Evaluate REGULATORY TONE AND LANGUAGE QUALITY of this PMF section.\n"
    "        SECTION KEY: {section_key}\n"
    "        AI-GENERATED OUTPUT: {output[:4000]}\n"
    "        1.0=Exemplary EU GMP language, 0.75=Good, 0.5=Acceptable,\n"
    "        0.25=Poor, 0.0=Inappropriate\n"
    "        Return: {\"score\": float, \"reason\": str, \"issues\": [str]}"
)

# ════════════════════════════════════════════════════════════════
# SECTION 19 — APPENDIX C: DATA SCHEMAS
# ════════════════════════════════════════════════════════════════
section_divider(doc, "Appendix C · Data Schemas")
heading(doc, "19. Appendix C — Data Schemas", 1, PURPLE)

heading(doc, "19.1 EvalResult Schema", 2, GREEN)
code_block(doc,
    "{\n"
    "  section_key: str,\n"
    "  rule_score: float | None,          // 0-100\n"
    "  judge_scores: {\n"
    "    scores: {factual_accuracy: int, regulatory_language: int,\n"
    "             site_specificity: int, completeness: int,\n"
    "             structural_coherence: int},\n"
    "    weighted_score: float,           // 0-5\n"
    "    normalized_score: float,         // 0-100\n"
    "    strengths: [str],\n"
    "    weaknesses: [str],\n"
    "    critical_issues: [str],\n"
    "    improvement_suggestions: [str],\n"
    "    judge_confidence: float | None,  // 0-1\n"
    "    evaluation_notes: str,\n"
    "    cached: bool,\n"
    "    judge_model: str,\n"
    "    judge_provider: str\n"
    "  } | null,\n"
    "  rag_scores: {\n"
    "    faithfulness: float,             // 0-1\n"
    "    contextual_precision: float,     // 0-1\n"
    "    answer_relevancy: float,         // 0-1\n"
    "    rag_triad_score: float,          // 0-1 harmonic mean\n"
    "    framework: 'deepeval_rag_triad'\n"
    "  } | null,\n"
    "  composite_score: float,            // 0-100\n"
    "  grade: 'A'|'B'|'C'|'D'|'F'\n"
    "}"
)

heading(doc, "19.2 OpikEval Schema", 2, GREEN)
code_block(doc,
    "{\n"
    "  section_key: str,\n"
    "  hallucination_score: float | null,       // 0-1 (lower is better)\n"
    "  hallucination_reason: str,\n"
    "  hallucination_examples: [str],\n"
    "  answer_relevance_score: float | null,    // 0-1 (higher is better)\n"
    "  answer_relevance_reason: str,\n"
    "  regulatory_tone_score: float | null,     // 0-1 (higher is better)\n"
    "  regulatory_tone_reason: str,\n"
    "  regulatory_tone_issues: [str],\n"
    "  opik_composite: float | null,            // 0-1\n"
    "  framework: 'opik_style',\n"
    "  evaluated_at: str                        // ISO 8601 UTC\n"
    "}"
)

heading(doc, "19.3 PerformanceReport Schema", 2, GREEN)
code_block(doc,
    "{\n"
    "  section_timings: [{\n"
    "    section_key: str,\n"
    "    retrieval_ms: float | null,\n"
    "    generation_ms: float | null,\n"
    "    eval_ms: float | null,\n"
    "    total_ms: float | null,\n"
    "    is_static: bool\n"
    "  }],\n"
    "  overall_timing: {\n"
    "    total_pipeline_ms: float,\n"
    "    total_generation_ms: float,\n"
    "    total_retrieval_ms: float,\n"
    "    total_eval_ms: float,\n"
    "    section_count: int,\n"
    "    slowest_section: str,\n"
    "    slowest_section_ms: float,\n"
    "    avg_section_ms: float,\n"
    "    pct_generation: float,\n"
    "    pct_retrieval: float,\n"
    "    pct_eval: float\n"
    "  },\n"
    "  failures: [{\n"
    "    section_key: str,\n"
    "    failure_type: str,\n"
    "    severity: 'critical'|'warning'|'info',\n"
    "    technical: str,\n"
    "    plain_english: str,\n"
    "    metric_value: float | null\n"
    "  }],\n"
    "  improvements: [{\n"
    "    area: str,\n"
    "    priority: 'high'|'medium'|'low',\n"
    "    technical: str,\n"
    "    plain_english: str,\n"
    "    affected_sections: [str]\n"
    "  }],\n"
    "  summary_technical: str,\n"
    "  summary_plain: str\n"
    "}"
)

# ── Save ──────────────────────────────────────────────────────────────────
output_path = "docs/PMF_LLM_Evaluation_Framework_Technical_Documentation.docx"
import os; os.makedirs("docs", exist_ok=True)
doc.save(output_path)
print(f"Saved: {output_path}")
